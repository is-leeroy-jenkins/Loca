'''
	******************************************************************************************
	    Assembly:                Loca LLama
	    Filename:                app.py
	    Author:                  Terry D. Eppler
	    Created:                 05-31-2024
	
	    Last Modified By:        Terry D. Eppler
	    Last Modified On:        05-01-2025
	******************************************************************************************
	<copyright file="app.py" company="Terry D. Eppler">
	
	           Loca is python application for running local LLMs.
	           Copyright ©  2023 Terry Eppler
	
	   Permission is hereby granted, free of charge, to any person obtaining a copy
	   of this software and associated documentation files (the “Software”),
	   to deal in the Software without restriction,
	   including without limitation the rights to use,
	   copy, modify, merge, publish, distribute, sublicense,
	   and/or sell copies of the Software,
	   and to permit persons to whom the Software is furnished to do so,
	   subject to the following conditions:
	
	   The above copyright notice and this permission notice shall be included in all
	   copies or substantial portions of the Software.
	
	   THE SOFTWARE IS PROVIDED “AS IS”, WITHOUT WARRANTY OF ANY KIND, EXPRESS OR IMPLIED,
	   INCLUDING BUT NOT LIMITED TO THE WARRANTIES OF MERCHANTABILITY,
	   FITNESS FOR A PARTICULAR PURPOSE AND NON-INFRINGEMENT.
	   IN NO EVENT SHALL THE AUTHORS OR COPYRIGHT HOLDERS BE LIABLE FOR ANY CLAIM,
	   DAMAGES OR OTHER LIABILITY, WHETHER IN AN ACTION OF CONTRACT, TORT OR OTHERWISE,
	   ARISING FROM, OUT OF OR IN CONNECTION WITH THE SOFTWARE OR THE USE OR OTHER
	   DEALINGS IN THE SOFTWARE.
	
	   You can contact me at:  terryeppler@gmail.com or eppler.terry@epa.gov
	
	</copyright>
	<summary>
	  app.py
	</summary>
	******************************************************************************************
'''
from __future__ import annotations

import base64
import hashlib
import re
import sqlite3
import time
from io import BytesIO
from pathlib import Path
from typing import Any, Dict, List, Tuple

import numpy as np
import pandas as pd
import plotly.express as px
import streamlit as st

import config as cfg

try:
	from docx import Document
except ImportError:
	Document = None

try:
	from llama_cpp import Llama
except ImportError:
	Llama = None

try:
	import fitz
except ImportError:
	fitz = None

# ==============================================================================
# STREAMLIT BOOTSTRAP
# ==============================================================================

st.set_page_config( page_title=cfg.APP_TITLE, layout='wide',
	page_icon=cfg.FAVICON )

st.caption( cfg.APP_SUBTITLE )

def is_docx_available( ) -> bool:
	"""
		Purpose:
		--------
		Determine whether python-docx is available for DOCX extraction.

		Parameters:
		-----------
		None

		Returns:
		--------
		bool
			True when python-docx is available; otherwise False.
	"""
	return Document is not None

def is_llama_cpp_available( ) -> bool:
	"""
		Purpose:
		--------
		Determine whether llama-cpp-python is available for local GGUF inference.

		Parameters:
		-----------
		None

		Returns:
		--------
		bool
			True when llama-cpp-python is available; otherwise False.
	"""
	return Llama is not None

def is_pymupdf_available( ) -> bool:
	"""
		Purpose:
		--------
		Determine whether PyMuPDF is available for native PDF text extraction.

		Parameters:
		-----------
		None

		Returns:
		--------
		bool
			True when PyMuPDF is available; otherwise False.
	"""
	return fitz is not None

def get_selected_model_name( ) -> str:
	"""
		Purpose:
		--------
		Return the currently selected local model name from Streamlit session state.

		Parameters:
		-----------
		None

		Returns:
		--------
		str
			Selected model name.
	"""
	model_name = str(
		st.session_state.get( 'selected_model_name', get_default_model_name( ) ) or
		get_default_model_name( ) )
	
	return model_name

def get_selected_model_path( ) -> str:
	"""
		Purpose:
		--------
		Return the currently selected local GGUF model path from Streamlit session state.

		Parameters:
		-----------
		None

		Returns:
		--------
		str
			Resolved local GGUF path for the selected model.
	"""
	model_name = get_selected_model_name( )
	model_path = str(
		st.session_state.get( 'selected_model_path', get_model_path_for_state( model_name ) ) or
		get_model_path_for_state( model_name ) )
	
	return model_path

def get_selected_model_spec( ) -> Dict[ str, Any ]:
	"""
		Purpose:
		--------
		Return the selected model specification from Streamlit session state.

		Parameters:
		-----------
		None

		Returns:
		--------
		Dict[str, Any]
			Selected model metadata.
	"""
	model_name = get_selected_model_name( )
	model_spec = st.session_state.get( 'selected_model_spec', None )
	
	if isinstance( model_spec, dict ) and len( model_spec ) > 0:
		return model_spec
	
	return get_model_spec_for_state( model_name )

def local_model_available( model_path: str | None = None ) -> bool:
	"""
		Purpose:
		--------
		Determine whether the selected or supplied local GGUF model file exists.

		Parameters:
		-----------
		model_path : str | None
			Optional GGUF model path. When omitted, the selected model path is used.

		Returns:
		--------
		bool
			True when the configured model file exists; otherwise False.
	"""
	try:
		path_value = str( model_path or get_selected_model_path( ) or '' ).strip( )
		
		if not path_value:
			return False
		
		return Path( path_value ).exists( )
	except Exception:
		return False
	
# ==============================================================================
# SESSION STATE INITIALIZATION
# ==============================================================================
if 'mode' not in st.session_state:
	st.session_state[ 'mode' ] = ''

if 'messages' not in st.session_state:
	st.session_state[ 'messages' ] = [ ]

if 'system_instructions' not in st.session_state:
	st.session_state[ 'system_instructions' ] = ''

if 'context_window' not in st.session_state:
	st.session_state[ 'context_window' ] = 0

if 'cpu_threads' not in st.session_state:
	st.session_state[ 'cpu_threads' ] = 0

if 'max_tokens' not in st.session_state:
	st.session_state[ 'max_tokens' ] = 0

if 'temperature' not in st.session_state:
	st.session_state[ 'temperature' ] = 0.0

if 'top_percent' not in st.session_state:
	st.session_state[ 'top_percent' ] = 0.0

if 'top_k' not in st.session_state:
	st.session_state[ 'top_k' ] = 0

if 'frequency_penalty' not in st.session_state:
	st.session_state[ 'frequency_penalty' ] = 0.0

if 'presense_penalty' not in st.session_state:
	st.session_state[ 'presense_penalty' ] = 0.0

if 'repeat_penalty' not in st.session_state:
	st.session_state[ 'repeat_penalty' ] = 0.0

if 'repeat_window' not in st.session_state:
	st.session_state[ 'repeat_window' ] = 0

if 'random_seed' not in st.session_state:
	st.session_state[ 'random_seed' ] = 0

if 'basic_docs' not in st.session_state:
	st.session_state[ 'basic_docs' ] = [ ]

if 'use_semantic' not in st.session_state:
	st.session_state[ 'use_semantic' ] = False

if 'is_grounded' not in st.session_state:
	st.session_state[ 'is_grounded' ] = False

if 'selected_prompt_id' not in st.session_state:
	st.session_state[ 'selected_prompt_id' ] = ''

if 'pending_system_prompt_name' not in st.session_state:
	st.session_state[ 'pending_system_prompt_name' ] = ''
	
# -------- TEXT GENERATION  ---------------------

if 'task_preset' not in st.session_state:
	st.session_state[ 'task_preset' ] = 'Chat'

if 'response_format' not in st.session_state:
	st.session_state[ 'response_format' ] = 'Markdown'

if 'use_chat_history' not in st.session_state:
	st.session_state[ 'use_chat_history' ] = True

if 'use_document_context' not in st.session_state:
	st.session_state[ 'use_document_context' ] = False

if 'reasoning_depth' not in st.session_state:
	st.session_state[ 'reasoning_depth' ] = 'Medium'

if 'answer_only' not in st.session_state:
	st.session_state[ 'answer_only' ] = False

if 'use_self_check' not in st.session_state:
	st.session_state[ 'use_self_check' ] = False

if 'deterministic_reasoning' not in st.session_state:
	st.session_state[ 'deterministic_reasoning' ] = False

if 'coding_language' not in st.session_state:
	st.session_state[ 'coding_language' ] = 'Python'

if 'coding_task' not in st.session_state:
	st.session_state[ 'coding_task' ] = 'Generate'

if 'coding_include_comments' not in st.session_state:
	st.session_state[ 'coding_include_comments' ] = True

if 'coding_editor_format' not in st.session_state:
	st.session_state[ 'coding_editor_format' ] = True

if 'coding_fenced_output' not in st.session_state:
	st.session_state[ 'coding_fenced_output' ] = True

if 'translation_target_language' not in st.session_state:
	st.session_state[ 'translation_target_language' ] = 'English'

if 'active_prompt_caption' not in st.session_state:
	st.session_state[ 'active_prompt_caption' ] = ''

if 'preview_effective_prompt' not in st.session_state:
	st.session_state[ 'preview_effective_prompt' ] = False

if 'last_preview_input' not in st.session_state:
	st.session_state[ 'last_preview_input' ] = ''
	
#-------- DOCQNA ---------------------

if 'uploaded' not in st.session_state:
	st.session_state[ 'uploaded' ] = [ ]

if 'active_docs' not in st.session_state:
	st.session_state[ 'active_docs' ] = [ ]

if 'doc_bytes' not in st.session_state:
	st.session_state[ 'doc_bytes' ] = { }
	
if 'doc_source' not in st.session_state:
	st.session_state[ 'doc_source' ] = 'uploadlocal'

if 'docqna_vec_ready' not in st.session_state:
	st.session_state[ 'docqna_vec_ready' ] = False

if 'docqna_fingerprint' not in st.session_state:
	st.session_state[ 'docqna_fingerprint' ] = ''

if 'docqna_chunk_count' not in st.session_state:
	st.session_state[ 'docqna_chunk_count' ] = 0
	
if 'docqna_fallback_rows' not in st.session_state:
	st.session_state[ 'docqna_fallback_rows' ] = [ ]
	
# -------- DOCUMENT Q&A EXTENSIONS ---------------------

if 'retrieval_k' not in st.session_state:
	st.session_state[ 'retrieval_k' ] = 6

if 'retrieval_chunk_size' not in st.session_state:
	st.session_state[ 'retrieval_chunk_size' ] = 1200

if 'retrieval_chunk_overlap' not in st.session_state:
	st.session_state[ 'retrieval_chunk_overlap' ] = 200

if 'show_retrieved_chunks' not in st.session_state:
	st.session_state[ 'show_retrieved_chunks' ] = True

if 'require_grounding' not in st.session_state:
	st.session_state[ 'require_grounding' ] = True

if 'answer_from_excerpts_only' not in st.session_state:
	st.session_state[ 'answer_from_excerpts_only' ] = True

if 'prefer_sqlite_vec' not in st.session_state:
	st.session_state[ 'prefer_sqlite_vec' ] = True

if 'allow_similarity_fallback' not in st.session_state:
	st.session_state[ 'allow_similarity_fallback' ] = True

if 'docqna_action' not in st.session_state:
	st.session_state[ 'docqna_action' ] = 'Answer Question'

if 'ocr_enabled' not in st.session_state:
	st.session_state[ 'ocr_enabled' ] = False

if 'prefer_native_pdf_text' not in st.session_state:
	st.session_state[ 'prefer_native_pdf_text' ] = True

if 'include_page_markers' not in st.session_state:
	st.session_state[ 'include_page_markers' ] = False

if 'show_docqna_diagnostics' not in st.session_state:
	st.session_state[ 'show_docqna_diagnostics' ] = False

if 'docqna_last_retrieval' not in st.session_state:
	st.session_state[ 'docqna_last_retrieval' ] = [ ]

if 'docqna_inventory_rows' not in st.session_state:
	st.session_state[ 'docqna_inventory_rows' ] = [ ]

# -------- SEMANTIC SEARCH ---------------------

if 'semantic_context_buffer' not in st.session_state:
	st.session_state[ 'semantic_context_buffer' ] = [ ]

if 'semantic_chunk_size' not in st.session_state:
	st.session_state[ 'semantic_chunk_size' ] = 1200

if 'semantic_chunk_overlap' not in st.session_state:
	st.session_state[ 'semantic_chunk_overlap' ] = 200

if 'semantic_top_k' not in st.session_state:
	st.session_state[ 'semantic_top_k' ] = 8

if 'semantic_min_similarity' not in st.session_state:
	st.session_state[ 'semantic_min_similarity' ] = 0.0

if 'semantic_group_by_document' not in st.session_state:
	st.session_state[ 'semantic_group_by_document' ] = False

if 'semantic_clear_existing' not in st.session_state:
	st.session_state[ 'semantic_clear_existing' ] = True

if 'semantic_append_existing' not in st.session_state:
	st.session_state[ 'semantic_append_existing' ] = False

if 'semantic_show_diagnostics' not in st.session_state:
	st.session_state[ 'semantic_show_diagnostics' ] = True

if 'semantic_uploaded_names' not in st.session_state:
	st.session_state[ 'semantic_uploaded_names' ] = [ ]

if 'semantic_result_rows' not in st.session_state:
	st.session_state[ 'semantic_result_rows' ] = [ ]

if 'semantic_selected_rows' not in st.session_state:
	st.session_state[ 'semantic_selected_rows' ] = [ ]

if 'semantic_index_chunk_count' not in st.session_state:
	st.session_state[ 'semantic_index_chunk_count' ] = 0

if 'semantic_index_dim' not in st.session_state:
	st.session_state[ 'semantic_index_dim' ] = 0

if 'semantic_index_doc_count' not in st.session_state:
	st.session_state[ 'semantic_index_doc_count' ] = 0

if 'semantic_last_query' not in st.session_state:
	st.session_state[ 'semantic_last_query' ] = ''
	
# -------- PROMPT ENGINEERING EXTENSIONS ---------------------

if 'prompt_category' not in st.session_state:
	st.session_state[ 'prompt_category' ] = 'General Chat'

if 'prompt_task' not in st.session_state:
	st.session_state[ 'prompt_task' ] = 'Chat'

if 'prompt_response_format' not in st.session_state:
	st.session_state[ 'prompt_response_format' ] = 'Markdown'

if 'pe_language' not in st.session_state:
	st.session_state[ 'pe_language' ] = 'English'

if 'pe_generator_goal' not in st.session_state:
	st.session_state[ 'pe_generator_goal' ] = ''

if 'pe_generator_constraints' not in st.session_state:
	st.session_state[ 'pe_generator_constraints' ] = ''

if 'pe_generator_style' not in st.session_state:
	st.session_state[ 'pe_generator_style' ] = 'Practical'

if 'pe_generated_template' not in st.session_state:
	st.session_state[ 'pe_generated_template' ] = ''
	
# -------- DATABASE  ---------------------

if 'dm_asset_sync_status' not in st.session_state:
	st.session_state[ 'dm_asset_sync_status' ] = ''

if 'dm_asset_counts' not in st.session_state:
	st.session_state[ 'dm_asset_counts' ] = { }

if 'dm_selected_asset_table' not in st.session_state:
	st.session_state[ 'dm_selected_asset_table' ] = 'documents'

if 'dm_register_uploaded_docs' not in st.session_state:
	st.session_state[ 'dm_register_uploaded_docs' ] = False

if 'dm_register_uploaded_images' not in st.session_state:
	st.session_state[ 'dm_register_uploaded_images' ] = False

# ==============================================================================
# MODEL / MODE SESSION STATE CONTRACT
# ==============================================================================

def get_default_model_name( ) -> str:
	"""
		Purpose:
		--------
		Return the configured default local model name from config.

		Parameters:
		-----------
		None

		Returns:
		--------
		str
			Default model name.
	"""
	default_model = str( getattr( cfg, 'DEFAULT_MODEL', '' ) or '' ).strip( )
	
	if default_model:
		return default_model
	
	if hasattr( cfg, 'get_model_names' ):
		model_names = cfg.get_model_names( )
	else:
		model_names = list( getattr( cfg, 'MODEL_MAP', { } ).keys( ) )
	
	return str( model_names[ 0 ] ) if model_names else ''

def get_model_names_for_state( ) -> List[ str ]:
	"""
		Purpose:
		--------
		Return configured model names from the config registry while preserving fallback
		compatibility with cfg.MODEL_MAP.

		Parameters:
		-----------
		None

		Returns:
		--------
		List[str]
			Configured model names.
	"""
	if hasattr( cfg, 'get_model_names' ):
		model_names = cfg.get_model_names( )
	else:
		model_names = list( getattr( cfg, 'MODEL_MAP', { } ).keys( ) )
	
	return [ str( name ) for name in model_names ]

def get_default_mode_name( model_name: str = '' ) -> str:
	"""
		Purpose:
		--------
		Return the default UI mode for the selected model.

		Parameters:
		-----------
		model_name : str
			Selected local model name.

		Returns:
		--------
		str
			Default UI mode name.
	"""
	model_value = str( model_name or get_default_model_name( ) ).strip( )
	
	if hasattr( cfg, 'get_model_modes' ):
		modes = cfg.get_model_modes( model_value )
	else:
		modes = getattr( cfg, 'MODES', [ ] )
	
	if isinstance( modes, list ) and len( modes ) > 0:
		return str( modes[ 0 ] )
	
	return str( getattr( cfg, 'DEFAULT_MODE', 'Text Generation' ) or 'Text Generation' )

def get_model_modes_for_state( model_name: str ) -> List[ str ]:
	"""
		Purpose:
		--------
		Return the supported modes for the selected model using the config model registry
		when available, while preserving fallback compatibility with cfg.MODES.

		Parameters:
		-----------
		model_name : str
			Selected local model name.

		Returns:
		--------
		List[str]
			Supported mode names.
	"""
	model_value = str( model_name or get_default_model_name( ) ).strip( )
	
	if hasattr( cfg, 'get_model_modes' ):
		modes = cfg.get_model_modes( model_value )
	else:
		modes = getattr( cfg, 'MODES', [ ] )
	
	if isinstance( modes, list ) and len( modes ) > 0:
		return [ str( mode_name ) for mode_name in modes ]
	
	return [ 'Text Generation' ]

def get_model_path_for_state( model_name: str ) -> str:
	"""
		Purpose:
		--------
		Return the selected model path using the config model registry when available,
		while preserving fallback compatibility with cfg.MODEL_MAP and cfg.MODEL_PATH.

		Parameters:
		-----------
		model_name : str
			Selected local model name.

		Returns:
		--------
		str
			Resolved GGUF model path.
	"""
	model_value = str( model_name or get_default_model_name( ) ).strip( )
	
	if hasattr( cfg, 'get_model_path' ):
		return str( cfg.get_model_path( model_value ) or '' )
	
	if hasattr( cfg, 'MODEL_MAP' ) and model_value in cfg.MODEL_MAP:
		return str( cfg.MODEL_MAP.get( model_value, '' ) or '' )
	
	return str( getattr( cfg, 'MODEL_PATH', '' ) or '' )

def get_model_spec_for_state( model_name: str ) -> Dict[ str, Any ]:
	"""
		Purpose:
		--------
		Return the selected model registry specification when available.

		Parameters:
		-----------
		model_name : str
			Selected local model name.

		Returns:
		--------
		Dict[str, Any]
			Model specification dictionary.
	"""
	model_value = str( model_name or get_default_model_name( ) ).strip( )
	
	if hasattr( cfg, 'get_model_spec' ):
		spec = cfg.get_model_spec( model_value )
		if isinstance( spec, dict ):
			return spec
	
	return {
			'path': get_model_path_for_state( model_value ),
			'modes': get_model_modes_for_state( model_value ),
			'family': '',
			'size': '',
			'chat_template': 'chatml',
			'description': ''
	}

def initialize_model_mode_state( ) -> None:
	"""
		Purpose:
		--------
		Initialize widget-owned and derived model/mode session-state keys before the
		sidebar widgets are created.

		Parameters:
		-----------
		None

		Returns:
		--------
		None
	"""
	model_names = get_model_names_for_state( )
	default_model = get_default_model_name( )
	
	if 'selected_model_name' not in st.session_state:
		st.session_state[ 'selected_model_name' ] = default_model
	
	model_name = str(
		st.session_state.get( 'selected_model_name', default_model ) or default_model )
	
	if model_names and model_name not in model_names:
		model_name = default_model
		st.session_state[ 'selected_model_name' ] = model_name
	
	model_modes = get_model_modes_for_state( model_name )
	
	if 'selected_mode' not in st.session_state:
		st.session_state[ 'selected_mode' ] = ( model_modes[ 0 ]
		                                        if model_modes
		                                        else get_default_mode_name( model_name ) )
	
	selected_mode = str( st.session_state.get( 'selected_mode', get_default_mode_name( model_name ) ) or
		get_default_mode_name( model_name ) )
	
	if selected_mode not in model_modes:
		selected_mode = model_modes[ 0 ] if model_modes else get_default_mode_name( model_name )
		st.session_state[ 'selected_mode' ] = selected_mode
	
	st.session_state[ 'selected_model_path' ] = get_model_path_for_state( model_name )
	st.session_state[ 'selected_model_modes' ] = model_modes
	st.session_state[ 'selected_model_spec' ] = get_model_spec_for_state( model_name )
	st.session_state[ 'active_model_name' ] = model_name
	st.session_state[ 'mode' ] = selected_mode
	
	if 'model_switch_counter' not in st.session_state:
		st.session_state[ 'model_switch_counter' ] = 0

def synchronize_model_derived_state( ) -> None:
	"""
		Purpose:
		--------
		Synchronize derived model state without modifying widget-owned keys after their
		widgets have been instantiated.

		Parameters:
		-----------
		None

		Returns:
		--------
		None
	"""
	model_name = str( st.session_state.get( 'selected_model_name', get_default_model_name( ) ) or
		get_default_model_name( ) )
	
	model_modes = get_model_modes_for_state( model_name )
	
	st.session_state[ 'selected_model_path' ] = get_model_path_for_state( model_name )
	st.session_state[ 'selected_model_modes' ] = model_modes
	st.session_state[ 'selected_model_spec' ] = get_model_spec_for_state( model_name )
	st.session_state[ 'active_model_name' ] = model_name
	
	selected_mode = str(
		st.session_state.get( 'selected_mode', get_default_mode_name( model_name ) ) or
		get_default_mode_name( model_name )
	)
	
	if selected_mode in model_modes:
		st.session_state[ 'mode' ] = selected_mode
	else:
		st.session_state[ 'pending_selected_mode' ] = (
				model_modes[ 0 ] if model_modes else get_default_mode_name( model_name )
		)

def on_selected_model_change( ) -> None:
	"""
		Purpose:
		--------
		Streamlit callback used by the LLM selector to resynchronize derived model values
		after the selected model changes without directly modifying selected_mode.

		Parameters:
		-----------
		None

		Returns:
		--------
		None
	"""
	previous_model = str( st.session_state.get( 'active_model_name', '' ) or '' )
	model_name = str(
		st.session_state.get( 'selected_model_name', get_default_model_name( ) ) or
		get_default_model_name( )
	)
	
	model_modes = get_model_modes_for_state( model_name )
	
	st.session_state[ 'selected_model_path' ] = get_model_path_for_state( model_name )
	st.session_state[ 'selected_model_modes' ] = model_modes
	st.session_state[ 'selected_model_spec' ] = get_model_spec_for_state( model_name )
	st.session_state[ 'active_model_name' ] = model_name
	
	if previous_model and previous_model != model_name:
		st.session_state[ 'model_switch_counter' ] = ( int( st.session_state.get(
			'model_switch_counter', 0 ) or 0 ) + 1 )
	
	current_mode = str( st.session_state.get( 'selected_mode', '' ) or '' )
	if current_mode not in model_modes:
		st.session_state[ 'pending_selected_mode' ] = ( model_modes[ 0 ]
		                                                if model_modes
	
		                                                else get_default_mode_name( model_name ) )
	try:
		refresh_capability_session_state( )
		apply_model_safe_retrieval_defaults( model_name )
	except Exception:
		pass
		
def on_selected_mode_change( ) -> None:
	"""
		Purpose:
		--------
		Streamlit callback used by the AI Mode selector to keep the legacy mode key
		aligned with selected_mode.

		Parameters:
		-----------
		None

		Returns:
		--------
		None
	"""
	selected_mode = str( st.session_state.get( 'selected_mode', '' ) or '' )
	st.session_state[ 'mode' ] = selected_mode

def get_mode_constant( constant_name: str, fallback: str ) -> str:
	"""
		Purpose:
		--------
		Return a mode constant from config with a stable fallback. This allows app.py to
		accept expanded config.py mode definitions without crashing while config updates
		are being staged.

		Parameters:
		-----------
		constant_name : str
			Name of the config.py constant.

		fallback : str
			Fallback mode name.

		Returns:
		--------
		str
			Resolved mode name.
	"""
	try:
		value = cfg.__dict__.get( constant_name, fallback )
		value = str( value or fallback ).strip( )
		return value if value else fallback
	except Exception:
		return fallback

def get_mode_definition_text( mode_name: str ) -> str:
	"""
		Purpose:
		--------
		Return descriptive config.py text for expanded API modes when available.

		Parameters:
		-----------
		mode_name : str
			UI mode name.

		Returns:
		--------
		str
			Mode description text.
	"""
	try:
		image_mode = get_mode_constant( 'IMAGE_MODE', 'Images API' )
		audio_mode = get_mode_constant( 'AUDIO_MODE', 'Audio API' )
		
		if mode_name == image_mode:
			return str( cfg.__dict__.get( 'IMAGES_API', '' ) or '' ).strip( )
		
		if mode_name == audio_mode:
			return str( cfg.__dict__.get( 'AUDIO_API', '' ) or '' ).strip( )
		
		return ''
	except Exception:
		return ''

def get_selected_base_model( ) -> str:
	"""
		Purpose:
		--------
		Return the selected model's configured base model name.

		Parameters:
		-----------
		None

		Returns:
		--------
		str
			Base model name.
	"""
	try:
		spec = get_selected_model_spec( )
		return str( spec.get( 'base_model', '' ) or '' ).strip( )
	except Exception:
		return ''

def get_selected_model_family( ) -> str:
	"""
		Purpose:
		--------
		Return the selected model's configured model family.

		Parameters:
		-----------
		None

		Returns:
		--------
		str
			Model family name.
	"""
	try:
		spec = get_selected_model_spec( )
		return str( spec.get( 'family', '' ) or '' ).strip( )
	except Exception:
		return ''

def get_selected_chat_template( ) -> str:
	"""
		Purpose:
		--------
		Return the selected model's configured chat template.

		Parameters:
		-----------
		None

		Returns:
		--------
		str
			Chat template name.
	"""
	try:
		spec = get_selected_model_spec( )
		return str( spec.get( 'chat_template', '' ) or '' ).strip( )
	except Exception:
		return ''

def is_buddy_model( ) -> bool:
	"""
		Purpose:
		--------
		Determine whether the selected model is Buddy or a Buddy base model.

		Parameters:
		-----------
		None

		Returns:
		--------
		bool
			True when Buddy is selected; otherwise False.
	"""
	model_name = get_selected_model_name( ).lower( )
	base_model = get_selected_base_model( ).lower( )
	return model_name == 'buddy' or base_model == 'gemma-3-270m-it'

def is_gipity_model( ) -> bool:
	"""
		Purpose:
		--------
		Determine whether the selected model is Gipity or a GPT-OSS base model.

		Parameters:
		-----------
		None

		Returns:
		--------
		bool
			True when Gipity is selected; otherwise False.
	"""
	model_name = get_selected_model_name( ).lower( )
	base_model = get_selected_base_model( ).lower( )
	return model_name == 'gipity' or base_model == 'gpt-oss-20b'

def is_gemma4_model( ) -> bool:
	"""
		Purpose:
		--------
		Determine whether the selected model uses the Gemma 4 E4B base model.

		Parameters:
		-----------
		None

		Returns:
		--------
		bool
			True when a Gemma 4 E4B model is selected; otherwise False.
	"""
	model_name = get_selected_model_name( ).lower( )
	base_model = get_selected_base_model( ).lower( )
	return model_name in ('jimi', 'nisty') or base_model == 'gemma-4-e4b-it'

def is_jimi_or_nisty_model( ) -> bool:
	"""
		Purpose:
		--------
		Determine whether the selected model is Jimi or Nisty.

		Parameters:
		-----------
		None

		Returns:
		--------
		bool
			True when Jimi or Nisty is selected; otherwise False.
	"""
	model_name = get_selected_model_name( ).lower( )
	return model_name in ('jimi', 'nisty')

def model_supports_mode( mode_name: str ) -> bool:
	"""
		Purpose:
		--------
		Determine whether the selected model registry advertises a specific UI mode.

		Parameters:
		-----------
		mode_name : str
			UI mode name.

		Returns:
		--------
		bool
			True when the mode is listed for the selected model; otherwise False.
	"""
	try:
		model_modes = st.session_state.get(
			'selected_model_modes',
			get_model_modes_for_state( get_selected_model_name( ) )
		)
		
		if not isinstance( model_modes, list ):
			return False
		
		return str( mode_name or '' ) in [ str( m ) for m in model_modes ]
	except Exception:
		return False

def get_runtime_multimodal_status( ) -> Dict[ str, Any ]:
	"""
		Purpose:
		--------
		Return the current runtime's multimodal adapter status. This detects whether app.py
		has an image/audio-capable local adapter configured separately from the model
		registry. The function fails closed so newly exposed modes cannot crash.

		Parameters:
		-----------
		None

		Returns:
		--------
		Dict[str, Any]
			Runtime multimodal status flags and message.
	"""
	status = {
			'image_runtime_available': False,
			'audio_runtime_available': False,
			'function_runtime_available': True,
			'web_runtime_available': False,
			'runtime_name': 'llama-cpp-python',
			'message': 'The current local runtime is treated as text-only until a '
			           'multimodal adapter is explicitly wired into app.py.'
	}
	
	try:
		if bool( cfg.__dict__.get( 'IMAGE_RUNTIME_AVAILABLE', False ) ):
			status[ 'image_runtime_available' ] = True
		
		if bool( cfg.__dict__.get( 'AUDIO_RUNTIME_AVAILABLE', False ) ):
			status[ 'audio_runtime_available' ] = True
		
		if bool( cfg.__dict__.get( 'WEB_RUNTIME_AVAILABLE', False ) ):
			status[ 'web_runtime_available' ] = True
		
		runtime_name = str( cfg.__dict__.get( 'MULTIMODAL_RUNTIME_NAME', '' ) or '' ).strip( )
		if runtime_name:
			status[ 'runtime_name' ] = runtime_name
		
		if status[ 'image_runtime_available' ] or status[ 'audio_runtime_available' ]:
			status[ 'message' ] = 'A multimodal runtime adapter is configured.'
		
		return status
	except Exception:
		return status

def get_active_model_capabilities( ) -> Dict[ str, Any ]:
	"""
		Purpose:
		--------
		Return selected model capability flags used by expanded Text, Image, Audio,
		Function Calling, Coding, Thinking, and Web Browsing workflows.

		Parameters:
		-----------
		None

		Returns:
		--------
		Dict[str, Any]
			Capability contract for the selected model.
	"""
	model_name = get_selected_model_name( )
	base_model = get_selected_base_model( )
	family = get_selected_model_family( )
	template = get_selected_chat_template( )
	image_mode = get_mode_constant( 'IMAGE_MODE', 'Images API' )
	audio_mode = get_mode_constant( 'AUDIO_MODE', 'Audio API' )
	docqna_mode = get_mode_constant( 'DOCQNA_MODE', 'Document Q&A' )
	semantic_mode = get_mode_constant( 'SEMANTIC_MODE', 'Semantic Search' )
	prompt_mode = get_mode_constant( 'PROMPT_MODE', 'Prompt Engineering' )
	data_mode = get_mode_constant( 'DATA_MODE', 'Data Management' )
	text_mode = get_mode_constant( 'TEXT_MODE', 'Text Generation' )
	runtime_status = get_runtime_multimodal_status( )
	
	capabilities: Dict[ str, Any ] = {
			'model_name': model_name,
			'base_model': base_model,
			'family': family,
			'chat_template': template,
			'text_generation': model_supports_mode( text_mode ),
			'document_qna': model_supports_mode( docqna_mode ),
			'semantic_search': model_supports_mode( semantic_mode ),
			'prompt_engineering': model_supports_mode( prompt_mode ),
			'data_management': model_supports_mode( data_mode ),
			'image_mode': model_supports_mode( image_mode ) and is_jimi_or_nisty_model( ),
			'audio_mode': model_supports_mode( audio_mode ) and is_jimi_or_nisty_model( ),
			'image_runtime_available': bool(
				runtime_status.get( 'image_runtime_available', False ) ),
			'audio_runtime_available': bool(
				runtime_status.get( 'audio_runtime_available', False ) ),
			'function_calling': is_gemma4_model( ) or is_gipity_model( ),
			'coding': is_gemma4_model( ),
			'thinking': is_gemma4_model( ),
			'web_browsing': is_gipity_model( ),
			'gipity': is_gipity_model( ),
			'gemma4': is_gemma4_model( ),
			'buddy': is_buddy_model( ),
			'runtime_status': runtime_status
	}
	
	return capabilities

def model_supports_capability( capability: str ) -> bool:
	"""
		Purpose:
		--------
		Determine whether the selected model supports a named expanded capability.

		Parameters:
		-----------
		capability : str
			Capability name.

		Returns:
		--------
		bool
			True when the selected model supports the capability; otherwise False.
	"""
	try:
		capabilities = get_active_model_capabilities( )
		return bool( capabilities.get( str( capability or '' ), False ) )
	except Exception:
		return False

def get_capability_status_message( capability: str ) -> str:
	"""
		Purpose:
		--------
		Return a user-facing status message for unsupported or unavailable capabilities.

		Parameters:
		-----------
		capability : str
			Capability name.

		Returns:
		--------
		str
			Status message.
	"""
	capabilities = get_active_model_capabilities( )
	model_name = str( capabilities.get( 'model_name', get_selected_model_name( ) ) or '' )
	runtime_status = capabilities.get( 'runtime_status', { } )
	
	if capability == 'image_mode':
		if not capabilities.get( 'image_mode', False ):
			return f'{model_name} is not configured for Image Mode.'
		
		if not capabilities.get( 'image_runtime_available', False ):
			return str( runtime_status.get( 'message', '' ) or
			            'Image Mode is configured, but no image-capable runtime is wired yet.' )
	
	if capability == 'audio_mode':
		if not capabilities.get( 'audio_mode', False ):
			return f'{model_name} is not configured for Audio Mode.'
		
		if not capabilities.get( 'audio_runtime_available', False ):
			return str( runtime_status.get( 'message', '' ) or
			            'Audio Mode is configured, but no audio-capable runtime is wired yet.' )
	
	if not bool( capabilities.get( capability, False ) ):
		return f'{model_name} does not advertise the "{capability}" capability.'
	
	return f'{model_name} supports the "{capability}" capability.'

def get_default_function_schema_text( ) -> str:
	"""
		Purpose:
		--------
		Return a safe starter JSON schema for function-calling workflows.

		Parameters:
		-----------
		None

		Returns:
		--------
		str
			Starter JSON function schema text.
	"""
	return '''{
	"name": "summarize_text",
	"description": "Summarize supplied text into concise bullet points.",
	"parameters": {
		"type": "object",
		"properties": {
			"text": {
				"type": "string",
				"description": "The text to summarize."
			},
			"max_bullets": {
				"type": "integer",
				"description": "Maximum number of bullets to return."
			}
		},
		"required": [
			"text"
		]
	}
}'''

def initialize_capability_session_state( ) -> None:
	"""
		Purpose:
		--------
		Initialize expanded capability session-state keys before Image, Audio, Function
		Calling, Coding, Thinking, and Web Browsing controls are introduced.

		Parameters:
		-----------
		None

		Returns:
		--------
		None
	"""
	defaults: Dict[ str, Any ] = {
			'image_prompt': '',
			'image_uploaded_name': '',
			'image_response': '',
			'image_status': '',
			'image_context_buffer': '',
			'image_send_to_text': False,
			'audio_prompt': '',
			'audio_uploaded_name': '',
			'audio_response': '',
			'audio_status': '',
			'audio_transcript': '',
			'audio_context_buffer': '',
			'audio_send_to_text': False,
			'function_schema_text': get_default_function_schema_text( ),
			'function_call_prompt': '',
			'function_call_response': '',
			'function_call_result': '',
			'function_call_status': '',
			'function_call_enabled': False,
			'function_call_model_json': '',
			'coding_mode_enabled': False,
			'coding_test_request': False,
			'coding_explain_request': False,
			'thinking_mode_enabled': False,
			'thinking_effort': 'Medium',
			'thinking_summary_enabled': True,
			'web_browse_url': '',
			'web_browse_allow_domain': '',
			'web_browse_prompt': '',
			'web_browse_result': '',
			'web_browse_status': '',
			'web_browse_context_buffer': '',
			'web_browse_send_to_text': False,
			'active_model_capabilities': { }
	}
	
	for key, value in defaults.items( ):
		if key not in st.session_state:
			st.session_state[ key ] = value
	
	st.session_state[ 'active_model_capabilities' ] = get_active_model_capabilities( )

def refresh_capability_session_state( ) -> None:
	"""
		Purpose:
		--------
		Refresh derived capability state after model or mode changes without clearing
		user-owned text, uploaded-file names, generated output, or existing chat state.

		Parameters:
		-----------
		None

		Returns:
		--------
		None
	"""
	try:
		st.session_state[ 'active_model_capabilities' ] = get_active_model_capabilities( )
	except Exception:
		st.session_state[ 'active_model_capabilities' ] = { }

initialize_capability_session_state( )

def get_model_retrieval_profile( model_name: str ) -> Dict[ str, Any ]:
	"""
		Purpose:
		--------
		Return model-safe retrieval defaults for Document Q&A and Semantic Search. Smaller
		models receive narrower retrieval windows so grounded prompts stay concise and
		less likely to exceed practical local runtime limits.

		Parameters:
		-----------
		model_name : str
			Selected local model name.

		Returns:
		--------
		Dict[str, Any]
			Retrieval profile values.
	"""
	model_value = str( model_name or get_selected_model_name( ) or '' ).strip( ).lower( )
	base_model = ''
	
	try:
		spec = get_model_spec_for_state( model_name )
		if isinstance( spec, dict ):
			base_model = str( spec.get( 'base_model', '' ) or '' ).strip( ).lower( )
	except Exception:
		base_model = ''
	
	if model_value == 'buddy' or base_model == 'gemma-3-270m-it':
		return {
				'profile_name': 'Buddy Compact Retrieval',
				'retrieval_k': 3,
				'retrieval_chunk_size': 800,
				'retrieval_chunk_overlap': 120,
				'semantic_top_k': 4,
				'semantic_chunk_size': 800,
				'semantic_chunk_overlap': 120,
				'semantic_min_similarity': 0.05,
				'require_grounding': True,
				'answer_from_excerpts_only': True,
				'show_retrieved_chunks': True,
				'prefer_sqlite_vec': True,
				'allow_similarity_fallback': True,
				'semantic_show_diagnostics': True,
				'semantic_group_by_document': False
		}
	
	return {
			'profile_name': 'Standard Retrieval',
			'retrieval_k': 6,
			'retrieval_chunk_size': 1200,
			'retrieval_chunk_overlap': 200,
			'semantic_top_k': 8,
			'semantic_chunk_size': 1200,
			'semantic_chunk_overlap': 200,
			'semantic_min_similarity': 0.0,
			'require_grounding': True,
			'answer_from_excerpts_only': True,
			'show_retrieved_chunks': True,
			'prefer_sqlite_vec': True,
			'allow_similarity_fallback': True,
			'semantic_show_diagnostics': True,
			'semantic_group_by_document': False
	}

def has_user_tuned_retrieval_controls( ) -> bool:
	"""
		Purpose:
		--------
		Determine whether the current retrieval controls have already been changed by the
		user or by a previously applied model profile. This prevents model-safe defaults
		from overwriting user-tuned values on every Streamlit rerun.

		Parameters:
		-----------
		None

		Returns:
		--------
		bool
			True when retrieval controls should be preserved; otherwise False.
	"""
	return bool( st.session_state.get( 'retrieval_controls_user_tuned', False ) )

def mark_retrieval_controls_user_tuned( ) -> None:
	"""
		Purpose:
		--------
		Mark retrieval controls as user-tuned. Later controls can call this callback if
		needed to permanently preserve manual user settings across model changes.

		Parameters:
		-----------
		None

		Returns:
		--------
		None
	"""
	st.session_state[ 'retrieval_controls_user_tuned' ] = True

def apply_retrieval_profile( profile: Dict[ str, Any ], force: bool = False ) -> None:
	"""
		Purpose:
		--------
		Apply a retrieval profile to Document Q&A and Semantic Search session-state keys.
		The profile is applied only when forced or when no user-tuned override exists.

		Parameters:
		-----------
		profile : Dict[str, Any]
			Retrieval profile to apply.

		force : bool
			When True, apply the profile even if retrieval_controls_user_tuned is True.

		Returns:
		--------
		None
	"""
	if not isinstance( profile, dict ) or len( profile ) == 0:
		return
	
	if has_user_tuned_retrieval_controls( ) and not force:
		return
	
	assignments = {
			'retrieval_k': int( profile.get( 'retrieval_k', 6 ) ),
			'retrieval_chunk_size': int( profile.get( 'retrieval_chunk_size', 1200 ) ),
			'retrieval_chunk_overlap': int( profile.get( 'retrieval_chunk_overlap', 200 ) ),
			'semantic_top_k': int( profile.get( 'semantic_top_k', 8 ) ),
			'semantic_chunk_size': int( profile.get( 'semantic_chunk_size', 1200 ) ),
			'semantic_chunk_overlap': int( profile.get( 'semantic_chunk_overlap', 200 ) ),
			'semantic_min_similarity': float( profile.get( 'semantic_min_similarity', 0.0 ) ),
			'require_grounding': bool( profile.get( 'require_grounding', True ) ),
			'answer_from_excerpts_only': bool( profile.get( 'answer_from_excerpts_only', True ) ),
			'show_retrieved_chunks': bool( profile.get( 'show_retrieved_chunks', True ) ),
			'prefer_sqlite_vec': bool( profile.get( 'prefer_sqlite_vec', True ) ),
			'allow_similarity_fallback': bool( profile.get( 'allow_similarity_fallback', True ) ),
			'semantic_show_diagnostics': bool( profile.get( 'semantic_show_diagnostics', True ) ),
			'semantic_group_by_document': bool( profile.get( 'semantic_group_by_document', False ) )
	}
	
	for key, value in assignments.items( ):
		st.session_state[ key ] = value
	
	st.session_state[ 'active_retrieval_profile' ] = str(
		profile.get( 'profile_name', 'Standard Retrieval' ) or 'Standard Retrieval' )
	st.session_state[ 'active_retrieval_profile_model' ] = get_selected_model_name( )

def apply_model_safe_retrieval_defaults( model_name: str = '' ) -> None:
	"""
		Purpose:
		--------
		Apply model-safe retrieval defaults when the selected model changes. Buddy receives
		compact retrieval settings suitable for a 270M model; other models receive standard
		settings unless the user has already tuned retrieval controls.

		Parameters:
		-----------
		model_name : str
			Optional selected model name. When omitted, the current selected model is used.

		Returns:
		--------
		None
	"""
	selected_model = str( model_name or get_selected_model_name( ) or '' ).strip( )
	if not selected_model:
		return
	
	last_profile_model = str(
		st.session_state.get( 'last_retrieval_profile_model', '' ) or '' ).strip( )
	
	if last_profile_model == selected_model:
		return
	
	profile = get_model_retrieval_profile( selected_model )
	force_apply = not has_user_tuned_retrieval_controls( )
	
	apply_retrieval_profile( profile=profile, force=force_apply )
	
	st.session_state[ 'last_retrieval_profile_model' ] = selected_model
	st.session_state[ 'retrieval_profile_status' ] = (
			f'Active retrieval profile: {st.session_state.get( "active_retrieval_profile", "" )}')

def reset_model_safe_retrieval_defaults( ) -> None:
	"""
		Purpose:
		--------
		Clear manual retrieval override state and reapply the selected model's recommended
		Document Q&A and Semantic Search retrieval profile.

		Parameters:
		-----------
		None

		Returns:
		--------
		None
	"""
	st.session_state[ 'retrieval_controls_user_tuned' ] = False
	st.session_state[ 'last_retrieval_profile_model' ] = ''
	apply_model_safe_retrieval_defaults( get_selected_model_name( ) )

def initialize_model_safe_retrieval_state( ) -> None:
	"""
		Purpose:
		--------
		Initialize retrieval-profile tracking keys and apply model-safe defaults once after
		capability session state is initialized.

		Parameters:
		-----------
		None

		Returns:
		--------
		None
	"""
	if 'retrieval_controls_user_tuned' not in st.session_state:
		st.session_state[ 'retrieval_controls_user_tuned' ] = False
	
	if 'last_retrieval_profile_model' not in st.session_state:
		st.session_state[ 'last_retrieval_profile_model' ] = ''
	
	if 'active_retrieval_profile' not in st.session_state:
		st.session_state[ 'active_retrieval_profile' ] = ''
	
	if 'active_retrieval_profile_model' not in st.session_state:
		st.session_state[ 'active_retrieval_profile_model' ] = ''
	
	if 'retrieval_profile_status' not in st.session_state:
		st.session_state[ 'retrieval_profile_status' ] = ''
	
	apply_model_safe_retrieval_defaults( get_selected_model_name( ) )

initialize_model_safe_retrieval_state( )

def extract_json_object_from_text( text: str ) -> Dict[ str, Any ]:
	"""
		Purpose:
		--------
		Extract the first valid JSON object from model-generated text. This supports
		function-calling outputs where the model may accidentally wrap the object in
		markdown fences or explanatory prose.

		Parameters:
		-----------
		text : str
			Model-generated text.

		Returns:
		--------
		Dict[str, Any]
			Parsed JSON object.
	"""
	import json
	
	text_value = str( text or '' ).strip( )
	if not text_value:
		raise ValueError( 'No function-call text was provided.' )
	
	if text_value.startswith( '```' ):
		text_value = re.sub( r'^```(?:json)?\s*', '', text_value, flags=re.IGNORECASE )
		text_value = re.sub( r'\s*```$', '', text_value )
		text_value = text_value.strip( )
	
	try:
		parsed = json.loads( text_value )
		if isinstance( parsed, dict ):
			return parsed
	except Exception:
		pass
	
	start = text_value.find( '{' )
	if start < 0:
		raise ValueError( 'No JSON object start marker was found.' )
	
	depth = 0
	in_string = False
	escape = False
	
	for idx in range( start, len( text_value ) ):
		char = text_value[ idx ]
		
		if escape:
			escape = False
			continue
		
		if char == '\\':
			escape = True
			continue
		
		if char == '"':
			in_string = not in_string
			continue
		
		if in_string:
			continue
		
		if char == '{':
			depth += 1
		elif char == '}':
			depth -= 1
			
			if depth == 0:
				candidate = text_value[ start:idx + 1 ]
				parsed = json.loads( candidate )
				if not isinstance( parsed, dict ):
					raise ValueError( 'The parsed function call was not a JSON object.' )
				
				return parsed
	
	raise ValueError( 'No complete JSON object was found.' )

def normalize_tool_call( tool_call: Dict[ str, Any ] ) -> Dict[ str, Any ]:
	"""
		Purpose:
		--------
		Normalize a model-generated tool call into the app contract:
		{"name": "...", "arguments": {...}}.

		Parameters:
		-----------
		tool_call : Dict[str, Any]
			Parsed tool-call JSON object.

		Returns:
		--------
		Dict[str, Any]
			Normalized tool-call object.
	"""
	if not isinstance( tool_call, dict ):
		raise ValueError( 'Tool call must be a dictionary.' )
	
	name = str( tool_call.get( 'name', '' ) or tool_call.get( 'function', '' ) or '' ).strip( )
	args = tool_call.get( 'arguments', { } )
	
	if not name:
		raise ValueError( 'Tool call is missing a function name.' )
	
	if isinstance( args, str ):
		try:
			args = extract_json_object_from_text( args )
		except Exception:
			args = { 'text': args }
	
	if not isinstance( args, dict ):
		raise ValueError( 'Tool-call arguments must be a dictionary.' )
	
	return {
			'name': name,
			'arguments': args
	}

def get_allowed_function_names( ) -> List[ str ]:
	"""
		Purpose:
		--------
		Return the function names that app.py is allowed to execute. This prevents model
		output from invoking arbitrary functions.

		Parameters:
		-----------
		None

		Returns:
		--------
		List[str]
			Allowlisted function names.
	"""
	return [
			'summarize_text',
			'extract_keywords',
			'web_browse_url'
	]

def summarize_text_tool( text: str, max_bullets: int = 5 ) -> str:
	"""
		Purpose:
		--------
		Summarize supplied text using a deterministic local sentence extraction fallback.
		This is intentionally non-agentic and does not execute arbitrary model code.

		Parameters:
		-----------
		text : str
			Text to summarize.

		max_bullets : int
			Maximum number of bullets to return.

		Returns:
		--------
		str
			Bullet summary.
	"""
	text_value = re.sub( r'\s+', ' ', str( text or '' ) ).strip( )
	if not text_value:
		return 'No text was provided.'
	
	try:
		bullet_count = int( max_bullets )
	except Exception:
		bullet_count = 5
	
	if bullet_count <= 0:
		bullet_count = 5
	
	sentences = re.split( r'(?<=[.!?])\s+', text_value )
	sentences = [ s.strip( ) for s in sentences if s and s.strip( ) ]
	selected = sentences[ :bullet_count ]
	
	if not selected:
		selected = [ text_value[ :800 ] ]
	
	return '\n'.join( [ f'- {sentence}' for sentence in selected ] )

def extract_keywords_tool( text: str, max_keywords: int = 15 ) -> str:
	"""
		Purpose:
		--------
		Extract simple frequency-ranked keywords from supplied text without external
		dependencies.

		Parameters:
		-----------
		text : str
			Text to analyze.

		max_keywords : int
			Maximum number of keywords to return.

		Returns:
		--------
		str
			Comma-separated keyword list.
	"""
	text_value = str( text or '' ).lower( )
	if not text_value:
		return ''
	
	try:
		keyword_count = int( max_keywords )
	except Exception:
		keyword_count = 15
	
	if keyword_count <= 0:
		keyword_count = 15
	
	stop_words = {
			'a', 'an', 'and', 'are', 'as', 'at', 'be', 'by', 'for', 'from',
			'has', 'have', 'in', 'is', 'it', 'its', 'of', 'on', 'or', 'that',
			'the', 'their', 'this', 'to', 'was', 'were', 'with', 'you', 'your'
	}
	
	words = re.findall( r'[a-zA-Z][a-zA-Z0-9_\-]{2,}', text_value )
	counts: Dict[ str, int ] = { }
	
	for word in words:
		if word in stop_words:
			continue
		
		counts[ word ] = counts.get( word, 0 ) + 1
	
	ranked = sorted( counts.items( ), key=lambda item: item[ 1 ], reverse=True )
	keywords = [ word for word, _ in ranked[ :keyword_count ] ]
	
	return ', '.join( keywords )

def is_private_or_local_hostname( hostname: str ) -> bool:
	"""
		Purpose:
		--------
		Determine whether a hostname resolves to a local, loopback, private, reserved, or
		link-local address. This blocks server-side requests to private network resources.

		Parameters:
		-----------
		hostname : str
			URL hostname.

		Returns:
		--------
		bool
			True when the hostname is private or local; otherwise False.
	"""
	import ipaddress
	import socket
	
	host_value = str( hostname or '' ).strip( ).lower( )
	if not host_value:
		return True
	
	if host_value in ('localhost', '0.0.0.0') or host_value.endswith( '.local' ):
		return True
	
	try:
		ip_value = ipaddress.ip_address( host_value )
		return bool(
			ip_value.is_private
			or ip_value.is_loopback
			or ip_value.is_link_local
			or ip_value.is_reserved
			or ip_value.is_multicast
			or ip_value.is_unspecified
		)
	except Exception:
		pass
	
	try:
		addresses = socket.getaddrinfo( host_value, None )
	except Exception:
		return True
	
	for address in addresses:
		try:
			ip_text = address[ 4 ][ 0 ]
			ip_value = ipaddress.ip_address( ip_text )
			if (
					ip_value.is_private
					or ip_value.is_loopback
					or ip_value.is_link_local
					or ip_value.is_reserved
					or ip_value.is_multicast
					or ip_value.is_unspecified
			):
				return True
		except Exception:
			return True
	
	return False

def validate_web_url( url: str, allowed_domain: str = '' ) -> str:
	"""
		Purpose:
		--------
		Validate an outbound web-browsing URL. Only HTTP and HTTPS URLs are allowed, and
		private/local network targets are blocked.

		Parameters:
		-----------
		url : str
			User-supplied URL.

		allowed_domain : str
			Optional allowed domain suffix.

		Returns:
		--------
		str
			Validated URL.
	"""
	from urllib.parse import urlparse
	
	url_value = str( url or '' ).strip( )
	if not url_value:
		raise ValueError( 'A URL is required.' )
	
	parsed = urlparse( url_value )
	if parsed.scheme.lower( ) not in ('http', 'https'):
		raise ValueError( 'Only http and https URLs are allowed.' )
	
	if not parsed.netloc or not parsed.hostname:
		raise ValueError( 'The URL must include a valid host.' )
	
	hostname = str( parsed.hostname or '' ).strip( ).lower( )
	
	if is_private_or_local_hostname( hostname ):
		raise ValueError( 'Private, local, loopback, reserved, and link-local hosts are blocked.' )
	
	domain_value = str( allowed_domain or '' ).strip( ).lower( )
	if domain_value:
		if domain_value.startswith( 'http://' ) or domain_value.startswith( 'https://' ):
			domain_value = str( urlparse( domain_value ).hostname or '' ).strip( ).lower( )
		
		domain_value = domain_value.lstrip( '.' )
		if hostname != domain_value and not hostname.endswith( f'.{domain_value}' ):
			raise ValueError( f'The URL host is not within the allowed domain: {domain_value}' )
	
	return url_value

def html_to_readable_text( html_text: str ) -> str:
	"""
		Purpose:
		--------
		Convert HTML to readable text using a dependency-free parser fallback.

		Parameters:
		-----------
		html_text : str
			Raw HTML text.

		Returns:
		--------
		str
			Readable extracted text.
	"""
	import html
	
	text_value = str( html_text or '' )
	text_value = re.sub( r'(?is)<(script|style|noscript).*?>.*?</\1>', ' ', text_value )
	text_value = re.sub( r'(?is)<br\s*/?>', '\n', text_value )
	text_value = re.sub( r'(?is)</p\s*>', '\n\n', text_value )
	text_value = re.sub( r'(?is)<[^>]+>', ' ', text_value )
	text_value = html.unescape( text_value )
	text_value = re.sub( r'[ \t\r\f\v]+', ' ', text_value )
	text_value = re.sub( r'\n\s+', '\n', text_value )
	text_value = re.sub( r'\n{3,}', '\n\n', text_value )
	
	return text_value.strip( )

def fetch_web_text( url: str, allowed_domain: str = '', timeout_seconds: int = 15,
		max_chars: int = 12000 ) -> Dict[ str, Any ]:
	"""
		Purpose:
		--------
		Fetch readable text from a public HTTP/HTTPS URL with timeout, size, and private
		network safeguards.

		Parameters:
		-----------
		url : str
			User-supplied URL.

		allowed_domain : str
			Optional allowed domain suffix.

		timeout_seconds : int
			Network timeout in seconds.

		max_chars : int
			Maximum readable text characters returned.

		Returns:
		--------
		Dict[str, Any]
			Web fetch result.
	"""
	from urllib.request import Request, urlopen
	
	try:
		timeout_value = int( timeout_seconds )
	except Exception:
		timeout_value = 15
	
	if timeout_value <= 0:
		timeout_value = 15
	
	try:
		max_char_value = int( max_chars )
	except Exception:
		max_char_value = 12000
	
	if max_char_value <= 0:
		max_char_value = 12000
	
	validated_url = validate_web_url( url=url, allowed_domain=allowed_domain )
	
	request = Request(
		validated_url,
		headers={
				'User-Agent': 'Loca-Llama/1.0 TextFetcher'
		}
	)
	
	with urlopen( request, timeout=timeout_value ) as response:
		content_type = str( response.headers.get( 'Content-Type', '' ) or '' )
		raw = response.read( max_char_value * 4 )
	
	text = raw.decode( 'utf-8', errors='ignore' )
	if 'html' in content_type.lower( ) or '<html' in text.lower( ):
		readable_text = html_to_readable_text( text )
	else:
		readable_text = re.sub( r'\s+', ' ', text ).strip( )
	
	if len( readable_text ) > max_char_value:
		readable_text = readable_text[ :max_char_value ].strip( )
	
	return {
			'url': validated_url,
			'content_type': content_type,
			'text': readable_text,
			'length': len( readable_text )
	}

def web_browse_url_tool( url: str, prompt: str = '', allowed_domain: str = '',
		max_chars: int = 12000 ) -> str:
	"""
		Purpose:
		--------
		Fetch a public web page and return bounded text suitable for model grounding.

		Parameters:
		-----------
		url : str
			User-supplied URL.

		prompt : str
			Optional user task for the fetched content.

		allowed_domain : str
			Optional allowed domain suffix.

		max_chars : int
			Maximum readable text characters returned.

		Returns:
		--------
		str
			Readable web context.
	"""
	if not model_supports_capability( 'web_browsing' ):
		return get_capability_status_message( 'web_browsing' )
	
	result = fetch_web_text(
		url=url,
		allowed_domain=allowed_domain,
		timeout_seconds=15,
		max_chars=max_chars
	)
	
	task_text = str( prompt or '' ).strip( )
	parts = [
			f'Web Source: {result.get( "url", "" )}',
			f'Content Type: {result.get( "content_type", "" )}',
			f'Characters: {result.get( "length", 0 )}'
	]
	
	if task_text:
		parts.append( f'User Web Task: {task_text}' )
	
	parts.append( 'Fetched Web Text:' )
	parts.append( str( result.get( 'text', '' ) or '' ) )
	
	return '\n\n'.join( parts ).strip( )

def execute_allowlisted_function( tool_call: Dict[ str, Any ] ) -> Dict[ str, Any ]:
	"""
		Purpose:
		--------
		Execute a normalized, allowlisted app function. Arbitrary model-generated function
		names are rejected.

		Parameters:
		-----------
		tool_call : Dict[str, Any]
			Normalized tool-call object.

		Returns:
		--------
		Dict[str, Any]
			Tool execution result.
	"""
	normalized = normalize_tool_call( tool_call )
	name = normalized[ 'name' ]
	args = normalized[ 'arguments' ]
	
	if name not in get_allowed_function_names( ):
		raise ValueError( f'Function "{name}" is not allowlisted.' )
	
	if name == 'summarize_text':
		result = summarize_text_tool(
			text=str( args.get( 'text', '' ) or '' ),
			max_bullets=int( args.get( 'max_bullets', 5 ) or 5 )
		)
	
	elif name == 'extract_keywords':
		result = extract_keywords_tool(
			text=str( args.get( 'text', '' ) or '' ),
			max_keywords=int( args.get( 'max_keywords', 15 ) or 15 )
		)
	
	elif name == 'web_browse_url':
		result = web_browse_url_tool(
			url=str( args.get( 'url', '' ) or '' ),
			prompt=str( args.get( 'prompt', '' ) or '' ),
			allowed_domain=str( args.get( 'allowed_domain', '' ) or '' ),
			max_chars=int( args.get( 'max_chars', 12000 ) or 12000 )
		)
	
	else:
		raise ValueError( f'Function "{name}" is not implemented.' )
	
	return {
			'name': name,
			'arguments': args,
			'result': result
	}

def build_tool_call_generation_prompt( user_task: str ) -> str:
	"""
		Purpose:
		--------
		Build a focused prompt that asks the selected model to emit one strict JSON
		function-call object.

		Parameters:
		-----------
		user_task : str
			User task to translate into a tool call.

		Returns:
		--------
		str
			Tool-call generation prompt.
	"""
	schema_text = str( st.session_state.get( 'function_schema_text', '' ) or '' ).strip( )
	available_functions = ', '.join( get_allowed_function_names( ) )
	
	return f'''Create one app-mediated function call for the user task below.
		
		Rules:
		- Return one strict JSON object only.
		- Do not wrap the object in markdown.
		- Use this shape: {{"name":"function_name","arguments":{{...}}}}
		- Use only one of these allowlisted functions: {available_functions}
		- Do not invent function names.
		- Do not include prose outside the JSON object.
		
		Current function schema:
		{schema_text}
		
		User task:
		{user_task}'''

def generate_function_call_json( user_task: str ) -> str:
	"""
		Purpose:
		--------
		Ask the selected local model to generate a strict JSON function-call object.

		Parameters:
		-----------
		user_task : str
			User task to convert into a function call.

		Returns:
		--------
		str
			Generated model text.
	"""
	if not model_supports_capability( 'function_calling' ):
		return get_capability_status_message( 'function_calling' )
	
	task_text = str( user_task or '' ).strip( )
	if not task_text:
		return 'No function-call task was provided.'
	
	prior_function_enabled = bool( st.session_state.get( 'function_call_enabled', False ) )
	st.session_state[ 'function_call_enabled' ] = True
	
	try:
		response = run_llm_turn(
			user_input=build_tool_call_generation_prompt( task_text ),
			temperature=0.0,
			top_p=1.0,
			repeat_penalty=float( st.session_state.get( 'repeat_penalty', 1.1 ) ),
			max_tokens=512,
			stream=False,
			output=None
		)
	finally:
		st.session_state[ 'function_call_enabled' ] = prior_function_enabled
	
	return str( response or '' ).strip( )

def execute_tool_call_text( tool_call_text: str ) -> Dict[ str, Any ]:
	"""
		Purpose:
		--------
		Parse and execute model-generated tool-call text through the app's allowlisted
		function layer.

		Parameters:
		-----------
		tool_call_text : str
			Model-generated function-call JSON text.

		Returns:
		--------
		Dict[str, Any]
			Tool execution result.
	"""
	parsed = extract_json_object_from_text( tool_call_text )
	normalized = normalize_tool_call( parsed )
	return execute_allowlisted_function( normalized )

def build_tool_result_final_prompt( user_task: str, tool_result: Dict[ str, Any ] ) -> str:
	"""
		Purpose:
		--------
		Build a final-answer prompt from a validated tool execution result.

		Parameters:
		-----------
		user_task : str
			Original user task.

		tool_result : Dict[str, Any]
			Executed tool result.

		Returns:
		--------
		str
			Final answer prompt.
	"""
	return f'''Use the validated app tool result below to answer the user.

		User task:
		{str( user_task or '' ).strip( )}
		
		Tool used:
		{tool_result.get( 'name', '' )}
		
		Tool arguments:
		{tool_result.get( 'arguments', { } )}
		
		Tool result:
		{tool_result.get( 'result', '' )}
		
		Return a useful final answer grounded in the tool result.'''

def generate_tool_grounded_final_answer( user_task: str, tool_result: Dict[ str, Any ] ) -> str:
	"""
		Purpose:
		--------
		Generate a final answer grounded in an executed tool result.

		Parameters:
		-----------
		user_task : str
			Original user task.

		tool_result : Dict[str, Any]
			Executed tool result.

		Returns:
		--------
		str
			Model-generated final answer.
	"""
	try:
		return run_llm_turn(
			user_input=build_tool_result_final_prompt(
				user_task=user_task,
				tool_result=tool_result
			),
			temperature=float( st.session_state.get( 'temperature', 0.0 ) ),
			top_p=float( st.session_state.get( 'top_percent', 0.95 ) ),
			repeat_penalty=float( st.session_state.get( 'repeat_penalty', 1.1 ) ),
			max_tokens=int( st.session_state.get( 'max_tokens', 1024 ) ) or 1024,
			stream=False,
			output=None
		)
	except Exception as e:
		return f'Final answer generation failed: {e}'

def send_web_context_to_text_generation( context_text: str ) -> None:
	"""
		Purpose:
		--------
		Send fetched web context into the shared Text Generation document context buffer.

		Parameters:
		-----------
		context_text : str
			Web context text.

		Returns:
		--------
		None
	"""
	context_value = str( context_text or '' ).strip( )
	if not context_value:
		return
	
	existing_docs = st.session_state.get( 'basic_docs', [ ] )
	if not isinstance( existing_docs, list ):
		existing_docs = [ ]
	
	existing_docs.append( context_value )
	st.session_state[ 'basic_docs' ] = existing_docs
	st.session_state[ 'use_document_context' ] = True
	
initialize_model_mode_state( )

# ==============================================================================
# UTILITIES
# ==============================================================================

def image_to_base64( path: str ) -> str:
    with open(path, "rb") as f:
        return base64.b64encode(f.read()).decode()

def get_model_logo_for_state( model_name: str ) -> str:
	"""
		Purpose:
		--------
		Return the logo path associated with the selected model.

		Parameters:
		-----------
		model_name : str
			Selected local model name.

		Returns:
		--------
		str
			Configured logo path.
	"""
	model_logo_map: Dict[ str, str ] = {
			'Bro': getattr( cfg, 'BRO_LOGO', '' ),
			'Gipity': getattr( cfg, 'GIPITY_LOGO', '' ),
			'Buddy': getattr( cfg, 'BUDDY_LOGO', '' ),
			'Boo': getattr( cfg, 'BOO_LOGO', '' ),
			'Jimi': getattr( cfg, 'JIMI_LOGO', '' ),
			'Leeroy': getattr( cfg, 'LEEROY_LOGO', '' ),
			'Nisty': getattr( cfg, 'NISTY_LOGO', '' )
	}
	
	return str( model_logo_map.get( str( model_name or '' ), '' ) or '' )

def resolve_resource_path( path: str ) -> Path:
	"""
		Purpose:
		--------
		Resolve a configured resource path relative to the application base directory
		when the supplied path is not already absolute.

		Parameters:
		-----------
		path : str
			Configured resource path.

		Returns:
		--------
		Path
			Resolved resource path.
	"""
	path_value = str( path or '' ).strip( )
	if not path_value:
		return Path( '' )
	
	resource_path = Path( path_value )
	if resource_path.is_absolute( ):
		return resource_path
	
	return Path( cfg.BASE_DIR ) / resource_path

def render_selected_model_logo( model_name: str, size: str='large' ) -> None:
	"""
		Purpose:
		--------
		Render the selected model logo using Streamlit's native logo API so the logo
		remains visible when the sidebar is collapsed.

		Parameters:
		-----------
		model_name : str
			Selected local model name.

		size : str
			Streamlit logo size. Expected values are 'small', 'medium', or 'large'.

		Returns:
		--------
		None
	"""
	logo_path = get_model_logo_for_state( model_name )
	resolved_logo_path = resolve_resource_path( logo_path )
	if logo_path and resolved_logo_path.exists( ):
		st.logo( image=str( resolved_logo_path ), icon_image=str( resolved_logo_path ), size=size )
		return
	
	default_logo_path = resolve_resource_path( getattr( cfg, 'LOGO', '' ) )
	if default_logo_path.exists( ):
		st.logo( image=str( default_logo_path ), icon_image=str( default_logo_path ), size=size )

def cosine_similarity( a: np.ndarray, b: np.ndarray ) -> float:
    denom = np.linalg.norm(a) * np.linalg.norm(b)
    return float( np.dot(a, b) / denom ) if denom else 0.0

# -------- CHAT/TEXT UTILITIES --------------------

def normalize_text( text: str ) -> str:
	"""
		
		Purpose
		-------
		Normalize text by:
			• Converting to lowercase
			• Removing punctuation except sentence delimiters (. ! ?)
			• Ensuring clean sentence boundary spacing
			• Collapsing whitespace
	
		Parameters
		----------
		text: str
	
		Returns
		-------
		str
		
	"""
	if not text:
		return ""
	
	# Lowercase
	text = text.lower( )
	
	# Remove punctuation except . ! ?
	text = re.sub( r"[^\w\s\.\!\?]", "", text )
	
	# Ensure single space after sentence delimiters
	text = re.sub( r"([.!?])\s*", r"\1 ", text )
	
	# Normalize whitespace
	text = re.sub( r"\s+", " ", text ).strip( )
	
	return text

def chunk_text( text: str, size: int=None, overlap: int=None ) -> List[ str ]:
	"""
		Purpose:
		--------
		Split text into overlapping chunks using session-state defaults when explicit values
		are not provided.

		Parameters:
		-----------
		text : str
		size : int | None
		overlap : int | None

		Returns:
		--------
		List[str]
	"""
	if not text:
		return [ ]
	
	chunk_size = int(
		size if size is not None else st.session_state.get( 'retrieval_chunk_size', 1200 ) )
	chunk_overlap = int(
		overlap if overlap is not None else st.session_state.get( 'retrieval_chunk_overlap', 200 ) )
	
	if chunk_size <= 0:
		chunk_size = 1200
	
	if chunk_overlap < 0:
		chunk_overlap = 0
	
	if chunk_overlap >= chunk_size:
		chunk_overlap = max( 0, chunk_size // 4 )
	
	chunks: List[ str ] = [ ]
	i = 0
	step = max( 1, chunk_size - chunk_overlap )

	while i < len( text ):
		chunk = text[ i:i + chunk_size ]
		if chunk and chunk.strip( ):
			chunks.append( chunk )
		i += step
	
	return chunks

def convert_xml( text: str ) -> str:
	"""
		
			Purpose:
			_________
			Convert XML-delimited prompt text into Markdown by treating XML-like
			tags as section delimiters, not as strict XML.
	
			Parameters:
			-----------
			text (str) - Prompt text containing XML-like opening and closing tags.
	
			Returns:
			---------
			Markdown-formatted text using level-2 headings (##).
	"""
	markdown_blocks: List[ str ] = [ ]
	for match in cfg.XML_BLOCK_PATTERN.finditer( text ):
		raw_tag: str = match.group( "tag" )
		body: str = match.group( "body" ).strip( )
		
		# Humanize tag name for Markdown heading
		heading: str = raw_tag.replace( "_", " " ).replace( "-", " " ).title( )
		markdown_blocks.append( f"## {heading}" )
		if body:
			markdown_blocks.append( body )
	return "\n\n".join( markdown_blocks )

def convert_markdown( text: Any ) -> str:
	"""
		Purpose:
		--------
		Convert between Markdown headings and simple XML-like heading tags.
	
		Behavior:
		---------
		Auto-detects direction:
		  - If <h1>...</h1> / <h2>...</h2> ... exist, converts to Markdown (# / ## / ###).
		  - Otherwise converts Markdown headings (# / ## / ###) to <hN>...</hN> tags.
	
		Parameters:
		-----------
		text : Any
			Source text. Non-string values return "".
	
		Returns:
		--------
		str
			Converted text.
	"""
	if not isinstance( text, str ) or not text.strip( ):
		return ""
	
	# Normalize newlines
	src = text.replace( "\r\n", "\n" ).replace( "\r", "\n" )
	
	htag_pattern = re.compile( r"<h([1-6])>(.*?)</h\1>", flags=re.IGNORECASE | re.DOTALL )
	md_heading_pattern = re.compile( r"^(#{1,6})[ \t]+(.+?)[ \t]*$", flags=re.MULTILINE )
	
	# ------------------------------------------------------------------
	# Direction detection
	# ------------------------------------------------------------------
	contains_htags = bool( htag_pattern.search( src ) )
	
	# ------------------------------------------------------------------
	# XML-like heading tags -> Markdown headings
	# ------------------------------------------------------------------
	if contains_htags:
		def _htag_to_md( match: re.Match ) -> str:
			level = int( match.group( 1 ) )
			content = match.group( 2 ).strip( )
			
			# Preserve inner newlines safely by collapsing interior whitespace
			# while keeping content readable.
			content = re.sub( r"[ \t]+\n", "\n", content )
			content = re.sub( r"\n[ \t]+", "\n", content )
			
			return f"{'#' * level} {content}"
		
		out = htag_pattern.sub( _htag_to_md, src )
		return out.strip( )
	
	# ------------------------------------------------------------------
	# Markdown headings -> XML-like heading tags
	# ------------------------------------------------------------------
	def _md_to_htag( match: re.Match ) -> str:
		hashes = match.group( 1 )
		content = match.group( 2 ).strip( )
		level = len( hashes )
		return f"<h{level}>{content}</h{level}>"
	
	out = md_heading_pattern.sub( _md_to_htag, src )
	return out.strip( )

def inject_response_css( ) -> None:
	"""
	
		Purpose:
		_________
		Set the the format via css.
		
	"""
	st.markdown(
		"""
		<style>
		/* Chat message text */
		.stChatMessage p {
			color: rgb(220, 220, 220);
			font-size: 1rem;
			line-height: 1.6;
		}

		/* Headings inside chat responses */
		.stChatMessage h1 {
			color: rgb(0, 120, 252); /* DoD Blue */
			font-size: 1.6rem;
		}

		.stChatMessage h2 {
			color: rgb(0, 120, 252);
			font-size: 1.35rem;
		}

		.stChatMessage h3 {
			color: rgb(0, 120, 252);
			font-size: 1.15rem;
		}
		
		.stChatMessage a {
			color: rgb(0, 120, 252); /* DoD Blue */
			text-decoration: underline;
		}
		
		.stChatMessage a:hover {
			color: rgb(80, 160, 255);
		}

		</style>
		""", unsafe_allow_html=True )

def style_subheaders( ) -> None:
	"""
	
		Purpose:
		_________
		Sets the style of subheaders in the main UI
		
	"""
	st.markdown(
		"""
		<style>
		div[data-testid="stMarkdownContainer"] h2,
		div[data-testid="stMarkdownContainer"] h3,
		div[data-testid="stChatMessage"] div[data-testid="stMarkdownContainer"] h2,
		div[data-testid="stChatMessage"] div[data-testid="stMarkdownContainer"] h3 {
			color: rgb(0, 120, 252) !important;
		}
		</style>
		""",
		unsafe_allow_html=True, )

def save_message( role: str, content: str ) -> None:
	with sqlite3.connect( cfg.DB_PATH ) as conn:
		conn.execute( 'INSERT INTO chat_history (role, content) VALUES (?, ?)', (role, content) )

def load_history( ) -> List[ Tuple[ str, str ] ]:
	with sqlite3.connect( cfg.DB_PATH ) as conn:
		return conn.execute( 'SELECT role, content FROM chat_history ORDER BY id' ).fetchall( )

def clear_history( ) -> None:
	with sqlite3.connect( cfg.DB_PATH ) as conn:
		conn.execute( "DELETE FROM chat_history" )

#-------- PROMPT ENGINEERING UTILITIES ----------------

def fetch_prompt_names( db_path: str ) -> list[ str ]:
	"""
		Purpose:
		--------
		Retrieve template names from Prompts table.
	
		Parameters:
		-----------
		db_path : str
			SQLite database path.
	
		Returns:
		--------
		list[str]
			Sorted prompt names.
	"""
	try:
		conn = sqlite3.connect( db_path )
		cur = conn.cursor( )
		cur.execute( "SELECT Caption FROM Prompts ORDER BY PromptsId;" )
		rows = cur.fetchall( )
		conn.close( )
		return [ r[ 0 ] for r in rows if r and r[ 0 ] is not None ]
	except Exception:
		return [ ]

def fetch_prompt_text( db_path: str, name: str ) -> str | None:
	"""
		Purpose:
		--------
		Retrieve template text by name.
	
		Parameters:
		-----------
		db_path : str
			SQLite database path.
		name : str
			Template name.
	
		Returns:
		--------
		str | None
			Prompt text if found.
	"""
	try:
		conn = sqlite3.connect( db_path )
		cur = conn.cursor( )
		cur.execute( "SELECT Text FROM Prompts WHERE Caption = ?;", (name,) )
		row = cur.fetchone( )
		conn.close( )
		return str( row[ 0 ] ) if row and row[ 0 ] is not None else None
	except Exception:
		return None

def fetch_prompts_df( ) -> pd.DataFrame:
	with sqlite3.connect( cfg.DB_PATH ) as conn:
		df = pd.read_sql_query(
			"SELECT PromptsId, Caption,  Name, Version, ID FROM Prompts ORDER BY PromptsId DESC",
			conn )
	df.insert( 0, "Selected", False )
	return df

def fetch_prompt_by_id( pid: int ) -> Dict[ str, Any ] | None:
	with sqlite3.connect( cfg.DB_PATH ) as conn:
		cur = conn.execute(
			"SELECT PromptsId, Caption, Name, Text, Version, ID FROM Prompts WHERE PromptsId=?",
			(pid,)
		)
		row = cur.fetchone( )
		return dict( zip( [ c[ 0 ] for c in cur.description ], row ) ) if row else None

def fetch_prompt_by_name( name: str ) -> Dict[ str, Any ] | None:
	with sqlite3.connect( cfg.DB_PATH ) as conn:
		cur = conn.execute(
			"SELECT PromptsId, Caption, Name, Text, Version, ID FROM Prompts WHERE Caption=?",
			(name,)
		)
		row = cur.fetchone( )
		return dict( zip( [ c[ 0 ] for c in cur.description ], row ) ) if row else None

def insert_prompt( data: Dict[ str, Any ] ) -> None:
	with sqlite3.connect( cfg.DB_PATH ) as conn:
		conn.execute(
			'INSERT INTO Prompts (Caption, Name, Text, Version, ID) VALUES (?, ?, ?, ?, ?)',
			(data[ 'Caption' ], data[ 'Name' ], data[ 'Text' ], data[ 'Version' ], data[ 'ID' ]) )
		
def update_prompt( pid: int, data: Dict[ str, Any ] ) -> None:
	with sqlite3.connect( cfg.DB_PATH ) as conn:
		conn.execute(
			"UPDATE Prompts SET Caption=?, Name=?, Text=?, Version=?, ID=? WHERE PromptsId=?",
			(data[ "Caption" ], data[ "Name" ], data[ "Text" ], data[ "Version" ], data[ "ID" ],
			 pid)
		)

def delete_prompt( pid: int ) -> None:
	with sqlite3.connect( cfg.DB_PATH ) as conn:
		conn.execute( "DELETE FROM Prompts WHERE PromptsId=?", (pid,) )

def get_effective_system_instructions( ) -> str:
	"""
		Purpose:
		--------
		Return the authoritative system instructions text from session state.

		Parameters:
		-----------
		None

		Returns:
		--------
		str
	"""
	text = st.session_state.get( 'system_instructions', '' )
	return str( text ).strip( ) if text is not None else ''

def build_task_instruction_block( ) -> str:
	"""
		Purpose:
		--------
		Build a task-specific instruction block for Text Generation mode, including
		model-gated Thinking, Coding, and Function Calling directives for Gemma 4 and
		GPT-OSS-aligned local models.

		Parameters:
		-----------
		None

		Returns:
		--------
		str
			Task instruction block.
	"""
	task_preset = str( st.session_state.get( 'task_preset', 'Chat' ) or 'Chat' ).strip( )
	response_format = str(
		st.session_state.get( 'response_format', 'Markdown' ) or 'Markdown' ).strip( )
	reasoning_depth = str(
		st.session_state.get( 'reasoning_depth', 'Medium' ) or 'Medium' ).strip( )
	answer_only = bool( st.session_state.get( 'answer_only', False ) )
	use_self_check = bool( st.session_state.get( 'use_self_check', False ) )
	deterministic_reasoning = bool( st.session_state.get( 'deterministic_reasoning', False ) )
	coding_language = str(
		st.session_state.get( 'coding_language', 'Python' ) or 'Python' ).strip( )
	coding_task = str( st.session_state.get( 'coding_task', 'Generate' ) or 'Generate' ).strip( )
	coding_include_comments = bool( st.session_state.get( 'coding_include_comments', True ) )
	coding_editor_format = bool( st.session_state.get( 'coding_editor_format', True ) )
	coding_fenced_output = bool( st.session_state.get( 'coding_fenced_output', True ) )
	translation_target_language = (
			str( st.session_state.get(
				'translation_target_language', 'English' ) or 'English' ).strip( ))
	
	thinking_mode_enabled = bool( st.session_state.get( 'thinking_mode_enabled', False ) )
	thinking_effort = str(
		st.session_state.get( 'thinking_effort', 'Medium' ) or 'Medium' ).strip( )
	thinking_summary_enabled = bool(
		st.session_state.get( 'thinking_summary_enabled', True ) )
	coding_mode_enabled = bool( st.session_state.get( 'coding_mode_enabled', False ) )
	coding_test_request = bool( st.session_state.get( 'coding_test_request', False ) )
	coding_explain_request = bool( st.session_state.get( 'coding_explain_request', False ) )
	function_call_enabled = bool( st.session_state.get( 'function_call_enabled', False ) )
	function_schema_text = str(
		st.session_state.get( 'function_schema_text', '' ) or '' ).strip( )
	
	lines: List[ str ] = [ ]
	lines.append( 'Task Preset:' )
	lines.append( f'- Active Task: {task_preset}' )
	lines.append( f'- Response Format: {response_format}' )
	
	if task_preset == 'Reasoning':
		lines.append( f'- Reasoning Depth: {reasoning_depth}' )
		lines.append(
			'- Use a careful analytical process internally and return a clear final answer.' )
		
		if answer_only:
			lines.append( '- Return the final answer without extra prefatory narration.' )
		if use_self_check:
			lines.append( '- Verify the conclusion against the prompt before answering.' )
		if deterministic_reasoning:
			lines.append( '- Prefer stable, conservative reasoning over creative variation.' )
	
	elif task_preset == 'Coding':
		lines.append( f'- Code Language: {coding_language}' )
		lines.append( f'- Coding Task: {coding_task}' )
		if coding_include_comments:
			lines.append(
				'- Include documentation comments and useful inline comments when appropriate.' )
		else:
			lines.append( '- Minimize comments unless required for clarity.' )
		if coding_editor_format:
			lines.append(
				'- Format the output as editor-ready source code, not as explanatory pseudo-code.' )
		if coding_fenced_output:
			lines.append(
				'- Return code inside fenced markdown code blocks when code is produced.' )
		else:
			lines.append(
				'- Return raw code without fenced markdown blocks when code is produced.' )
	
	elif task_preset == 'Translation':
		lines.append( f'- Translate the user content into {translation_target_language}.' )
		lines.append( '- Preserve original meaning, tone, and structure where practical.' )
	
	elif task_preset == 'Summarization':
		lines.append( '- Summarize the user content clearly and faithfully.' )
		lines.append( '- Preserve key facts, names, dates, and conclusions.' )
	
	elif task_preset == 'Extraction':
		lines.append( '- Extract the requested facts faithfully and do not invent missing values.' )
		if response_format == 'JSON':
			lines.append( '- Return valid JSON only.' )
	
	else:
		lines.append( '- Respond as a general-purpose assistant.' )
	
	if thinking_mode_enabled and model_supports_capability( 'thinking' ):
		lines.append( '' )
		lines.append( 'Thinking Capability:' )
		lines.append( f'- Thinking Effort: {thinking_effort}' )
		lines.append(
			'- Use private internal reasoning to solve the task, but do not expose hidden '
			'chain-of-thought.' )
		
		if thinking_summary_enabled:
			lines.append(
				'- Provide a concise reasoning summary only when it improves the usefulness '
				'of the final answer.' )
		else:
			lines.append(
				'- Return the final answer without a separate reasoning summary.' )
	
	elif thinking_mode_enabled:
		lines.append( '' )
		lines.append( 'Thinking Capability:' )
		lines.append(
			f'- {get_selected_model_name( )} does not advertise the Thinking capability. '
			'Use the standard reasoning controls only.' )
	
	if coding_mode_enabled and model_supports_capability( 'coding' ):
		lines.append( '' )
		lines.append( 'Advanced Coding Capability:' )
		lines.append( f'- Primary Language: {coding_language}' )
		lines.append( f'- Requested Coding Operation: {coding_task}' )
		lines.append(
			'- Preserve existing application structure and generate paste-ready source code.' )
		
		if coding_test_request:
			lines.append(
				'- Include a minimal verification or test strategy when appropriate.' )
		
		if coding_explain_request:
			lines.append(
				'- Include a concise explanation of the implementation after the code.' )
	
	elif coding_mode_enabled:
		lines.append( '' )
		lines.append( 'Advanced Coding Capability:' )
		lines.append(
			f'- {get_selected_model_name( )} does not advertise the advanced Coding '
			'capability. Use the standard coding task controls only.' )
	
	if function_call_enabled and model_supports_capability( 'function_calling' ):
		lines.append( '' )
		lines.append( 'Function Calling Capability:' )
		lines.append(
			'- When a function call is required, return a single strict JSON object and no '
			'extra prose.' )
		lines.append(
			'- The JSON object must use this shape: '
			'{"name":"function_name","arguments":{...}}' )
		lines.append(
			'- Do not invent functions. Use only the function schema supplied below.' )
		
		if is_gipity_model( ):
			lines.append(
				'- For GPT-OSS/Gipity, treat function calling as an app-mediated tool call. '
				'The app will validate and execute only allowlisted functions.' )
		
		if is_gemma4_model( ):
			lines.append(
				'- For Gemma 4 models, generate a valid function-call object only when the '
				'user request clearly requires structured tool invocation.' )
		
		if function_schema_text:
			lines.append( '' )
			lines.append( 'Available Function Schema:' )
			lines.append( function_schema_text )
	
	elif function_call_enabled:
		lines.append( '' )
		lines.append( 'Function Calling Capability:' )
		lines.append(
			f'- {get_selected_model_name( )} does not advertise function calling. '
			'Answer normally and do not emit tool-call JSON.' )
	
	return '\n'.join( lines ).strip( )

def build_effective_prompt_preview( user_input: str ) -> str:
	"""
		Purpose:
		--------
		Build a readable preview of the effective prompt content used for generation.

		Parameters:
		-----------
		user_input : str

		Returns:
		--------
		str
	"""
	system_instructions = get_effective_system_instructions( )
	task_block = build_task_instruction_block( )
	preview_parts: List[ str ] = [ ]
	
	if system_instructions:
		preview_parts.append( '[System Instructions]' )
		preview_parts.append( system_instructions )
	
	if task_block:
		preview_parts.append( '[Task Instructions]' )
		preview_parts.append( task_block )
	
	preview_parts.append( '[User Input]' )
	preview_parts.append( user_input or '' )
	
	return '\n\n'.join( preview_parts ).strip( )

# ==============================================================================
# SYSTEM INSTRUCTIONS RENDERER
# ==============================================================================

def get_preset_system_instruction( task_preset: str ) -> str:
	"""
		Purpose:
		--------
		Return a starter system-instruction preset for the selected task type.

		Parameters:
		-----------
		task_preset : str
			Selected task preset.

		Returns:
		--------
		str
			System-instruction preset text.
	"""
	preset_name = str( task_preset or 'Chat' ).strip( )
	
	preset_map: Dict[ str, str ] = {
			'Chat':
				'You are Loca, a helpful local assistant. Be accurate, practical, and concise.',
			'Reasoning':
				'Solve the task carefully. Use a careful internal process, then return a clear '
				'final answer.',
			'Coding':
				'Produce correct, editor-ready code. Preserve the requested language, structure, '
				'and implementation intent.',
			'Translation':
				'Translate faithfully while preserving meaning, tone, and structure.',
			'Summarization':
				'Summarize faithfully and preserve key facts, names, dates, and conclusions.',
			'Extraction':
				'Extract only supported facts. Do not invent missing values.'
	}
	
	return preset_map.get( preset_name, preset_map[ 'Chat' ] )

def get_system_instruction_action_key( prefix: str ) -> str:
	"""
		Purpose:
		--------
		Return the pending action key used by a specific System Instructions renderer.

		Parameters:
		-----------
		prefix : str
			Renderer prefix.

		Returns:
		--------
		str
			Pending action session-state key.
	"""
	return f'{prefix}_pending_system_instruction_action'

def get_system_instruction_template_key( prefix: str ) -> str:
	"""
		Purpose:
		--------
		Return the template-select widget key used by a specific System Instructions
		renderer.

		Parameters:
		-----------
		prefix : str
			Renderer prefix.

		Returns:
		--------
		str
			Template widget session-state key.
	"""
	return f'{prefix}_instructions_template'

def get_system_instruction_pending_template_key( prefix: str ) -> str:
	"""
		Purpose:
		--------
		Return the pending template key used by a specific System Instructions renderer.

		Parameters:
		-----------
		prefix : str
			Renderer prefix.

		Returns:
		--------
		str
			Pending template session-state key.
	"""
	return f'{prefix}_pending_system_template_name'

def request_system_template_change( prefix: str ) -> None:
	"""
		Purpose:
		--------
		Request a system-instruction template change without directly modifying the
		widget-owned system_instructions key.

		Parameters:
		-----------
		prefix : str
			Renderer prefix.

		Returns:
		--------
		None
	"""
	template_key = get_system_instruction_template_key( prefix )
	pending_key = get_system_instruction_pending_template_key( prefix )
	name = st.session_state.get( template_key, None )
	
	if name:
		st.session_state[ pending_key ] = str( name )

def request_system_instruction_action( prefix: str, action: str ) -> None:
	"""
		Purpose:
		--------
		Request a pending system-instruction action without directly modifying the
		widget-owned system_instructions key.

		Parameters:
		-----------
		prefix : str
			Renderer prefix.

		action : str
			Requested action name.

		Returns:
		--------
		None
	"""
	action_key = get_system_instruction_action_key( prefix )
	st.session_state[ action_key ] = str( action or '' )

def process_pending_system_instruction_requests( prefix: str ) -> None:
	"""
		Purpose:
		--------
		Process pending System Instructions requests before the system_instructions text
		area is instantiated. This is the only safe place to modify the shared
		system_instructions widget-owned key.

		Parameters:
		-----------
		prefix : str
			Renderer prefix.

		Returns:
		--------
		None
	"""
	template_key = get_system_instruction_template_key( prefix )
	pending_template_key = get_system_instruction_pending_template_key( prefix )
	action_key = get_system_instruction_action_key( prefix )
	
	pending_template_name = st.session_state.pop( pending_template_key, None )
	if pending_template_name:
		template_text = fetch_prompt_text( cfg.DB_PATH, str( pending_template_name ) )
		
		if template_text is not None:
			st.session_state[ 'system_instructions' ] = str( template_text )
			st.session_state[ 'active_prompt_caption' ] = str( pending_template_name )
	
	pending_action = st.session_state.pop( action_key, None )
	if pending_action == 'clear':
		st.session_state[ 'system_instructions' ] = ''
		st.session_state[ 'active_prompt_caption' ] = ''
		
		if template_key in st.session_state:
			del st.session_state[ template_key ]
	
	elif pending_action == 'convert':
		text = st.session_state.get( 'system_instructions', '' )
		
		if isinstance( text, str ) and text.strip( ):
			src = text.strip( )
			
			if cfg.XML_BLOCK_PATTERN.search( src ):
				converted = convert_xml( src )
			else:
				converted = convert_markdown( src )
			
			st.session_state[ 'system_instructions' ] = converted
	
	elif pending_action == 'apply_preset':
		task_preset = str( st.session_state.get( 'task_preset', 'Chat' ) or 'Chat' ).strip( )
		st.session_state[ 'system_instructions' ] = get_preset_system_instruction( task_preset )
		st.session_state[ 'active_prompt_caption' ] = f'{task_preset} Preset'

def render_system_instructions( prefix: str, include_apply_preset: bool = False,
		include_preview: bool = False ) -> None:
	"""
		Purpose:
		--------
		Render a Streamlit-safe shared System Instructions control surface. All writes to
		the widget-owned system_instructions key are processed before the text-area widget
		is instantiated.

		Parameters:
		-----------
		prefix : str
			Unique renderer prefix, such as 'text' or 'docqna'.

		include_apply_preset : bool
			When True, show the Apply Preset button.

		include_preview : bool
			When True, show the Preview Prompt button and preview text area.

		Returns:
		--------
		None
	"""
	process_pending_system_instruction_requests( prefix )
	
	template_key = get_system_instruction_template_key( prefix )
	prompt_names = fetch_prompt_names( cfg.DB_PATH )
	
	if not prompt_names:
		prompt_names = [ '' ]
	
	in_left, in_right = st.columns( [ 0.8, 0.2 ] )
	
	with in_left:
		st.text_area( label='Enter Text', height=120,
			width='stretch', help=cfg.SYSTEM_INSTRUCTIONS, key='system_instructions' )
	
	with in_right:
		st.selectbox( label='Use Template', options=prompt_names,
			index=None, key=template_key,
			on_change=request_system_template_change, args=(prefix,) )
	
	if include_apply_preset and include_preview:
		btn_c1, btn_c2, btn_c3, btn_c4 = st.columns( [ 0.35, 0.2, 0.2, 0.25 ] )
		
		with btn_c1:
			st.button( label='Clear Instructions', width='stretch',
				on_click=request_system_instruction_action, args=(prefix, 'clear') )
		
		with btn_c2:
			st.button( label='XML <-> Markdown', width='stretch',
				on_click=request_system_instruction_action, args=(prefix, 'convert') )
		
		with btn_c3:
			st.button( label='Apply Preset', width='stretch',
				on_click=request_system_instruction_action, args=(prefix, 'apply_preset') )
		
		with btn_c4:
			if st.button( label='Preview Prompt', width='stretch', key=f'{prefix}_preview_prompt' ):
				st.session_state[ 'preview_effective_prompt' ] = not bool(
					st.session_state.get( 'preview_effective_prompt', False ) )
	
	else:
		btn_c1, btn_c2 = st.columns( [ 0.8, 0.2 ] )
		
		with btn_c1:
			st.button( label='Clear Instructions', width='stretch',
				on_click=request_system_instruction_action, args=(prefix, 'clear') )
		
		with btn_c2:
			st.button( label='XML <-> Markdown', width='stretch', on_click=request_system_instruction_action,
				args=(prefix, 'convert') )
	
	if include_preview and bool( st.session_state.get( 'preview_effective_prompt', False ) ):
		user_preview_input = str( st.session_state.get( 'last_preview_input', '' ) or '' )
		
		st.text_area( label='Effective Prompt Preview',
			value=build_effective_prompt_preview( user_preview_input ), height=220,
			disabled=True, key=f'{prefix}_effective_prompt_preview' )
		
def get_runtime_llm( ) -> Any | None:
	"""
		Purpose:
		--------
		Load the selected llama.cpp model using the currently selected model path and
		runtime settings.

		Parameters:
		-----------
		None

		Returns:
		--------
		Any | None
			Loaded llama.cpp model instance when available; otherwise None.
	"""
	synchronize_model_derived_state( )
	
	model_path = str(
		st.session_state.get(
			'selected_model_path',
			get_model_path_for_state( get_selected_model_name( ) )
		) or get_model_path_for_state( get_selected_model_name( ) )
	)
	
	ctx_value = int( st.session_state.get( 'context_window', cfg.DEFAULT_CTX ) or cfg.DEFAULT_CTX )
	thread_value = int( st.session_state.get( 'cpu_threads', cfg.CORES ) or cfg.CORES )
	seed_value = int( st.session_state.get( 'random_seed', -1 ) or -1 )
	
	if ctx_value <= 0:
		ctx_value = int( cfg.DEFAULT_CTX )
	
	if thread_value <= 0:
		thread_value = int( cfg.CORES )
	
	return load_llm( model_path=model_path, ctx=ctx_value,
		threads=thread_value, seed=seed_value )

def build_prompt( user_input: str ) -> str:
	"""
		Purpose:
		--------
		Build a llama.cpp-compatible prompt using unified system instructions, task-specific
		Text Generation settings, optional semantic/basic context, and chat history. Semantic
		context retrieval is guarded so Text Generation cannot crash when the embedder,
		database, embeddings table, or stored vectors are unavailable or inconsistent.

		Parameters:
		-----------
		user_input : str
			User prompt text supplied by the Text Generation mode.

		Returns:
		--------
		str
			Prompt text formatted for the local llama.cpp chat template.
	"""
	global embedder
	
	system_instructions = get_effective_system_instructions( )
	task_block = build_task_instruction_block( )
	use_semantic = bool( st.session_state.get( 'use_semantic', False ) )
	use_chat_history = bool( st.session_state.get( 'use_chat_history', True ) )
	use_document_context = bool( st.session_state.get( 'use_document_context', False ) )
	basic_docs = st.session_state.get( 'basic_docs', [ ] )
	messages = st.session_state.get( 'messages', [ ] )
	user_text = str( user_input or '' )
	
	top_k_value = int( st.session_state.get( 'top_k', 0 ) or 0 )
	if top_k_value <= 0:
		top_k_value = 4
	
	system_parts: List[ str ] = [ ]
	if system_instructions:
		system_parts.append( str( system_instructions ) )
	if task_block:
		system_parts.append( str( task_block ) )
	
	system_text = '\n\n'.join( [ p for p in system_parts if p ] ).strip( )
	
	prompt = ''
	if system_text:
		prompt += f'<|system|>\n{system_text}\n</s>\n'
	
	if use_semantic:
		try:
			if not is_embedder_available( globals( ).get( 'embedder', None ) ):
				st.session_state[ 'semantic_status' ] = get_embedder_unavailable_message( )
			else:
				with sqlite3.connect( cfg.DB_PATH ) as conn:
					rows = conn.execute(
						'SELECT chunk, vector FROM embeddings' ).fetchall( )
				
				if rows:
					q_raw = embedder.encode( [ user_text ], show_progress_bar=False )[ 0 ]
					q = np.asarray( q_raw, dtype=np.float32 ).reshape( -1 )
					scored: List[ Tuple[ str, float ] ] = [ ]
					
					for chunk, vector_blob in rows:
						if not chunk or vector_blob is None:
							continue
						
						vector = np.frombuffer( vector_blob, dtype=np.float32 )
						if vector.size == 0 or vector.size != q.size:
							continue
						
						score = cosine_similarity( q, vector )
						scored.append( (str( chunk ), score) )
					
					if scored:
						scored = sorted( scored, key=lambda x: x[ 1 ], reverse=True )
						for chunk, score in scored[ :top_k_value ]:
							prompt += f'<|system|>\nSemantic Context:\n{chunk}\n</s>\n'
						
						st.session_state[ 'semantic_status' ] = (
								f'Loaded {min( len( scored ), top_k_value )} semantic context '
								f'chunk(s) for Text Generation.')
					else:
						st.session_state[ 'semantic_status' ] = (
								'Semantic context was enabled, but no compatible embedding '
								'vectors were available for Text Generation.')
		except Exception as e:
			st.session_state[ 'semantic_status' ] = (
					f'Semantic context was skipped because retrieval failed: {e}')
	
	if use_document_context and isinstance( basic_docs, list ):
		for document_text in basic_docs[ :6 ]:
			if document_text:
				prompt += f'<|system|>\nDocument Context:\n{document_text}\n</s>\n'
	
	if use_chat_history and isinstance( messages, list ):
		for msg in messages:
			role = ''
			content = ''
			
			if isinstance( msg, tuple ) or isinstance( msg, list ):
				if len( msg ) == 2:
					role = str( msg[ 0 ] or '' ).strip( )
					content = str( msg[ 1 ] or '' )
			elif isinstance( msg, dict ):
				role = str( msg.get( 'role', '' ) or '' ).strip( )
				content = str( msg.get( 'content', '' ) or '' )
			
			if role in ('user', 'assistant', 'system'):
				prompt += f'<|{role}|>\n{content}\n</s>\n'
	
	prompt += f'<|user|>\n{user_text}\n</s>\n<|assistant|>\n'
	return prompt

def build_llama_call_args( max_tokens: int, temperature: float, top_p: float,
		repeat_penalty: float, stream: bool ) -> Dict[ str, Any ]:
	"""
		Purpose:
		--------
		Build llama.cpp generation arguments from the current Streamlit runtime settings.

		Parameters:
		-----------
		max_tokens : int
			Maximum number of generated tokens.

		temperature : float
			Sampling temperature.

		top_p : float
			Nucleus sampling value.

		repeat_penalty : float
			Repeat penalty value.

		stream : bool
			Whether streaming output is requested.

		Returns:
		--------
		Dict[str, Any]
			Generation argument dictionary for llama.cpp.
	"""
	max_token_value = int( max_tokens ) if int( max_tokens or 0 ) > 0 else 1024
	temperature_value = float( temperature ) if temperature is not None else 0.0
	top_p_value = float( top_p ) if top_p is not None else 0.95
	repeat_penalty_value = float( repeat_penalty ) if repeat_penalty is not None else 1.1
	top_k_value = int( st.session_state.get( 'top_k', 0 ) or 0 )
	repeat_window_value = int( st.session_state.get( 'repeat_window', 0 ) or 0 )
	
	call_args: Dict[ str, Any ] = {
			'stream': bool( stream ),
			'max_tokens': max_token_value,
			'temperature': temperature_value,
			'top_p': top_p_value,
			'repeat_penalty': repeat_penalty_value,
			'stop': [ '</s>' ]
	}
	
	if top_k_value > 0:
		call_args[ 'top_k' ] = top_k_value
	
	if repeat_window_value > 0:
		call_args[ 'repeat_last_n' ] = repeat_window_value
	
	return call_args

def get_missing_model_message( ) -> str:
	"""
		Purpose:
		--------
		Build a clear user-facing message when the selected local GGUF model or required
		llama.cpp dependency is not available.

		Parameters:
		-----------
		None

		Returns:
		--------
		str
			Model availability message.
	"""
	model_name = get_selected_model_name( )
	model_path = get_selected_model_path( )
	
	if Llama is None:
		return (
				'Local model inference is unavailable because `llama-cpp-python` is not '
				'installed or could not be imported.\n\n'
				f'**Selected Model:** {model_name}\n\n'
				'Install or repair `llama-cpp-python`, then restart the Streamlit app.'
		)
	
	if model_path:
		return (
				'The selected local GGUF model is unavailable.\n\n'
				f'**Model:** {model_name}\n\n'
				f'**Configured Path:** `{model_path}`\n\n'
				'Verify that the file exists or update the model path in `config.py` '
				'or the corresponding environment variable.'
		)
	
	return (
			'The selected local GGUF model is not configured.\n\n'
			f'**Model:** {model_name}\n\n'
			'Update the model registry path in `config.py` or provide the corresponding '
			'environment variable.'
	)

def run_llm_turn( user_input: str, temperature: float, top_p: float, repeat_penalty: float,
		max_tokens: int, stream: bool, output: Any = None ) -> str:
	"""
		Purpose:
		--------
		Run a single local llama.cpp LLM turn using the selected model, current prompt
		contract, and runtime settings. Missing or unavailable GGUF models are handled
		safely with a user-facing diagnostic response instead of a callable None failure.

		Parameters:
		-----------
		user_input : str
			User prompt or prepared application prompt.

		temperature : float
			Sampling temperature.

		top_p : float
			Nucleus sampling value.

		repeat_penalty : float
			Repeat penalty value.

		max_tokens : int
			Maximum number of generated tokens.

		stream : bool
			Whether to stream output tokens into the supplied Streamlit placeholder.

		output : Any | None
			Optional Streamlit output placeholder.

		Returns:
		--------
		str
			Generated response text or diagnostic message.
	"""
	if user_input is None:
		return ''
	
	user_text = str( user_input or '' ).strip( )
	if not user_text:
		return ''
	
	synchronize_model_derived_state( )
	
	runtime_llm = get_runtime_llm( )
	if runtime_llm is None:
		message = get_missing_model_message( )
		if output is not None:
			output.markdown( message )
		return message
	
	prompt = build_prompt( user_text )
	call_args = build_llama_call_args(
		max_tokens=max_tokens,
		temperature=temperature,
		top_p=top_p,
		repeat_penalty=repeat_penalty,
		stream=stream
	)
	
	try:
		if not stream:
			resp = runtime_llm( prompt, **call_args )
			text = (
					resp.get( 'choices', [ { 'text': '' } ] )[ 0 ]
					.get( 'text', '' ) or ''
			)
			
			return str( text ).strip( )
		
		buf = ''
		if output is None:
			output = st.empty( )
		
		for chunk in runtime_llm( prompt, **call_args ):
			try:
				token = chunk[ 'choices' ][ 0 ][ 'text' ]
			except Exception:
				token = ''
			
			if token:
				buf += token
				output.markdown( buf + '▌' )
		
		output.markdown( buf )
		return buf.strip( )
	
	except TypeError:
		fallback_args: Dict[ str, Any ] = {
				'stream': bool( stream ),
				'max_tokens': int( max_tokens ) if int( max_tokens or 0 ) > 0 else 1024,
				'temperature': float( temperature ) if temperature is not None else 0.0,
				'top_p': float( top_p ) if top_p is not None else 0.95,
				'repeat_penalty': float( repeat_penalty ) if repeat_penalty is not None else 1.1,
				'stop': [ '</s>' ]
		}
		
		if not stream:
			resp = runtime_llm( prompt, **fallback_args )
			text = ( resp.get( 'choices', [ { 'text': '' } ] )[ 0 ].get( 'text', '' ) or '' )
			
			return str( text ).strip( )
		
		buf = ''
		if output is None:
			output = st.empty( )
		
		for chunk in runtime_llm( prompt, **fallback_args ):
			try:
				token = chunk[ 'choices' ][ 0 ][ 'text' ]
			except Exception:
				token = ''
			
			if token:
				buf += token
				output.markdown( buf + '▌' )
		
		output.markdown( buf )
		return buf.strip( )
	
	except Exception as e:
		message = (
				'Local model generation failed.\n\n'
				f'**Model:** {get_selected_model_name( )}\n\n'
				f'**Path:** `{get_selected_model_path( )}`\n\n'
				f'**Error:** {e}'
		)
		
		if output is not None:
			output.markdown( message )
		
		return message

def get_prompt_categories( ) -> List[ str ]:
	"""
		Purpose:
		--------
		Return supported prompt categories.

		Parameters:
		-----------
		None

		Returns:
		--------
		List[str]
	"""
	return [
			'General Chat',
			'Reasoning',
			'Coding',
			'Translation',
			'Summarization',
			'Extraction',
			'Document Extraction',
			'OCR',
			'JSON Output'
	]

def get_prompt_task_types( ) -> List[ str ]:
	"""
		Purpose:
		--------
		Return supported task types.

		Parameters:
		-----------
		None

		Returns:
		--------
		List[str]
	"""
	return [
			'Chat',
			'Reasoning',
			'Coding',
			'Translation',
			'Summarization',
			'Extraction'
	]

def infer_prompt_category( prompt_row: Dict[ str, Any ] | None ) -> str:
	"""
		Purpose:
		--------
		Infer a prompt category from the prompt row content.

		Parameters:
		-----------
		prompt_row : Dict[str, Any] | None

		Returns:
		--------
		str
	"""
	if not isinstance( prompt_row, dict ):
		return 'General Chat'
	
	caption = str( prompt_row.get( 'Caption', '' ) or '' ).lower( )
	name = str( prompt_row.get( 'Name', '' ) or '' ).lower( )
	text = str( prompt_row.get( 'Text', '' ) or '' ).lower( )
	
	blob = f'{caption} {name} {text}'
	
	if 'json' in blob:
		return 'JSON Output'
	if 'ocr' in blob:
		return 'OCR'
	if 'document' in blob and 'extract' in blob:
		return 'Document Extraction'
	if 'extract' in blob:
		return 'Extraction'
	if 'summar' in blob:
		return 'Summarization'
	if 'translat' in blob:
		return 'Translation'
	if 'coding' in blob or 'code' in blob or 'debug' in blob or 'refactor' in blob:
		return 'Coding'
	if 'reason' in blob or 'analysis' in blob:
		return 'Reasoning'
	
	return 'General Chat'

def build_starter_prompt_template( category: str, task_type: str, response_format: str,
		language: str ) -> str:
	"""
		Purpose:
		--------
		Build a starter prompt template from high-level prompt metadata.

		Parameters:
		-----------
		category : str
		task_type : str
		response_format : str
		language : str

		Returns:
		--------
		str
	"""
	category_value = str( category or 'General Chat' ).strip( )
	task_value = str( task_type or 'Chat' ).strip( )
	format_value = str( response_format or 'Markdown' ).strip( )
	language_value = str( language or 'English' ).strip( )
	lines: List[ str ] = [ ]
	lines.append( f'You are Loca, a local AI assistant operating in the category "{category_value}".' )
	lines.append( f'Primary task type: {task_value}.' )
	lines.append( f'Response format: {format_value}.' )
	lines.append( f'Preferred language: {language_value}.' )
	
	if category_value == 'Reasoning':
		lines.append(
			'Provide careful, structured analytical answers grounded in the supplied information.' )
	elif category_value == 'Coding':
		lines.append(
			'Produce editor-ready code and explain only what is necessary for correct implementation.' )
	elif category_value == 'Translation':
		lines.append( 'Translate faithfully while preserving meaning, tone, and structure.' )
	elif category_value == 'Summarization':
		lines.append( 'Summarize faithfully and preserve key facts, names, and dates.' )
	elif category_value == 'Extraction':
		lines.append( 'Extract only supported facts. Do not invent missing values.' )
	elif category_value == 'Document Extraction':
		lines.append(
			'Use the document content as the evidence base and extract structured facts faithfully.' )
	elif category_value == 'OCR':
		lines.append( 'Extract visible text accurately and preserve structural cues where possible.' )
	elif category_value == 'JSON Output':
		lines.append( 'Return valid JSON only, matching the requested structure exactly.' )
	else:
		lines.append( 'Respond helpfully, accurately, and concisely.' )
	
	lines.append( 'If information is missing, state that clearly.' )
	return '\n'.join( lines ).strip( )

def generate_prompt_template_draft( goal: str, constraints: str, style: str,
		category: str, task_type: str, response_format: str, language: str ) -> str:
	"""
		Purpose:
		--------
		Generate a draft system prompt using the local model.

		Parameters:
		-----------
		goal : str
		constraints : str
		style : str
		category : str
		task_type : str
		response_format : str
		language : str

		Returns:
		--------
		str
	"""
	prompt = f"""
	Create a strong system prompt for the Loca local AI application.
	
	Category: {category}
	Task Type: {task_type}
	Response Format: {response_format}
	Language: {language}
	Goal: {goal}
	Constraints: {constraints}
	Style: {style}
	
	Write only the system prompt text. Do not add explanation.
	""".strip( )
	
	return run_llm_turn( user_input=prompt,
		temperature=float( st.session_state.get( 'temperature', 0.2 ) ),
		top_p=float( st.session_state.get( 'top_percent', 0.95 ) ),
		repeat_penalty=float( st.session_state.get( 'repeat_penalty', 1.05 ) ),
		max_tokens=512, stream=False, output=None )

def apply_prompt_to_text_generation( prompt_text: str ) -> None:
	"""
		Purpose:
		--------
		Apply a prompt to shared Text Generation settings.

		Parameters:
		-----------
		prompt_text : str

		Returns:
		--------
		None
	"""
	st.session_state[ 'system_instructions' ] = str( prompt_text or '' )

def apply_prompt_to_document_qna( prompt_text: str ) -> None:
	"""
		Purpose:
		--------
		Apply a prompt to shared Document Q&A settings.

		Parameters:
		-----------
		prompt_text : str

		Returns:
		--------
		None
	"""
	st.session_state[ 'system_instructions' ] = str( prompt_text or '' )
	st.session_state[ 'require_grounding' ] = True
	st.session_state[ 'answer_from_excerpts_only' ] = True

def apply_prompt_metadata_to_shared_state( category: str, task_type: str,
		response_format: str, language: str ) -> None:
	"""
		Purpose:
		--------
		Apply prompt metadata to the shared app contract.

		Parameters:
		-----------
		category : str
		task_type : str
		response_format : str
		language : str

		Returns:
		--------
		None
	"""
	st.session_state[ 'task_preset' ] = str( task_type or 'Chat' )
	st.session_state[ 'response_format' ] = str( response_format or 'Markdown' )
	st.session_state[ 'translation_target_language' ] = str( language or 'English' )

def clone_prompt_record( source_prompt: Dict[ str, Any ] | None ) -> None:
	"""
		Purpose:
		--------
		Clone a selected prompt into the edit surface as a new prompt draft.

		Parameters:
		-----------
		source_prompt : Dict[str, Any] | None

		Returns:
		--------
		None
	"""
	if not isinstance( source_prompt, dict ):
		return
	
	st.session_state.pe_selected_id = None
	st.session_state.pe_caption = f'{str( source_prompt.get( "Caption", "" ) )} Copy'.strip( )
	st.session_state.pe_name = str( source_prompt.get( 'Name', '' ) or '' )
	st.session_state.pe_text = str( source_prompt.get( 'Text', '' ) or '' )
	st.session_state.pe_version = str( source_prompt.get( 'Version', '' ) or '' )
	st.session_state.pe_id = source_prompt.get( 'ID', 0 )

# ----------- DATABASE UTILITIES -------------------------

def initialize_database( ) -> None:
	"""
		Purpose:
		--------
		Ensure required SQLite tables exist and that the Prompts table contains the
		columns required by the prompt utilities, Prompt Engineering mode, and
		AI-asset governance features.

		Parameters:
		-----------
		None

		Returns:
		--------
		None
	"""
	Path( 'stores/sqlite' ).mkdir( parents=True, exist_ok=True )
	
	with sqlite3.connect( cfg.DB_PATH ) as conn:
		conn.execute( """
            CREATE TABLE IF NOT EXISTS chat_history
            (
                id
                INTEGER
                PRIMARY
                KEY
                AUTOINCREMENT,
                role
                TEXT,
                content
                TEXT
            )
			""" )
		
		conn.execute( """
            CREATE TABLE IF NOT EXISTS embeddings
            (
                id
                INTEGER
                PRIMARY
                KEY
                AUTOINCREMENT,
                chunk
                TEXT,
                vector
                BLOB
            )
			""" )
		
		conn.execute( """
            CREATE TABLE IF NOT EXISTS Prompts
            (
                PromptsId
                INTEGER
                NOT
                NULL
                PRIMARY
                KEY
                AUTOINCREMENT,
                Caption
                TEXT,
                Name
                TEXT
            (
                80
            ),
                Text TEXT,
                Version TEXT
            (
                80
            ),
                ID TEXT
            (
                80
            )
                )
			""" )
		
		conn.execute( """
            CREATE TABLE IF NOT EXISTS documents
            (
                DocumentId
                INTEGER
                PRIMARY
                KEY
                AUTOINCREMENT,
                Name
                TEXT
                NOT
                NULL,
                Type
                TEXT,
                SizeBytes
                INTEGER,
                Source
                TEXT,
                Fingerprint
                TEXT,
                TextLength
                INTEGER,
                ChunkCount
                INTEGER,
                CreatedOn
                TEXT
            )
			""" )
		
		conn.execute( """
            CREATE TABLE IF NOT EXISTS document_chunks
            (
                ChunkId
                INTEGER
                PRIMARY
                KEY
                AUTOINCREMENT,
                DocumentName
                TEXT
                NOT
                NULL,
                ChunkIndex
                INTEGER,
                ChunkText
                TEXT,
                ChunkLength
                INTEGER,
                Fingerprint
                TEXT,
                CreatedOn
                TEXT
            )
			""" )
		
		conn.execute( """
            CREATE TABLE IF NOT EXISTS document_embeddings
            (
                EmbeddingId
                INTEGER
                PRIMARY
                KEY
                AUTOINCREMENT,
                DocumentName
                TEXT
                NOT
                NULL,
                ChunkIndex
                INTEGER,
                VectorDim
                INTEGER,
                Fingerprint
                TEXT,
                CreatedOn
                TEXT
            )
			""" )
		
		conn.execute( """
            CREATE TABLE IF NOT EXISTS images
            (
                ImageId
                INTEGER
                PRIMARY
                KEY
                AUTOINCREMENT,
                Name
                TEXT
                NOT
                NULL,
                MimeType
                TEXT,
                SizeBytes
                INTEGER,
                Fingerprint
                TEXT,
                Source
                TEXT,
                CreatedOn
                TEXT
            )
			""" )
		
		prompt_columns = [ row[ 1 ] for row in conn.execute( 'PRAGMA table_info("Prompts");' ).fetchall( ) ]
		
		if 'Caption' not in prompt_columns:
			conn.execute( 'ALTER TABLE "Prompts" ADD COLUMN "Caption" TEXT;' )
		
		conn.commit( )
		
def create_connection( ) -> sqlite3.Connection:
	return sqlite3.connect( cfg.DB_PATH )

def list_tables( ) -> List[ str ]:
	with create_connection( ) as conn:
		_query = "SELECT name FROM sqlite_master WHERE type='table' ORDER BY name;"
		rows = conn.execute( _query ).fetchall( )
		return [ r[ 0 ] for r in rows ]

def create_schema( table: str ) -> List[ Tuple ]:
	with create_connection( ) as conn:
		return conn.execute( f'PRAGMA table_info("{table}");' ).fetchall( )

def read_table( table: str, limit: int=None, offset: int=0 ) -> pd.DataFrame:
	query = f'SELECT rowid, * FROM "{table}"'
	if limit:
		query += f" LIMIT {limit} OFFSET {offset}"
	with create_connection( ) as conn:
		return pd.read_sql_query( query, conn )

def drop_table( table: str ) -> None:
	"""
		Purpose:
		--------
		Safely drop a table if it exists.
	
		Parameters:
		-----------
		table : str
			Table name.
	"""
	if not table:
		return
	
	with create_connection( ) as conn:
		conn.execute( f'DROP TABLE IF EXISTS "{table}";' )
		conn.commit( )

def rename_table( old_name: str, new_name: str ) -> None:
	"""
		Purpose:
		--------
		Rename an existing SQLite table. Attempts native ALTER TABLE rename first; if it fails,
		falls back to a schema-safe rebuild using the original CREATE TABLE statement and
		preserves indexes.

		Parameters:
		-----------
		old_name : str
			Existing table name.

		new_name : str
			New table name.

		Returns:
		--------
		None
	"""
	if not old_name or not new_name:
		return
	
	with create_connection( ) as conn:
		try:
			conn.execute( f'ALTER TABLE "{old_name}" RENAME TO "{new_name}";' )
			conn.commit( )
			return
		except Exception:
			pass
		
		row = conn.execute(
			"""
            SELECT sql
            FROM sqlite_master
            WHERE type ='table' AND name =?
			""",
			(old_name,)
		).fetchone( )
		
		if not row or not row[ 0 ]:
			raise ValueError( "Table definition not found." )
		
		create_sql = row[ 0 ]
		indexes = conn.execute(
			"""
            SELECT sql
            FROM sqlite_master
            WHERE type ='index' AND tbl_name=? AND sql IS NOT NULL
			""",
			(old_name,)
		).fetchall( )
		
		open_paren = create_sql.find( "(" )
		if open_paren == -1:
			raise ValueError( "Malformed CREATE TABLE statement." )
		
		temp_name = f"{new_name}__rebuild_temp"
		conn.execute( "BEGIN" )
		conn.execute( f'CREATE TABLE "{temp_name}" {create_sql[ open_paren: ]}' )
		cols = [ r[ 1 ] for r in conn.execute( f'PRAGMA table_info("{old_name}");' ).fetchall( ) ]
		col_list = ", ".join( [ f'"{c}"' for c in cols ] )
		
		conn.execute(
			f'INSERT INTO "{temp_name}" ({col_list}) SELECT {col_list} FROM "{old_name}";'
		)
		
		conn.execute( f'DROP TABLE "{old_name}";' )
		conn.execute( f'ALTER TABLE "{temp_name}" RENAME TO "{new_name}";' )
		
		for idx in indexes:
			idx_sql = idx[ 0 ]
			if idx_sql:
				idx_sql = idx_sql.replace( f'ON "{old_name}"', f'ON "{new_name}"' )
				conn.execute( idx_sql )
		
		conn.commit( )

def rename_column( table_name: str, old_name: str, new_name: str ) -> None:
	"""
		Purpose:
		--------
		Rename a column within an existing SQLite table. Attempts native ALTER TABLE rename
		first; if it fails, falls back to a schema-safe rebuild preserving column order, data,
		and indexes.

		Parameters:
		-----------
		table_name : str
			Table containing the column.

		old_name : str
			Existing column name.

		new_name : str
			New column name.

		Returns:
		--------
		None
	"""
	if not table_name or not old_name or not new_name:
		return
	
	with create_connection( ) as conn:
		try:
			conn.execute(
				f'ALTER TABLE "{table_name}" RENAME COLUMN "{old_name}" TO "{new_name}";'
			)
			conn.commit( )
			return
		except Exception:
			pass
		
		row = conn.execute( """
            SELECT sql
            FROM sqlite_master
            WHERE type ='table' AND name =?
			""", (table_name,) ).fetchone( )
		
		if not row or not row[ 0 ]:
			raise ValueError( "Table definition not found." )
		
		create_sql = row[ 0 ]
		indexes = conn.execute( """
            SELECT sql
            FROM sqlite_master
            WHERE type ='index' AND tbl_name=? AND sql IS NOT NULL
			""", (table_name,) ).fetchall( )
		
		schema = conn.execute( f'PRAGMA table_info("{table_name}");' ).fetchall( )
		cols = [ r[ 1 ] for r in schema ]
		if old_name not in cols:
			raise ValueError( "Column not found." )
		
		mapped_cols = [ (new_name if c == old_name else c) for c in cols ]
		temp_table = f"{table_name}__rebuild_temp"
		col_defs: List[ str ] = [ ]
		pk_cols = [ r for r in schema if int( r[ 5 ] or 0 ) > 0 ]
		single_pk = len( pk_cols ) == 1
		
		for row in schema:
			col_name = row[ 1 ]
			col_type = row[ 2 ] or ''
			not_null = int( row[ 3 ] or 0 )
			default_value = row[ 4 ]
			pk = int( row[ 5 ] or 0 )
			
			out_name = new_name if col_name == old_name else col_name
			col_def = f'"{out_name}" {col_type}'.strip( )
			
			if not_null:
				col_def += ' NOT NULL'
			
			if default_value is not None:
				col_def += f' DEFAULT {default_value}'
			
			if single_pk and pk == 1:
				col_def += ' PRIMARY KEY'
			
			col_defs.append( col_def )
		
		new_create_sql = f'CREATE TABLE "{temp_table}" ({", ".join( col_defs )});'
		
		old_select = ", ".join( [ f'"{c}"' for c in cols ] )
		new_insert = ", ".join( [ f'"{c}"' for c in mapped_cols ] )
		
		conn.execute( "BEGIN" )
		conn.execute( new_create_sql )
		conn.execute(
			f'INSERT INTO "{temp_table}" ({new_insert}) SELECT {old_select} FROM "{table_name}";'
		)
		
		conn.execute( f'DROP TABLE "{table_name}";' )
		conn.execute( f'ALTER TABLE "{temp_table}" RENAME TO "{table_name}";' )
		
		for idx in indexes:
			idx_sql = idx[ 0 ]
			if idx_sql:
				idx_sql = idx_sql.replace( f'"{old_name}"', f'"{new_name}"' )
				conn.execute( idx_sql )
		
		conn.commit( )
		
def create_index( table: str, column: str ) -> None:
	"""
		Purpose:
		--------
		Create a safe SQLite index on a specified table column.
	
		Handles:
			- Spaces in column names
			- Special characters
			- Reserved words
			- Duplicate index names
			- Validation against actual table schema
	
		Parameters:
		-----------
		table : str
			Table name.
		column : str
			Column name to index.
	"""
	if not table or not column:
		return
	
	# ----------  Validate table exists
	tables = list_tables( )
	if table not in tables:
		raise ValueError( "Invalid table name." )
	
	# ----------  Validate column exists
	schema = create_schema( table )
	valid_columns = [ col[ 1 ] for col in schema ]
	
	if column not in valid_columns:
		raise ValueError( "Invalid column name." )
	
	# ----------  Sanitize index name (identifier only)
	safe_index_name = re.sub( r"[^0-9a-zA-Z_]+", "_", f"idx_{table}_{column}" )
	
	# ----------  Create index safely (quote identifiers)
	sql = f'CREATE INDEX IF NOT EXISTS "{safe_index_name}" ON "{table}"("{column}");'
	
	with create_connection( ) as conn:
		conn.execute( sql )
		conn.commit( )

def apply_filters( df: pd.DataFrame ) -> pd.DataFrame:
	st.subheader( 'Advanced Filters' )
	col1, col2, col3 = st.columns( 3 )
	column = col1.selectbox( 'Column', df.columns )
	operator = col2.selectbox( 'Operator', [ '=', '!=', '>', '<', '>=', '<=', 'contains' ] )
	value = col3.text_input( 'Value' )
	if value:
		if operator == '=':
			df = df[ df[ column ] == value ]
		elif operator == '!=':
			df = df[ df[ column ] != value ]
		elif operator == '>':
			df = df[ df[ column ].astype( float ) > float( value ) ]
		elif operator == '<':
			df = df[ df[ column ].astype( float ) < float( value ) ]
		elif operator == '>=':
			df = df[ df[ column ].astype( float ) >= float( value ) ]
		elif operator == '<=':
			df = df[ df[ column ].astype( float ) <= float( value ) ]
		elif operator == 'contains':
			df = df[ df[ column ].astype( str ).str.contains( value ) ]
	
	return df

def create_aggregation( df: pd.DataFrame ):
	st.subheader( 'Aggregation Engine' )
	
	numeric_cols = df.select_dtypes( include=[ 'number' ] ).columns.tolist( )
	
	if not numeric_cols:
		st.info( 'No numeric columns available.' )
		return
	
	col = st.selectbox( 'Column', numeric_cols )
	agg = st.selectbox( 'Aggregation', [ 'COUNT', 'SUM', 'AVG', 'MIN', 'MAX', 'MEDIAN' ] )
	
	if agg == 'COUNT':
		result = df[ col ].count( )
	elif agg == 'SUM':
		result = df[ col ].sum( )
	elif agg == 'AVG':
		result = df[ col ].mean( )
	elif agg == 'MIN':
		result = df[ col ].min( )
	elif agg == 'MAX':
		result = df[ col ].max( )
	elif agg == 'MEDIAN':
		result = df[ col ].median( )
	
	st.metric( 'Result', result )

def create_visualization( df: pd.DataFrame ):
	st.subheader( 'Visualization Engine' )
	numeric_cols = df.select_dtypes( include=[ 'number' ] ).columns.tolist( )
	categorical_cols = df.select_dtypes( include=[ 'object' ] ).columns.tolist( )
	chart = st.selectbox( 'Chart Type',
		[ 'Histogram', 'Bar', 'Line', 'Scatter', 'Box', 'Pie', 'Correlation' ] )
	
	if chart == 'Histogram' and numeric_cols:
		col = st.selectbox( 'Column', numeric_cols )
		fig = px.histogram( df, x=col )
		st.plotly_chart( fig, use_container_width=True )
	
	elif chart == 'Bar':
		x = st.selectbox( 'X', df.columns )
		y = st.selectbox( 'Y', numeric_cols )
		fig = px.bar( df, x=x, y=y )
		st.plotly_chart( fig, use_container_width=True )
	
	elif chart == 'Line':
		x = st.selectbox( 'X', df.columns )
		y = st.selectbox( 'Y', numeric_cols )
		fig = px.line( df, x=x, y=y )
		st.plotly_chart( fig, use_container_width=True )
	
	elif chart == 'Scatter':
		x = st.selectbox( 'X', numeric_cols )
		y = st.selectbox( 'Y', numeric_cols )
		fig = px.scatter( df, x=x, y=y )
		st.plotly_chart( fig, use_container_width=True )
	
	elif chart == 'Box':
		col = st.selectbox( 'Column', numeric_cols )
		fig = px.box( df, y=col )
		st.plotly_chart( fig, use_container_width=True )
	
	elif chart == 'Pie':
		col = st.selectbox( 'Category Column', categorical_cols )
		fig = px.pie( df, names=col )
		st.plotly_chart( fig, use_container_width=True )
	
	elif chart == 'Correlation' and len( numeric_cols ) > 1:
		corr = df[ numeric_cols ].corr( )
		fig = px.imshow( corr, text_auto=True )
		st.plotly_chart( fig, use_container_width=True )

def convert_dataframe( table_name: str, df: pd.DataFrame ):
	columns = [ ]
	for col in df.columns:
		sql_type = get_sqlite_type( df[ col ].dtype )
		safe_col = col.replace( ' ', '_' )
		columns.append( f'{safe_col} {sql_type}' )
	
	create_stmt = f'CREATE TABLE IF NOT EXISTS {table_name} ({", ".join( columns )});'
	
	with create_connection( ) as conn:
		conn.execute( create_stmt )
		conn.commit( )

def insert_data( table_name: str, df: pd.DataFrame ):
	df = df.copy( )
	df.columns = [ c.replace( ' ', '_' ) for c in df.columns ]
	
	placeholders = ', '.join( [ '?' ] * len( df.columns ) )
	stmt = f'INSERT INTO {table_name} VALUES ({placeholders});'
	
	with create_connection( ) as conn:
		conn.executemany( stmt, df.values.tolist( ) )
		conn.commit( )

def get_sqlite_type( dtype ) -> str:
	"""
		Purpose:
		--------
		Map a pandas dtype to an appropriate SQLite column type.
	
		Parameters:
		-----------
		dtype : pandas dtype
			The dtype of a pandas Series.
	
		Returns:
		--------
		str
			SQLite column type.
	"""
	dtype_str = str( dtype ).lower( )
	
	# ----------  Integer Types
	if "int" in dtype_str:
		return "INTEGER"
	
	# ----------  Float Types
	if "float" in dtype_str:
		return "REAL"
	
	# ----------  Boolean
	if "bool" in dtype_str:
		return "INTEGER"
	
	# ----------  Datetime
	if "datetime" in dtype_str:
		return "TEXT"
	
	# ----------  Categorical
	if "category" in dtype_str:
		return "TEXT"
	
	# ----------  Default fallback
	return "TEXT"

def create_custom_table( table_name: str, columns: list ) -> None:
	"""
		Purpose:
		--------
		Create a custom SQLite table from column definitions.
	
		Parameters:
		-----------
		table_name : str
			Name of table.
	
		columns : list of dict
			[
				{
					"name": str,
					"type": str,
					"not_null": bool,
					"primary_key": bool,
					"auto_increment": bool
				}
			]
	"""
	if not table_name:
		raise ValueError( "Table name required." )
	
	# ----------  Validate identifier
	if not re.match( r"^[A-Za-z_][A-Za-z0-9_]*$", table_name ):
		raise ValueError( "Invalid table name." )
	
	col_defs = [ ]
	for col in columns:
		col_name = col[ "name" ]
		col_type = col[ "type" ].upper( )
		if not re.match( r"^[A-Za-z_][A-Za-z0-9_]*$", col_name ):
			raise ValueError( f"Invalid column name: {col_name}" )
		
		definition = f'"{col_name}" {col_type}'
		if col[ "primary_key" ]:
			definition += " PRIMARY KEY"
			if col[ "auto_increment" ] and col_type == "INTEGER":
				definition += " AUTOINCREMENT"
		
		if col[ "not_null" ]:
			definition += " NOT NULL"
	
		col_defs.append( definition )
	
	sql = f'CREATE TABLE IF NOT EXISTS "{table_name}" ({", ".join( col_defs )});'
	with create_connection( ) as conn:
		conn.execute( sql )
		conn.commit( )

def is_safe_query( query: str ) -> bool:
	"""
	
		Purpose:
		--------
		Determine whether a SQL query is read-only and safe to execute.
	
		Allows:
			SELECT
			WITH (CTE returning SELECT)
			EXPLAIN SELECT
			PRAGMA (read-only)
	
		Blocks:
			INSERT, UPDATE, DELETE, DROP, ALTER, CREATE, ATTACH,
			DETACH, VACUUM, REPLACE, TRIGGER, and multiple statements.
			
	"""
	if not query or not isinstance( query, str ):
		return False
	
	q = query.strip( ).lower( )
	
	# ----------  Block multiple statements
	if ';' in q[ :-1 ]:
		return False
	
	# ----------  Remove SQL comments
	q = re.sub( r"--.*?$", "", q, flags=re.MULTILINE )
	q = re.sub( r"/\*.*?\*/", "", q, flags=re.DOTALL )
	q = q.strip( )
	
	# ----------  Allowed starting keywords
	allowed_starts = ('select', 'with', 'explain', 'pragma')
	if not q.startswith( allowed_starts ):
		return False
	
	# ----------  Block dangerous keywords anywhere
	blocked_keywords = ('insert ', 'update ', 'delete ', 'drop ', 'alter ',
	                    'create ', 'attach ', 'detach ', 'vacuum ', 'replace ', 'trigger ')
	
	for keyword in blocked_keywords:
		if keyword in q:
			return False
	
	return True

def create_identifier( name: str ) -> str:
	"""
	
		Purpose:
		--------
		Sanitize a string into a safe SQLite identifier.
	
		- Replaces invalid characters with underscores
		- Ensures it starts with a letter or underscore
		- Prevents empty names
		
	"""
	if not name or not isinstance( name, str ):
		raise ValueError( 'Invalid Identifier.' )
	
	safe = re.sub( r'[^0-9a-zA-Z_]', '_', name.strip( ) )
	if not re.match( r'^[A-Za-z_]', safe ):
		safe = f'_{safe}'
	
	if not safe:
		raise ValueError( 'Invalid identifier after sanitization.' )
	
	return safe

def get_indexes( table: str ):
	with create_connection( ) as conn:
		rows = conn.execute( f'PRAGMA index_list("{table}");' ).fetchall( )
		return rows

def add_column( table: str, column: str, col_type: str ):
	column = create_identifier( column )
	col_type = col_type.upper( )
	
	with create_connection( ) as conn:
		conn.execute(
			f'ALTER TABLE "{table}" ADD COLUMN "{column}" {col_type};' )
		conn.commit( )

def create_profile_table( table: str ):
	df = read_table( table )
	profile_rows = [ ]
	total_rows = len( df )
	for col in df.columns:
		series = df[ col ]
		null_count = series.isna( ).sum( )
		distinct_count = series.nunique( dropna=True )
		row = \
			{
					'column': col, 'dtype': str( series.dtype ),
					'null_%': round( (null_count / total_rows) * 100, 2 ) if total_rows else 0,
					'distinct_%': round( (
								                     distinct_count / total_rows) * 100, 2 ) if total_rows else 0,
			}
		
		if pd.api.types.is_numeric_dtype( series ):
			row[ "min" ] = series.min( )
			row[ "max" ] = series.max( )
			row[ "mean" ] = series.mean( )
		else:
			row[ "min" ] = None
			row[ "max" ] = None
			row[ "mean" ] = None
		
		profile_rows.append( row )
	
	return pd.DataFrame( profile_rows )

def drop_column( table: str, column: str ):
	if not table or not column:
		raise ValueError( "Table and column required." )
	
	with create_connection( ) as conn:
		schema = conn.execute( f'PRAGMA table_info("{table}");' ).fetchall( )
		if not schema:
			raise ValueError( "Table definition not found." )
		
		col_names = [ r[ 1 ] for r in schema ]
		if column not in col_names:
			raise ValueError( "Column not found." )
		
		remaining = [ r for r in schema if r[ 1 ] != column ]
		if not remaining:
			raise ValueError( "Cannot drop the only remaining column." )
		
		temp_table = f"{table}_rebuild_temp"
		
		pk_cols = [ r for r in remaining if int( r[ 5 ] or 0 ) > 0 ]
		single_pk = len( pk_cols ) == 1
		
		new_defs: List[ str ] = [ ]
		for row in remaining:
			col_name = row[ 1 ]
			col_type = row[ 2 ] or ''
			not_null = int( row[ 3 ] or 0 )
			default_value = row[ 4 ]
			pk = int( row[ 5 ] or 0 )
			
			col_def = f'"{col_name}" {col_type}'.strip( )
			
			if not_null:
				col_def += ' NOT NULL'
			
			if default_value is not None:
				col_def += f' DEFAULT {default_value}'
			
			if single_pk and pk == 1:
				col_def += ' PRIMARY KEY'
			
			new_defs.append( col_def )
		
		new_create_sql = f'CREATE TABLE "{temp_table}" ({", ".join( new_defs )});'
		
		conn.execute( "BEGIN" )
		conn.execute( new_create_sql )
		
		remaining_cols = [ r[ 1 ] for r in remaining ]
		col_list = ", ".join( [ f'"{c}"' for c in remaining_cols ] )
		
		conn.execute(
			f'INSERT INTO "{temp_table}" ({col_list}) '
			f'SELECT {col_list} FROM "{table}";'
		)
		
		indexes = conn.execute(
			"""
            SELECT sql
            FROM sqlite_master
            WHERE type ='index' AND tbl_name=? AND sql IS NOT NULL
			""",
			(table,)
		).fetchall( )
		
		conn.execute( f'DROP TABLE "{table}";' )
		conn.execute( f'ALTER TABLE "{temp_table}" RENAME TO "{table}";' )
		
		for idx in indexes:
			idx_sql = idx[ 0 ]
			if idx_sql and column not in idx_sql:
				conn.execute( idx_sql )
		
		conn.commit( )
	
def load_prompt( pid: int ) -> None:
	with create_connection( ) as conn:
		_select = f"SELECT Caption, Name, Text, Version, ID FROM {TABLE} WHERE PromptsId=?"
		cur = conn.execute( _select, (pid,), )
		row = cur.fetchone( )
		if not row:
			return
		st.session_state.pe_caption = row[ 0 ]
		st.session_state.pe_name = row[ 1 ]
		st.session_state.pe_text = row[ 2 ]
		st.session_state.pe_version = row[ 3 ]
		st.session_state.pe_id = row[ 4 ]

def get_ai_asset_tables( ) -> List[ str ]:
	"""
		Purpose:
		--------
		Return the AI-asset governance table names.

		Parameters:
		-----------
		None

		Returns:
		--------
		List[str]
			AI-asset table names.
	"""
	return [
			'documents',
			'document_chunks',
			'document_embeddings',
			'images'
	]

def get_table_row_count( conn: sqlite3.Connection, table_name: str ) -> int:
	"""
		Purpose:
		--------
		Return the row count for a SQLite table.

		Parameters:
		-----------
		conn : sqlite3.Connection
			Open SQLite connection.

		table_name : str
			Table name.

		Returns:
		--------
		int
			Number of rows in the table.
	"""
	if not table_name:
		return 0
	
	try:
		row = conn.execute(
			f'SELECT COUNT(*) FROM "{table_name}";'
		).fetchone( )
		
		return int( row[ 0 ] ) if row else 0
	except Exception:
		return 0

def get_ai_asset_counts( ) -> Dict[ str, int ]:
	"""
		Purpose:
		--------
		Count rows in the AI asset governance tables used by Document Q&A,
		Semantic Search, and Data Management.

		Parameters:
		-----------
		None

		Returns:
		--------
		Dict[str, int]
			Dictionary keyed by AI asset table name with row counts as values.
	"""
	counts: Dict[ str, int ] = { }
	tables = list_tables( )
	
	with create_connection( ) as conn:
		for table_name in get_ai_asset_tables( ):
			if table_name not in tables:
				counts[ table_name ] = 0
				continue
			
			counts[ table_name ] = get_table_row_count( conn, table_name )
	
	return counts

def purge_orphaned_document_chunks( conn: sqlite3.Connection ) -> int:
	"""
		Purpose:
		--------
		Delete document chunk rows whose DocumentName no longer exists in the documents
		table.

		Parameters:
		-----------
		conn : sqlite3.Connection
			Open SQLite connection.

		Returns:
		--------
		int
			Deleted row count.
	"""
	try:
		cur = conn.execute(
			"""
            DELETE
            FROM document_chunks
            WHERE DocumentName NOT IN
                  (SELECT Name
                   FROM documents);
			"""
		)
		
		return int( cur.rowcount if cur.rowcount is not None else 0 )
	except Exception:
		return 0

def purge_orphaned_document_embeddings( conn: sqlite3.Connection ) -> int:
	"""
		Purpose:
		--------
		Delete document embedding metadata rows whose DocumentName no longer exists in
		the documents table.

		Parameters:
		-----------
		conn : sqlite3.Connection
			Open SQLite connection.

		Returns:
		--------
		int
			Deleted row count.
	"""
	try:
		cur = conn.execute(
			"""
            DELETE
            FROM document_embeddings
            WHERE DocumentName NOT IN
                  (SELECT Name
                   FROM documents);
			"""
		)
		
		return int( cur.rowcount if cur.rowcount is not None else 0 )
	except Exception:
		return 0

def purge_orphaned_ai_assets( ) -> Dict[ str, int ]:
	"""
		Purpose:
		--------
		Delete orphaned AI asset rows that depend on the governed documents table.

		Parameters:
		-----------
		None

		Returns:
		--------
		Dict[str, int]
			Dictionary containing deleted chunk and embedding row counts.
	"""
	result: Dict[ str, int ] = {
			'deleted_chunks': 0,
			'deleted_embeddings': 0
	}
	
	tables = list_tables( )
	
	if 'documents' not in tables:
		return result
	
	with create_connection( ) as conn:
		if 'document_chunks' in tables:
			result[ 'deleted_chunks' ] = purge_orphaned_document_chunks( conn )
		
		if 'document_embeddings' in tables:
			result[ 'deleted_embeddings' ] = purge_orphaned_document_embeddings( conn )
		
		conn.commit( )
	
	return result

def get_timestamp_text( ) -> str:
	"""
		Purpose:
		--------
		Return a UTC-like timestamp string for metadata rows.

		Parameters:
		-----------
		None

		Returns:
		--------
		str
	"""
	return time.strftime( '%Y-%m-%d %H:%M:%S' )

def register_session_documents( ) -> Dict[ str, int ]:
	"""
		Purpose:
		--------
		Register active uploaded documents into the governed documents table.

		Parameters:
		-----------
		None

		Returns:
		--------
		Dict[str, int]
	"""
	active_docs = st.session_state.get( 'active_docs', [ ] )
	doc_bytes = st.session_state.get( 'doc_bytes', { } )
	
	inserted = 0
	updated = 0
	
	with create_connection( ) as conn:
		for name in active_docs:
			file_bytes = doc_bytes.get( name, b'' )
			if not file_bytes:
				continue
			
			text = extract_text( file_bytes, name )
			chunks = chunk_text( text ) if text else [ ]
			fingerprint = hashlib.sha256( file_bytes ).hexdigest( )
			file_type = Path( name ).suffix.lower( ).replace( '.', '' )
			created_on = get_timestamp_text( )
			
			existing = conn.execute(
				'''
                SELECT DocumentId
                FROM documents
                WHERE Name = ?
                  AND Fingerprint = ?
				''',
				(name, fingerprint)
			).fetchone( )
			
			if existing:
				conn.execute(
					'''
                    UPDATE documents
                    SET Type       = ?,
                        SizeBytes  = ?,
                        Source     = ?,
                        TextLength = ?,
                        ChunkCount = ?,
                        CreatedOn  = ?
                    WHERE DocumentId = ?
					''',
					(
							file_type,
							len( file_bytes ),
							'uploadlocal',
							len( text ),
							len( chunks ),
							created_on,
							existing[ 0 ]
					)
				)
				updated += 1
			else:
				conn.execute(
					'''
                    INSERT INTO documents
                    (Name,
                     Type,
                     SizeBytes,
                     Source,
                     Fingerprint,
                     TextLength,
                     ChunkCount,
                     CreatedOn)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?)
					''',
					(
							name,
							file_type,
							len( file_bytes ),
							'uploadlocal',
							fingerprint,
							len( text ),
							len( chunks ),
							created_on
					)
				)
				inserted += 1
		
		conn.commit( )
	
	return { 'inserted': inserted, 'updated': updated }

def register_session_chunks( ) -> Dict[ str, int ]:
	"""
		Purpose:
		--------
		Register active document chunks into the governed document_chunks table.

		Parameters:
		-----------
		None

		Returns:
		--------
		Dict[str, int]
	"""
	active_docs = st.session_state.get( 'active_docs', [ ] )
	doc_bytes = st.session_state.get( 'doc_bytes', { } )
	inserted = 0
	
	with create_connection( ) as conn:
		for name in active_docs:
			file_bytes = doc_bytes.get( name, b'' )
			if not file_bytes:
				continue
			
			text = extract_text( file_bytes, name )
			chunks = chunk_text( text ) if text else [ ]
			file_fingerprint = hashlib.sha256( file_bytes ).hexdigest( )
			created_on = get_timestamp_text( )
			
			conn.execute(
				'DELETE FROM document_chunks WHERE DocumentName = ? AND Fingerprint = ?',
				(name, file_fingerprint)
			)
			
			for idx, chunk_value in enumerate( chunks ):
				conn.execute(
					'''
                    INSERT INTO document_chunks
                    (DocumentName,
                     ChunkIndex,
                     ChunkText,
                     ChunkLength,
                     Fingerprint,
                     CreatedOn)
                    VALUES (?, ?, ?, ?, ?, ?)
					''',
					(
							name,
							idx,
							chunk_value,
							len( chunk_value ),
							file_fingerprint,
							created_on
					)
				)
				inserted += 1
		
		conn.commit( )
	
	return { 'inserted': inserted }

def register_session_embeddings( ) -> Dict[ str, int ]:
	"""
		Purpose:
		--------
		Register active document embedding metadata into the governed
		document_embeddings table.

		Parameters:
		-----------
		None

		Returns:
		--------
		Dict[str, int]
	"""
	active_docs = st.session_state.get( 'active_docs', [ ] )
	doc_bytes = st.session_state.get( 'doc_bytes', { } )
	inserted = 0
	
	if embedder is None:
		return { 'inserted': 0 }
	
	vector_dim = getattr( embedder, 'get_sentence_embedding_dimension', lambda: 384 )( )
	vector_dim = int( vector_dim ) if vector_dim else 384
	with create_connection( ) as conn:
		for name in active_docs:
			file_bytes = doc_bytes.get( name, b'' )
			if not file_bytes:
				continue
			
			text = extract_text( file_bytes, name )
			chunks = chunk_text( text ) if text else [ ]
			file_fingerprint = hashlib.sha256( file_bytes ).hexdigest( )
			created_on = get_timestamp_text( )
			
			conn.execute(
				'DELETE FROM document_embeddings WHERE DocumentName = ? AND Fingerprint = ?',
				(name, file_fingerprint) )
			
			for idx, _chunk_value in enumerate( chunks ):
				conn.execute( '''
                    INSERT INTO document_embeddings
                    (DocumentName,
                     ChunkIndex,
                     VectorDim,
                     Fingerprint,
                     CreatedOn)
                    VALUES (?, ?, ?, ?, ?)
					''', (name, idx, vector_dim, file_fingerprint, created_on) )
				inserted += 1
		
		conn.commit( )
	
	return { 'inserted': inserted }

def register_upload_images( uploaded_files: List[ Any ] ) -> Dict[ str, int ]:
	"""
		Purpose:
		--------
		Register uploaded image metadata into the governed images table.

		Parameters:
		-----------
		uploaded_files : List[Any]

		Returns:
		--------
		Dict[str, int]
	"""
	inserted = 0
	updated = 0
	
	with create_connection( ) as conn:
		for f in uploaded_files:
			try:
				name = str( getattr( f, 'name', '' ) or '' ).strip( )
				file_bytes = f.getvalue( )
				mime_type = str( getattr( f, 'type', '' ) or '' ).strip( )
			except Exception:
				continue
			
			if not name or not file_bytes:
				continue
			
			fingerprint = hashlib.sha256( file_bytes ).hexdigest( )
			created_on = get_timestamp_text( )
			
			existing = conn.execute(
				'''
                SELECT ImageId
                FROM images
                WHERE Name = ?
                  AND Fingerprint = ?
				''',
				(name, fingerprint)
			).fetchone( )
			
			if existing:
				conn.execute(
					'''
                    UPDATE images
                    SET MimeType  = ?,
                        SizeBytes = ?,
                        Source    = ?,
                        CreatedOn = ?
                    WHERE ImageId = ?
					''',
					(
							mime_type,
							len( file_bytes ),
							'uploadlocal',
							created_on,
							existing[ 0 ]
					)
				)
				updated += 1
			else:
				conn.execute(
					'''
                    INSERT INTO images
                    (Name,
                     MimeType,
                     SizeBytes,
                     Fingerprint,
                     Source,
                     CreatedOn)
                    VALUES (?, ?, ?, ?, ?, ?)
					''',
					(
							name,
							mime_type,
							len( file_bytes ),
							fingerprint,
							'uploadlocal',
							created_on
					)
				)
				inserted += 1
		
		conn.commit( )
	
	return { 'inserted': inserted, 'updated': updated }

# -------------- LLM  UTILITIES -------------------

@st.cache_resource
def load_llm( model_path: str, ctx: int, threads: int, seed: int ) -> Any | None:
	"""
		Purpose:
		--------
		Lazily load the selected local llama.cpp GGUF model using the supplied runtime
		settings. The model path is part of the cache key so switching models creates a
		distinct cached model resource.

		Parameters:
		-----------
		model_path : str
			Local GGUF model path.

		ctx : int
			Context window size.

		threads : int
			CPU thread count.

		seed : int
			Random seed used by llama.cpp.

		Returns:
		--------
		Any | None
			Loaded llama.cpp model instance when available; otherwise None.
	"""
	try:
		if Llama is None:
			return None
		
		model_path_value = str( model_path or '' ).strip( )
		
		if not model_path_value:
			return None
		
		if not Path( model_path_value ).exists( ):
			return None
		
		ctx_value = int( ctx ) if int( ctx ) > 0 else int( cfg.DEFAULT_CTX )
		thread_value = int( threads ) if int( threads ) > 0 else int( cfg.CORES )
		seed_value = int( seed ) if seed is not None else -1
		
		return Llama(
			model_path=model_path_value,
			n_ctx=ctx_value,
			n_threads=thread_value,
			n_batch=512,
			seed=seed_value,
			verbose=False
		)
	
	except Exception:
		return None

@st.cache_resource
def load_embedder( ) -> Any | None:
	"""
		Purpose:
		--------
		Lazily load the sentence embedding model when the dependency is available.

		Parameters:
		-----------
		None

		Returns:
		--------
		Any | None
			A sentence-transformer model instance when available; otherwise None.
	"""
	try:
		from sentence_transformers import SentenceTransformer
		
		return SentenceTransformer( 'all-MiniLM-L6-v2' )
	except Exception:
		return None

# ------------- DOCQNA UTILITIES ----------------------

def create_docqna_instruction( action_name: str ) -> str:
	"""
		Purpose:
		--------
		Return an instruction block for a selected document action.

		Parameters:
		-----------
		action_name : str

		Returns:
		--------
		str
	"""
	action = str( action_name or 'Answer Question' ).strip( )
	action_map = {
			'Answer Question':
				'Answer the user question directly using the retrieved excerpts.',
			'Summarize Active Document':
				'Provide a clear, structured summary of the active document.',
			'Extract Key Points':
				'Extract the most important points as a concise bullet list.',
			'Generate Outline':
				'Generate a structured outline of the document.',
			'Extract Entities':
				'Extract named entities, important organizations, dates, and references.',
			'Extract Tables':
				'Describe tabular information or structured fields present in the excerpts.',
			'Compare Active Documents':
				'Compare the active documents, noting agreements, differences, and gaps.'
	}
	
	return action_map.get( action, action_map[ 'Answer Question' ] )

def build_instruction_block( ) -> str:
	"""
		Purpose:
		--------
		Build a unified instruction block for document-grounded answering.

		Parameters:
		-----------
		None

		Returns:
		--------
		str
	"""
	system_instructions = get_effective_system_instructions( )
	require_grounding = bool( st.session_state.get( 'require_grounding', True ) )
	answer_from_excerpts_only = bool( st.session_state.get( 'answer_from_excerpts_only', True ) )
	response_format = str( st.session_state.get( 'response_format', 'Markdown' ) or 'Markdown' ).strip( )
	doc_action = str( st.session_state.get( 'docqna_action', 'Answer Question' ) or 'Answer Question' )
	lines: List[ str ] = [ ]
	if system_instructions:
		lines.append( system_instructions )
	
	lines.append( 'Document Q&A Instructions:' )
	lines.append( f'- Action: {doc_action}' )
	lines.append( f'- Response Format: {response_format}' )
	lines.append( f'- Action Guidance: {create_docqna_instruction( doc_action )}' )
	if require_grounding:
		lines.append( '- Ground every answer in the retrieved document excerpts.' )
	
	if answer_from_excerpts_only:
		lines.append(
			'- If the retrieved excerpts do not contain the answer, '
			'state clearly that there is not enough information.' )
	
	if response_format == 'JSON':
		lines.append( '- Return valid JSON only.' )
	
	return '\n'.join( lines ).strip( )

def extract_text_from_pdf_bytes( file_bytes: bytes, include_page_markers: bool = False ) -> str:
	"""
		Purpose:
		--------
		Extract native text from PDF bytes using PyMuPDF when it is available.

		Parameters:
		-----------
		file_bytes : bytes
			PDF file bytes.

		include_page_markers : bool
			When True, include page markers before each extracted page.

		Returns:
		--------
		str
			Extracted PDF text.
	"""
	if not file_bytes:
		return ''
	
	if fitz is None:
		return ''
	
	try:
		doc = fitz.open( stream=file_bytes, filetype='pdf' )
		parts: List[ str ] = [ ]
		
		for page_index, page in enumerate( doc, start=1 ):
			page_text = page.get_text( 'text' ) or ''
			
			if include_page_markers:
				parts.append( f'[Page {page_index}]' )
			
			if page_text:
				parts.append( page_text )
		
		doc.close( )
		return '\n'.join( parts ).strip( )
	except Exception:
		return ''

def extract_text_from_docx_bytes( file_bytes: bytes ) -> str:
	"""
		Purpose:
		--------
		Extract text from DOCX bytes using python-docx when it is available.

		Parameters:
		-----------
		file_bytes : bytes
			DOCX file bytes.

		Returns:
		--------
		str
			Extracted DOCX text.
	"""
	if not file_bytes:
		return ''

	if Document is None:
		return ''
	
	try:
		buffer = BytesIO( file_bytes )
		document = Document( buffer )
		parts: List[ str ] = [ ]
		
		for paragraph in document.paragraphs:
			text = str( paragraph.text or '' ).strip( )
			if text:
				parts.append( text )
		
		for table in document.tables:
			for row in table.rows:
				values: List[ str ] = [ ]
				for cell in row.cells:
					cell_text = str( cell.text or '' ).strip( )
					values.append( cell_text )
				
				row_text = ' | '.join( values ).strip( )
				if row_text:
					parts.append( row_text )
		
		return '\n'.join( parts ).strip( )
	except Exception:
		return ''

def compute_fingerprint( active_docs: List[ str ], doc_bytes: Dict[ str, bytes ] ) -> str:
	'''
		
		Purpose:
		--------
		Computes a stable fingerprint for the currently selected active documents and their byte contents.
	
		Parameters:
		-----------
		active_docs:
			A List[ str ] of active document names.
		doc_bytes:
			A Dict[ str, bytes ] mapping document name to file bytes.
	
		Returns:
		--------
		A str fingerprint suitable for cache invalidation.
	
	'''
	h = hashlib.sha256( )
	for name in sorted( active_docs ):
		b = doc_bytes.get( name, b'' )
		h.update( name.encode( 'utf-8', errors='ignore' ) )
		h.update( len( b ).to_bytes( 8, 'little', signed=False ) )
		h.update( hashlib.sha256( b ).digest( ) )
	return h.hexdigest( )

def decode_text_bytes( file_bytes: bytes ) -> str:
	"""
		Purpose:
		--------
		Decode text-like document bytes using common encodings and a permissive fallback.

		Parameters:
		-----------
		file_bytes : bytes
			File bytes to decode.

		Returns:
		--------
		str
			Decoded text.
	"""
	if not file_bytes:
		return ''
	
	encodings = [ 'utf-8', 'utf-8-sig', 'cp1252', 'latin-1' ]
	
	for encoding in encodings:
		try:
			return file_bytes.decode( encoding ).strip( )
		except Exception:
			continue
	
	try:
		return file_bytes.decode( errors='ignore' ).strip( )
	except Exception:
		return ''

def extract_text_from_bytes( file_bytes: bytes, file_name: str = '' ) -> str:
	"""
		Purpose:
		--------
		Extract text from supported document bytes using the file name extension and
		current parsing preferences.

		Parameters:
		-----------
		file_bytes : bytes
			Source document bytes.

		file_name : str
			Source document name.

		Returns:
		--------
		str
			Extracted text.
	"""
	if not file_bytes:
		return ''
	
	file_name_value = str( file_name or '' ).lower( ).strip( )
	include_page_markers = bool( st.session_state.get( 'include_page_markers', False ) )
	prefer_native_pdf_text = bool( st.session_state.get( 'prefer_native_pdf_text', True ) )
	
	if file_name_value.endswith( '.pdf' ):
		if prefer_native_pdf_text:
			text = extract_text_from_pdf_bytes(
				file_bytes=file_bytes,
				include_page_markers=include_page_markers
			)
			
			if text:
				return text
		
		return decode_text_bytes( file_bytes )
	
	if file_name_value.endswith( '.docx' ):
		text = extract_text_from_docx_bytes( file_bytes )
		
		if text:
			return text
		
		return decode_text_bytes( file_bytes )
	
	if (
			file_name_value.endswith( '.txt' )
			or file_name_value.endswith( '.md' )
			or file_name_value.endswith( '.csv' )
			or file_name_value.endswith( '.json' )
			or file_name_value.endswith( '.xml' )
			or file_name_value.endswith( '.html' )
			or file_name_value.endswith( '.htm' )
	):
		return decode_text_bytes( file_bytes )
	
	if not file_name_value:
		pdf_text = extract_text_from_pdf_bytes(
			file_bytes=file_bytes,
			include_page_markers=include_page_markers
		)
		
		if pdf_text:
			return pdf_text
	
	return decode_text_bytes( file_bytes )

def extract_text_bytes( file_bytes: bytes, file_name: str = '' ) -> str:
	"""
		Purpose:
		--------
		Backward-compatible wrapper for extracting text from document bytes.

		Parameters:
		-----------
		file_bytes : bytes
			Source document bytes.

		file_name : str
			Source document name.

		Returns:
		--------
		str
			Extracted text.
	"""
	return extract_text_from_bytes( file_bytes=file_bytes, file_name=file_name )

def extract_text( file_bytes: bytes, file_name: str = '' ) -> str:
	"""
		Purpose:
		--------
		Extract document text using the configured parsing behavior.

		Parameters:
		-----------
		file_bytes : bytes
			Source document bytes.

		file_name : str
			Source document name.

		Returns:
		--------
		str
			Extracted text.
	"""
	return extract_text_from_bytes( file_bytes=file_bytes, file_name=file_name )

def load_sqlite_vec( conn: sqlite3.Connection ) -> bool:
	'''
		
		Purpose:
		--------
		Attempts to load sqlite-vec into the provided SQLite connection.
	
		Parameters:
		-----------
		conn:
			The sqlite3.Connection.
	
		Returns:
		--------
		True if sqlite-vec loaded successfully; otherwise False.
		
	'''
	try:
		import sqlite_vec
		
		sqlite_vec.load( conn )
		return True
	except Exception:
		return False

def ensure_schema( dim: int ) -> bool:
	'''
	
		Purpose:
		--------
		Creates the sqlite-vec virtual table used for Document Q&A embeddings if possible.
	
		Parameters:
		-----------
		dim:
			The embedding dimension (e.g., 384 for all-MiniLM-L6-v2).
	
		Returns:
		--------
		True if the schema exists and is usable; otherwise False.
	
	'''
	conn = create_connection( )
	try:
		ok = load_sqlite_vec( conn )
		if not ok:
			return False
		
		cur = conn.cursor( )
		cur.execute(
			f'''
			CREATE VIRTUAL TABLE IF NOT EXISTS docqna_vec
			USING vec0(
				embedding float[{int( dim )}],
				doc_name TEXT,
				chunk TEXT
			);
			'''
		)
		conn.commit( )
		return True
	except Exception:
		return False
	finally:
		conn.close( )

def build_docqna_inventory( ) -> List[ Dict[ str, Any ] ]:
	"""
		Purpose:
		--------
		Build inventory rows for the currently active uploaded documents.

		Parameters:
		-----------
		None

		Returns:
		--------
		List[Dict[str, Any]]
	"""
	rows: List[ Dict[ str, Any ] ] = [ ]
	active_docs = st.session_state.get( 'active_docs', [ ] )
	doc_bytes = st.session_state.get( 'doc_bytes', { } )
	for name in active_docs:
		b = doc_bytes.get( name, b'' )
		text = extract_text( b, name ) if b else ''
		chunks = chunk_text( text ) if text else [ ]
		rows.append( {
					'Name': name,
					'SizeBytes': len( b ) if b else 0,
					'TextLength': len( text ) if text else 0,
					'ChunkCount': len( chunks ),
					'Loaded': bool( b )
			} )
	
	return rows

def get_docqna_names( ) -> str:
	"""
		Purpose:
		--------
		Build a human-readable string of active document names.

		Parameters:
		-----------
		None

		Returns:
		--------
		str
	"""
	active_docs = st.session_state.get( 'active_docs', [ ] )
	if not isinstance( active_docs, list ) or len( active_docs ) == 0:
		return 'No active documents'
	return ', '.join( [ str( name ) for name in active_docs ] )

def is_embedder_available( candidate: Any | None = None ) -> bool:
	"""
		Purpose:
		--------
		Determine whether a sentence embedding model is available and usable.

		Parameters:
		-----------
		candidate : Any | None
			Optional embedding model instance. When omitted, the global embedder is used.

		Returns:
		--------
		bool
			True when an embedder with an encode method is available; otherwise False.
	"""
	model = candidate if candidate is not None else globals( ).get( 'embedder', None )
	return model is not None and hasattr( model, 'encode' )

def get_embedder_unavailable_message( ) -> str:
	"""
		Purpose:
		--------
		Return a standard diagnostic message when sentence-transformer embeddings are
		unavailable.

		Parameters:
		-----------
		None

		Returns:
		--------
		str
			Diagnostic message.
	"""
	return (
			'Embedding retrieval is unavailable because the sentence-transformer model '
			'could not be loaded. Install or repair `sentence-transformers`, then restart '
			'the Streamlit app.'
	)

def decode_embedding_vector( vector_blob: bytes | memoryview | None ) -> np.ndarray:
	"""
		Purpose:
		--------
		Decode a stored embedding vector BLOB into a NumPy float32 array.

		Parameters:
		-----------
		vector_blob : bytes | memoryview | None
			Stored vector BLOB.

		Returns:
		--------
		np.ndarray
			Decoded vector. Empty array when decoding fails.
	"""
	if not vector_blob:
		return np.asarray( [ ], dtype=np.float32 )
	
	try:
		return np.frombuffer( vector_blob, dtype=np.float32 )
	except Exception:
		return np.asarray( [ ], dtype=np.float32 )

def rebuild_index( embedder: Any | None ) -> None:
	"""
		Purpose:
		--------
		Build or refresh the Document Q&A vector index when active documents or chunk
		settings change. The function fails closed when embeddings are unavailable instead
		of raising an AttributeError from embedder.encode(...).

		Parameters:
		-----------
		embedder : Any | None
			Sentence embedding model instance.

		Returns:
		--------
		None
	"""
	if not is_embedder_available( embedder ):
		st.session_state[ 'docqna_vec_ready' ] = False
		st.session_state[ 'docqna_fallback_rows' ] = [ ]
		st.session_state[ 'docqna_chunk_count' ] = 0
		st.session_state[ 'docqna_last_retrieval' ] = [ ]
		st.session_state[ 'docqna_inventory_rows' ] = build_docqna_inventory( )
		st.session_state[ 'docqna_retrieval_status' ] = get_embedder_unavailable_message( )
		return
	
	active_docs: List[ str ] = st.session_state.get( 'active_docs', [ ] )
	doc_bytes: Dict[ str, bytes ] = st.session_state.get( 'doc_bytes', { } )
	retrieval_chunk_size = int( st.session_state.get( 'retrieval_chunk_size', 1200 ) )
	retrieval_chunk_overlap = int( st.session_state.get( 'retrieval_chunk_overlap', 200 ) )
	
	fp_seed = f'{retrieval_chunk_size}|{retrieval_chunk_overlap}|'
	fp_seed += compute_fingerprint( active_docs, doc_bytes )
	fp = hashlib.sha256( fp_seed.encode( 'utf-8', errors='ignore' ) ).hexdigest( )
	
	if fp and fp == st.session_state.get( 'docqna_fingerprint', '' ):
		st.session_state[ 'docqna_inventory_rows' ] = build_docqna_inventory( )
		return
	
	st.session_state[ 'docqna_fingerprint' ] = fp
	st.session_state[ 'docqna_chunk_count' ] = 0
	st.session_state[ 'docqna_fallback_rows' ] = [ ]
	st.session_state[ 'docqna_inventory_rows' ] = build_docqna_inventory( )
	st.session_state[ 'docqna_retrieval_status' ] = ''
	
	try:
		dim_value = getattr( embedder, 'get_sentence_embedding_dimension', lambda: 384 )( )
		dim = int( dim_value ) if dim_value else 384
	except Exception:
		dim = 384
	
	prefer_sqlite_vec = bool( st.session_state.get( 'prefer_sqlite_vec', True ) )
	vec_ready = False
	
	if prefer_sqlite_vec:
		vec_ready = ensure_schema( dim )
	
	st.session_state[ 'docqna_vec_ready' ] = bool( vec_ready )
	
	conn = create_connection( )
	try:
		cur = conn.cursor( )
		
		if vec_ready:
			try:
				cur.execute( 'DELETE FROM docqna_vec;' )
				conn.commit( )
			except Exception:
				st.session_state[ 'docqna_vec_ready' ] = False
				vec_ready = False
		
		total_chunks = 0
		fallback_rows: List[ Tuple[ str, str, bytes ] ] = [ ]
		
		for name in active_docs:
			b = doc_bytes.get( name )
			if not b:
				continue
			
			text = extract_text( b, name )
			if not text:
				continue
			
			chunks = chunk_text(
				text,
				size=retrieval_chunk_size,
				overlap=retrieval_chunk_overlap
			)
			
			if not chunks:
				continue
			
			try:
				vecs = embedder.encode( chunks, show_progress_bar=False )
				vecs = np.asarray( vecs, dtype=np.float32 )
			except Exception as e:
				st.session_state[ 'docqna_retrieval_status' ] = (
						f'Embedding generation failed for {name}: {e}'
				)
				continue
			
			if vec_ready:
				for chunk_text_value, v in zip( chunks, vecs ):
					cur.execute(
						'INSERT INTO docqna_vec ( embedding, doc_name, chunk ) VALUES ( ?, ?, ? );',
						(v.tobytes( ), name, chunk_text_value)
					)
			else:
				for chunk_text_value, v in zip( chunks, vecs ):
					fallback_rows.append( (name, chunk_text_value, v.tobytes( )) )
			
			total_chunks += int( len( chunks ) )
		
		conn.commit( )
		st.session_state[ 'docqna_chunk_count' ] = total_chunks
		
		if not vec_ready:
			st.session_state[ 'docqna_fallback_rows' ] = fallback_rows
		else:
			st.session_state[ 'docqna_fallback_rows' ] = [ ]
	
	except Exception as e:
		st.session_state[ 'docqna_vec_ready' ] = False
		st.session_state[ 'docqna_fallback_rows' ] = [ ]
		st.session_state[ 'docqna_chunk_count' ] = 0
		st.session_state[ 'docqna_retrieval_status' ] = f'Document index rebuild failed: {e}'
	
	finally:
		conn.close( )

def retrieve_chunks( query: str, k: int = None ) -> List[ Tuple[ str, str, float ] ]:
	"""
		Purpose:
		--------
		Retrieve top-k document chunks relevant to the query using sqlite-vec when available,
		with optional cosine-similarity fallback. Missing embeddings fail safely.

		Parameters:
		-----------
		query : str
			User query.

		k : int | None
			Number of chunks to retrieve.

		Returns:
		--------
		List[Tuple[str, str, float]]
			Ranked retrieval results as document name, chunk text, and score or distance.
	"""
	if not query or not query.strip( ):
		return [ ]
	
	if not is_embedder_available( globals( ).get( 'embedder', None ) ):
		st.session_state[ 'docqna_last_retrieval' ] = [ ]
		st.session_state[ 'docqna_retrieval_status' ] = get_embedder_unavailable_message( )
		return [ ]
	
	rebuild_index( embedder )
	
	k_value = int( k ) if k is not None else int( st.session_state.get( 'retrieval_k', 6 ) )
	if k_value <= 0:
		k_value = 6
	
	try:
		qv = embedder.encode( [ query ], show_progress_bar=False )
		qv = np.asarray( qv, dtype=np.float32 )[ 0 ]
	except Exception as e:
		st.session_state[ 'docqna_last_retrieval' ] = [ ]
		st.session_state[ 'docqna_retrieval_status' ] = f'Query embedding failed: {e}'
		return [ ]
	
	if bool( st.session_state.get( 'docqna_vec_ready', False ) ):
		conn = create_connection( )
		try:
			load_sqlite_vec( conn )
			cur = conn.cursor( )
			cur.execute(
				'''
                SELECT doc_name, chunk, distance
                FROM docqna_vec
                WHERE embedding MATCH ?
                ORDER BY distance ASC LIMIT ?;
				''',
				(qv.tobytes( ), int( k_value ))
			)
			rows = cur.fetchall( )
			results = [ (r[ 0 ], r[ 1 ], float( r[ 2 ] )) for r in rows ]
			st.session_state[ 'docqna_last_retrieval' ] = results
			return results
		
		except Exception as e:
			st.session_state[ 'docqna_vec_ready' ] = False
			st.session_state[ 'docqna_retrieval_status' ] = (
					f'sqlite-vec retrieval failed; using fallback when enabled. Error: {e}'
			)
		
		finally:
			conn.close( )
	
	if not bool( st.session_state.get( 'allow_similarity_fallback', True ) ):
		st.session_state[ 'docqna_last_retrieval' ] = [ ]
		return [ ]
	
	fallback_rows: List[ Tuple[ str, str, bytes ] ] = st.session_state.get(
		'docqna_fallback_rows',
		[ ]
	)
	
	results: List[ Tuple[ str, str, float ] ] = [ ]
	
	for doc_name, chunk_text_value, vec_blob in fallback_rows:
		v = decode_embedding_vector( vec_blob )
		
		if v.size == 0:
			continue
		
		score = cosine_similarity( qv, v )
		results.append( (doc_name, chunk_text_value, float( score )) )
	
	results.sort( key=lambda r: r[ 2 ], reverse=True )
	results = results[ : int( k_value ) ]
	st.session_state[ 'docqna_last_retrieval' ] = results
	return results

def build_docqna_input( user_query: str, k: int = None ) -> str:
	"""
		Purpose:
		--------
		Build a document-grounded prompt using retrieved excerpts and the current document
		action. Missing retrieval returns a safe prompt rather than failing.

		Parameters:
		-----------
		user_query : str
			User request.

		k : int | None
			Number of chunks to retrieve.

		Returns:
		--------
		str
			Document-grounded LLM prompt.
	"""
	doc_instruction_block = build_instruction_block( )
	hits = retrieve_chunks( user_query, k=k )
	st.session_state[ 'docqna_last_retrieval' ] = hits
	
	context_blocks: List[ str ] = [ ]
	
	for doc_name, chunk, score in hits:
		context_blocks.append( f'[Document: {doc_name}]\n{chunk}'.strip( ) )
	
	semantic_context_buffer = st.session_state.get( 'semantic_context_buffer', [ ] )
	if isinstance( semantic_context_buffer, list ):
		for value in semantic_context_buffer:
			if isinstance( value, str ) and value.strip( ):
				context_blocks.append( f'[Semantic Context]\n{value.strip( )}' )
	
	context = '\n\n'.join( context_blocks ).strip( )
	active_doc_names = get_docqna_names( )
	retrieval_status = str( st.session_state.get( 'docqna_retrieval_status', '' ) or '' ).strip( )
	
	prompt_parts: List[ str ] = [ ]
	
	if doc_instruction_block:
		prompt_parts.append( doc_instruction_block )
	
	prompt_parts.append( f'Active Documents:\n{active_doc_names}' )
	
	if context:
		prompt_parts.append(
			'Use the following retrieved document excerpts as the evidence base for your answer.\n\n'
			f'{context}'
		)
	else:
		if retrieval_status:
			prompt_parts.append(
				'No retrieved document excerpts were available.\n\n'
				f'Retrieval Status: {retrieval_status}'
			)
		else:
			prompt_parts.append(
				'No retrieved document excerpts were available for this question.'
			)
	
	prompt_parts.append( f'User Request:\n{user_query}\n\nAnswer:' )
	return '\n\n'.join( prompt_parts ).strip( )

# ------------- SEMANTIC SEARCH UTILITIES ----------------------

def decode_embedding_rows( ) -> List[ Tuple[ str, np.ndarray ] ]:
	"""
		Purpose:
		--------
		Read and decode rows from the semantic embeddings table. Database failures,
		missing tables, corrupt blobs, and empty vectors fail closed so Semantic Search
		and Text Generation context reuse cannot crash.

		Parameters:
		-----------
		None

		Returns:
		--------
		List[Tuple[str, np.ndarray]]
			Decoded chunk/vector rows.
	"""
	rows_out: List[ Tuple[ str, np.ndarray ] ] = [ ]
	
	try:
		initialize_database( )
		
		with sqlite3.connect( cfg.DB_PATH, timeout=30 ) as conn:
			rows = conn.execute( 'SELECT chunk, vector FROM embeddings' ).fetchall( )
	except Exception as e:
		st.session_state[ 'semantic_status' ] = (
				f'Semantic index could not be read: {e}')
		return rows_out
	
	for chunk_text_value, vector_blob in rows:
		try:
			chunk = str( chunk_text_value or '' )
			vec = decode_embedding_vector( vector_blob )
			
			if not chunk or vec.size == 0:
				continue
			
			rows_out.append( (chunk, vec) )
		except Exception:
			continue
	
	return rows_out

def clear_semantic_index( ) -> None:
	"""
		Purpose:
		--------
		Clear the semantic embeddings table and reset Semantic Search diagnostics without
		raising hard database errors into the Streamlit UI execution path.

		Parameters:
		-----------
		None

		Returns:
		--------
		None
	"""
	try:
		initialize_database( )
		
		with sqlite3.connect( cfg.DB_PATH, timeout=30 ) as conn:
			conn.execute( 'DELETE FROM embeddings' )
			conn.commit( )
		
		st.session_state[ 'semantic_result_rows' ] = [ ]
		st.session_state[ 'semantic_selected_rows' ] = [ ]
		st.session_state[ 'semantic_index_chunk_count' ] = 0
		st.session_state[ 'semantic_index_dim' ] = 0
		st.session_state[ 'semantic_index_doc_count' ] = 0
		st.session_state[ 'semantic_uploaded_names' ] = [ ]
		st.session_state[ 'semantic_last_query' ] = ''
		st.session_state[ 'semantic_status' ] = 'Semantic index deleted.'
	except Exception as e:
		st.session_state[ 'semantic_status' ] = (
				f'Semantic index could not be deleted: {e}')

def build_semantic_index( uploaded_files: List[ Any ] ) -> Dict[ str, Any ]:
	"""
		Purpose:
		--------
		Build or append a semantic chunk index from uploaded files. The function preserves
		the existing embeddings table contract while guarding extraction, embedding, vector
		shape, and SQLite write failures.

		Parameters:
		-----------
		uploaded_files : List[Any]
			Uploaded files from Streamlit.

		Returns:
		--------
		Dict[str, Any]
			Index build result.
	"""
	global embedder
	
	if not is_embedder_available( globals( ).get( 'embedder', None ) ):
		message = get_embedder_unavailable_message( )
		st.session_state[ 'semantic_status' ] = message
		
		return {
				'success': False,
				'message': message,
				'doc_count': 0,
				'chunk_count': 0,
				'vector_dim': 0
		}
	
	if not isinstance( uploaded_files, list ) or len( uploaded_files ) == 0:
		message = 'No files were provided for semantic indexing.'
		st.session_state[ 'semantic_status' ] = message
		
		return {
				'success': False,
				'message': message,
				'doc_count': 0,
				'chunk_count': 0,
				'vector_dim': 0
		}
	
	try:
		chunk_size = int( st.session_state.get( 'semantic_chunk_size', 1200 ) )
		chunk_overlap = int( st.session_state.get( 'semantic_chunk_overlap', 200 ) )
	except Exception:
		chunk_size = 1200
		chunk_overlap = 200
	
	clear_existing = bool( st.session_state.get( 'semantic_clear_existing', True ) )
	append_existing = bool( st.session_state.get( 'semantic_append_existing', False ) )
	
	if append_existing:
		clear_existing = False
	
	all_chunks: List[ str ] = [ ]
	doc_names: List[ str ] = [ ]
	
	for f in uploaded_files:
		try:
			file_name = str( getattr( f, 'name', '' ) or '' ).strip( )
			file_bytes = f.getvalue( )
		except Exception:
			continue
		
		if not file_name or not file_bytes:
			continue
		
		try:
			text = extract_text( file_bytes=file_bytes, file_name=file_name )
		except Exception:
			text = ''
		
		if not text:
			try:
				text = file_bytes.decode( errors='ignore' )
			except Exception:
				text = ''
		
		if not text:
			continue
		
		try:
			chunks = chunk_text( text=text, size=chunk_size, overlap=chunk_overlap )
		except Exception:
			chunks = [ ]
		
		if not chunks:
			continue
		
		all_chunks.extend( [ str( chunk ) for chunk in chunks if chunk ] )
		doc_names.append( file_name )
	
	if len( all_chunks ) == 0:
		message = 'No extractable text was found in the uploaded files.'
		st.session_state[ 'semantic_status' ] = message
		
		return {
				'success': False,
				'message': message,
				'doc_count': 0,
				'chunk_count': 0,
				'vector_dim': 0
		}
	
	try:
		vecs = embedder.encode( all_chunks, show_progress_bar=False )
		vecs = np.asarray( vecs, dtype=np.float32 )
		
		if len( vecs.shape ) != 2 or vecs.shape[ 0 ] != len( all_chunks ):
			message = (
					'Semantic embedding generation returned an unexpected vector shape.')
			st.session_state[ 'semantic_status' ] = message
			
			return {
					'success': False,
					'message': message,
					'doc_count': len( doc_names ),
					'chunk_count': len( all_chunks ),
					'vector_dim': 0
			}
	except Exception as e:
		message = f'Semantic embedding generation failed: {e}'
		st.session_state[ 'semantic_status' ] = message
		
		return {
				'success': False,
				'message': message,
				'doc_count': len( doc_names ),
				'chunk_count': len( all_chunks ),
				'vector_dim': 0
		}
	
	try:
		initialize_database( )
		
		with sqlite3.connect( cfg.DB_PATH, timeout=30 ) as conn:
			if clear_existing:
				conn.execute( 'DELETE FROM embeddings' )
			
			for chunk_text_value, vec in zip( all_chunks, vecs ):
				vec = np.asarray( vec, dtype=np.float32 ).reshape( -1 )
				
				if vec.size == 0:
					continue
				
				conn.execute(
					'INSERT INTO embeddings (chunk, vector) VALUES (?, ?)',
					(chunk_text_value, vec.tobytes( ))
				)
			
			conn.commit( )
	except Exception as e:
		message = f'Semantic index database write failed: {e}'
		st.session_state[ 'semantic_status' ] = message
		
		return {
				'success': False,
				'message': message,
				'doc_count': len( doc_names ),
				'chunk_count': len( all_chunks ),
				'vector_dim': int( vecs.shape[ 1 ] ) if len( vecs.shape ) == 2 else 0
		}
	
	vector_dim = int( vecs.shape[ 1 ] ) if len( vecs.shape ) == 2 else 0
	st.session_state[ 'semantic_uploaded_names' ] = doc_names
	st.session_state[ 'semantic_index_doc_count' ] = len( doc_names )
	st.session_state[ 'semantic_index_chunk_count' ] = len( all_chunks )
	st.session_state[ 'semantic_index_dim' ] = vector_dim
	st.session_state[ 'semantic_result_rows' ] = [ ]
	st.session_state[ 'semantic_selected_rows' ] = [ ]
	st.session_state[ 'semantic_status' ] = 'Semantic index built successfully.'
	
	return {
			'success': True,
			'message': 'Semantic index built successfully.',
			'doc_count': len( doc_names ),
			'chunk_count': len( all_chunks ),
			'vector_dim': vector_dim
	}

def query_semantic_index( query_text: str ) -> List[ Dict[ str, Any ] ]:
	"""
		Purpose:
		--------
		Query the semantic index and return ranked chunk results. Missing embeddings,
		database failures, empty indexes, malformed vectors, and vector dimension mismatches
		fail closed instead of raising hard runtime errors.

		Parameters:
		-----------
		query_text : str
			Query text.

		Returns:
		--------
		List[Dict[str, Any]]
			Ranked semantic result rows.
	"""
	global embedder
	
	if not query_text or not str( query_text ).strip( ):
		st.session_state[ 'semantic_result_rows' ] = [ ]
		st.session_state[ 'semantic_status' ] = 'Enter a semantic query before searching.'
		return [ ]
	
	if not is_embedder_available( globals( ).get( 'embedder', None ) ):
		st.session_state[ 'semantic_result_rows' ] = [ ]
		st.session_state[ 'semantic_status' ] = get_embedder_unavailable_message( )
		return [ ]
	
	try:
		top_k = int( st.session_state.get( 'semantic_top_k', 8 ) )
	except Exception:
		top_k = 8
	
	try:
		min_similarity = float( st.session_state.get( 'semantic_min_similarity', 0.0 ) )
	except Exception:
		min_similarity = 0.0
	
	rows = decode_embedding_rows( )
	if not rows:
		st.session_state[ 'semantic_result_rows' ] = [ ]
		if not st.session_state.get( 'semantic_status', '' ):
			st.session_state[ 'semantic_status' ] = 'Semantic index is empty.'
		return [ ]
	
	try:
		q = embedder.encode( [ str( query_text ).strip( ) ], show_progress_bar=False )[ 0 ]
		q = np.asarray( q, dtype=np.float32 ).reshape( -1 )
	except Exception as e:
		st.session_state[ 'semantic_result_rows' ] = [ ]
		st.session_state[ 'semantic_status' ] = f'Semantic query embedding failed: {e}'
		return [ ]
	
	if q.size == 0:
		st.session_state[ 'semantic_result_rows' ] = [ ]
		st.session_state[ 'semantic_status' ] = (
				'Semantic query embedding returned an empty vector.')
		return [ ]
	
	scored_rows: List[ Dict[ str, Any ] ] = [ ]
	skipped_rows = 0
	
	for idx, (chunk_text_value, vec) in enumerate( rows, start=1 ):
		try:
			vec = np.asarray( vec, dtype=np.float32 ).reshape( -1 )
			
			if vec.size == 0 or vec.size != q.size:
				skipped_rows += 1
				continue
			
			score = cosine_similarity( q, vec )
		except Exception:
			skipped_rows += 1
			continue
		
		if score < min_similarity:
			continue
		
		scored_rows.append(
			{
					'Selected': False,
					'Rank': idx,
					'Score': float( score ),
					'Chunk': str( chunk_text_value or '' ),
					'Length': len( str( chunk_text_value or '' ) )
			}
		)
	
	scored_rows.sort( key=lambda r: r[ 'Score' ], reverse=True )
	scored_rows = scored_rows[ :top_k ]
	
	st.session_state[ 'semantic_last_query' ] = str( query_text ).strip( )
	st.session_state[ 'semantic_result_rows' ] = scored_rows
	
	if scored_rows:
		st.session_state[ 'semantic_status' ] = 'Semantic search completed.'
	elif skipped_rows:
		st.session_state[ 'semantic_status' ] = (
				'No semantic matches found. Some stored vectors were skipped because '
				'their dimensions did not match the active embedding model.')
	else:
		st.session_state[ 'semantic_status' ] = 'No semantic matches found.'
	
	return scored_rows

def create_semantic_context( ) -> str:
	"""
		Purpose:
		--------
		Build a semantic-context text block from selected search rows.

		Parameters:
		-----------
		None

		Returns:
		--------
		str
			Semantic context text.
	"""
	selected_rows = st.session_state.get( 'semantic_selected_rows', [ ] )
	
	if not isinstance( selected_rows, list ) or len( selected_rows ) == 0:
		return ''
	
	context_parts: List[ str ] = [ ]
	
	for idx, row in enumerate( selected_rows, start=1 ):
		chunk_text_value = str( row.get( 'Chunk', '' ) or '' ).strip( )
		score_value = row.get( 'Score', '' )
		
		if not chunk_text_value:
			continue
		
		context_parts.append( f'[Semantic Chunk {idx} | Score: {score_value}]\n{chunk_text_value}' )
	
	return '\n\n'.join( context_parts ).strip( )

def extract_selected_rows( edited_rows: Any ) -> List[ Dict[ str, Any ] ]:
	"""
		Purpose:
		--------
		Extract selected semantic rows from a data_editor result payload.

		Parameters:
		-----------
		edited_rows : Any
			Data editor result payload.

		Returns:
		--------
		List[Dict[str, Any]]
			Selected rows.
	"""
	selected: List[ Dict[ str, Any ] ] = [ ]
	
	if isinstance( edited_rows, pd.DataFrame ):
		for _, row in edited_rows.iterrows( ):
			row_dict = row.to_dict( )
			if bool( row_dict.get( 'Selected', False ) ):
				selected.append( row_dict )
		
		return selected
	
	if not isinstance( edited_rows, list ):
		return selected
	
	for row in edited_rows:
		if isinstance( row, dict ) and bool( row.get( 'Selected', False ) ):
			selected.append( row )
	
	return selected

def send_text_chunks( ) -> None:
	"""
		Purpose:
		--------
		Push selected semantic chunks into the shared basic document context buffer.

		Parameters:
		-----------
		None

		Returns:
		--------
		None
	"""
	context_text = create_semantic_context( )
	
	if not context_text:
		return
	
	existing_docs = st.session_state.get( 'basic_docs', [ ] )
	
	if not isinstance( existing_docs, list ):
		existing_docs = [ ]
	
	existing_docs.append( context_text )
	st.session_state[ 'basic_docs' ] = existing_docs
	st.session_state[ 'use_semantic' ] = True

def send_docqna_chunks( ) -> None:
	"""
		Purpose:
		--------
		Push selected semantic chunks into the shared document context buffer used by
		Document Q&A prompts.

		Parameters:
		-----------
		None

		Returns:
		--------
		None
	"""
	context_text = create_semantic_context( )
	
	if not context_text:
		return
	
	buffer_rows = st.session_state.get( 'semantic_context_buffer', [ ] )
	
	if not isinstance( buffer_rows, list ):
		buffer_rows = [ ]
	
	buffer_rows.append( context_text )
	st.session_state[ 'semantic_context_buffer' ] = buffer_rows

# ==============================================================================
# Init
# ==============================================================================
initialize_database( )
embedder = load_embedder( )

if not isinstance( st.session_state.get( 'messages' ), list ):
	st.session_state[ 'messages' ] = [ ]

if len( st.session_state[ 'messages' ] ) == 0:
	st.session_state[ 'messages' ] = load_history( )

if 'system_instructions' not in st.session_state:
	st.session_state[ 'system_instructions' ] = ''

# ==============================================================================
# SIDEBAR
# ==============================================================================
sidebar_model_name = str( st.session_state.get( 'selected_model_name', get_default_model_name( ) ) or
	get_default_model_name( ) )

render_selected_model_logo( sidebar_model_name, size='large' )

with st.sidebar:
	style_subheaders( )
	
	c1, c2 = st.columns( [ 0.05, 0.95 ] )
	with c2:
		st.divider( )
		st.markdown( '#### ⚙️ Configuration' )
		
		with st.expander( 'LLM', expanded=True ):
			model_names = get_model_names_for_state( )
			default_model_name = get_default_model_name( )
			
			if default_model_name in model_names:
				default_model_index = model_names.index( default_model_name )
			else:
				default_model_index = 0
			
			current_model_name = str(
				st.session_state.get( 'selected_model_name', default_model_name ) or
				default_model_name
			)
			
			if current_model_name in model_names:
				current_model_index = model_names.index( current_model_name )
			else:
				current_model_index = default_model_index
			
			st.radio( label='Select', options=model_names, index=current_model_index, key='selected_model_name',
				on_change=on_selected_model_change )
			
			selected_model_name = str( st.session_state.get( 'selected_model_name',
				default_model_name ) or default_model_name )
			
			selected_model_path = get_model_path_for_state( selected_model_name )
			selected_model_spec = get_model_spec_for_state( selected_model_name )
			selected_model_modes = get_model_modes_for_state( selected_model_name )
			
			st.session_state[ 'selected_model_path' ] = selected_model_path
			st.session_state[ 'selected_model_spec' ] = selected_model_spec
			st.session_state[ 'selected_model_modes' ] = selected_model_modes
			
			if isinstance( selected_model_spec, dict ):
				model_family = str( selected_model_spec.get( 'family', '' ) or '' )
				base_model = str( selected_model_spec.get( 'base_model', '' ) or '' )
				model_size = str( selected_model_spec.get( 'size', '' ) or '' )
				model_description = str( selected_model_spec.get( 'description', '' ) or '' )
			else:
				model_family = ''
				model_size = ''
				model_description = ''
			
			if model_family or model_size:
				st.caption( f'Model Family: '
				            f'{model_family}'.strip( ) )
				st.caption( f'Base Model: '
				            f'{base_model}'.strip( ) )
				st.caption(f'Parameters: '
				           f'{model_size}'.strip( ) )
			
			if model_description:
				st.caption( model_description )
			
			if selected_model_path:
				if Path( selected_model_path ).exists( ):
					st.caption( 'Model File: Available' )
				else:
					st.caption( 'Model File: Missing' )
			else:
				st.caption( 'Model File: Not Configured' )
		
		with st.expander( 'Mode', expanded=False ):
			selected_model_name = str( st.session_state.get( 'selected_model_name',
				get_default_model_name( ) ) or get_default_model_name( ) )
			
			model_modes = get_model_modes_for_state( selected_model_name )
			
			pending_selected_mode = st.session_state.pop( 'pending_selected_mode', None )
			
			if pending_selected_mode and pending_selected_mode in model_modes:
				st.session_state[ 'selected_mode' ] = pending_selected_mode
				st.session_state[ 'mode' ] = pending_selected_mode
			
			current_mode = str( st.session_state.get( 'selected_mode',
				get_default_mode_name( selected_model_name ) ) or
			                    get_default_mode_name( selected_model_name ) )
			
			if current_mode not in model_modes:
				current_mode = model_modes[ 0 ] if model_modes else get_default_mode_name( selected_model_name )
				
				st.session_state[ 'selected_mode' ] = current_mode
				st.session_state[ 'mode' ] = current_mode
			
			current_mode_index = ( model_modes.index(
				current_mode ) if current_mode in model_modes else 0 )
			
			st.radio( label='Select', options=model_modes, index=current_mode_index, key='selected_mode',
				on_change=on_selected_mode_change )
			
			mode = str( st.session_state.get( 'selected_mode',
					get_default_mode_name( selected_model_name ) )
			            or get_default_mode_name( selected_model_name ) )
			
			st.session_state[ 'mode' ] = mode
			
# ==============================================================================
# TEXT GENERATION MODE
# ==============================================================================
if mode == 'Text Generation':
	messages = st.session_state.get( 'messages', [ ] )
	max_tokens = st.session_state.get( 'max_tokens', 0 )
	top_percent = st.session_state.get( 'top_percent', 0.0 )
	top_k = st.session_state.get( 'top_k', 0 )
	temperature = st.session_state.get( 'temperature', 0.0 )
	is_grounded = st.session_state.get( 'is_grounded', False )
	frequency_penalty = st.session_state.get( 'frequency_penalty', 0.0 )
	presense_penalty = st.session_state.get( 'presense_penalty', 0.0 )
	repeat_penalty = st.session_state.get( 'repeat_penalty', 0.0 )
	repeat_window = st.session_state.get( 'repeat_window', 0.0 )
	cpu_threads = st.session_state.get( 'cpu_threads', cfg.CORES )
	context_window = st.session_state.get( 'context_window', cfg.DEFAULT_CTX )
	left, center, right = st.columns( [ 0.05, 0.9, 0.05 ] )
	with center:
		st.subheader( '💬 Text Generation', help=cfg.TEXT_GENERATION )
		st.divider( )
		
		# ------------------------------------------------------------------
		# Expander — Mind Controls
		# ------------------------------------------------------------------
		with st.expander( label='Mind Controls', icon='🧠', expanded=False ):
			text_reset_defaults: Dict[ str, Dict[ str, Any ] ] = {
					'task_preset_reset':
						{
								'task_preset': 'Chat',
								'response_format': 'Markdown',
								'use_chat_history': True,
								'use_document_context': False
						},
					'reasoning_controls_reset':
						{
								'reasoning_depth': 'Medium',
								'answer_only': False,
								'use_self_check': False,
								'deterministic_reasoning': False
						},
					'coding_controls_reset':
						{
								'coding_language': 'Python',
								'coding_task': 'Generate',
								'coding_include_comments': True,
								'coding_editor_format': True,
								'coding_fenced_output': True,
								'translation_target_language': 'English'
						},
					'response_controls_reset':
						{
								'top_k': 0,
								'top_percent': 0.95,
								'temperature': 0.0,
								'is_grounded': False
						},
					'probability_controls_reset':
						{
								'frequency_penalty': 0.0,
								'presense_penalty': 0.0,
								'repeat_penalty': 1.1,
								'repeat_window': 0
						},
					'context_controls_reset':
						{
								'random_seed': 0,
								'max_tokens': 1024,
								'cpu_threads': int( cfg.CORES ),
								'context_window': int( cfg.DEFAULT_CTX )
						}
			}
			
			def request_text_generation_reset( reset_name: str ) -> None:
				"""
				Purpose:
				--------
				Request a Text Generation control reset without directly modifying any
				widget-owned keys after their widgets have been instantiated.
		
				Parameters:
				-----------
				reset_name : str
					Name of the reset group to process on the next safe script pass.
		
				Returns:
				--------
				None
				"""
				st.session_state[ 'pending_text_generation_reset' ] = str( reset_name or '' )
			pending_text_reset = st.session_state.pop( 'pending_text_generation_reset', None )
			if pending_text_reset:
				reset_values = text_reset_defaults.get( str( pending_text_reset ), { } )
				if isinstance( reset_values, dict ):
					for reset_key, reset_value in reset_values.items( ):
						st.session_state[ reset_key ] = reset_value
			
			# ------------------------------------------------------------------
			# Task Preset
			# ------------------------------------------------------------------
			with st.expander( label='Task Preset', icon='🧭', expanded=False ):
				task_c1, task_c2, task_c3, task_c4 = st.columns( [ 0.25, 0.25, 0.25, 0.25 ], border=True,
					gap='medium' )
				
				with task_c1:
					st.selectbox( label='Task Type', options=[ 'Chat', 'Reasoning', 'Coding',
							'Translation', 'Summarization', 'Extraction' ], key='task_preset' )
				
				with task_c2:
					st.selectbox( label='Response Format',
						options=[ 'Plain Text', 'Markdown', 'Bullet Summary', 'JSON' ],
						key='response_format' )
				
				with task_c3:
					st.toggle( label='Use Conversation History',
						value=bool( st.session_state.get( 'use_chat_history', True ) ),
						key='use_chat_history' )
				
				with task_c4:
					st.toggle( label='Use Document Context',
						value=bool( st.session_state.get( 'use_document_context', False ) ),
						key='use_document_context' )
				
				st.button( label='Reset', key='task_preset_reset', width='stretch',
					on_click=request_text_generation_reset, args=('task_preset_reset',) )
			
			# ------------------------------------------------------------------
			# Reasoning Controls
			# ------------------------------------------------------------------
			with st.expander( label='Reasoning Controls', icon='🧩', expanded=False ):
				reason_c1, reason_c2, reason_c3, reason_c4 = st.columns( [ 0.25, 0.25, 0.25, 0.25 ],
					border=True, gap='medium' )
				
				with reason_c1:
					st.selectbox( label='Reasoning Depth',
						options=[ 'Low', 'Medium', 'High' ], key='reasoning_depth' )
				
				with reason_c2:
					st.toggle( label='Answer Only',
						value=bool( st.session_state.get( 'answer_only', False ) ), key='answer_only'  )
				
				with reason_c3:
					st.toggle( label='Use Self-Check',
						value=bool( st.session_state.get( 'use_self_check', False ) ),
						key='use_self_check' )
				
				with reason_c4:
					st.toggle( label='Prefer Deterministic Reasoning', value=bool(
						st.session_state.get( 'deterministic_reasoning', False ) ),
						key='deterministic_reasoning' )
				
				st.button( label='Reset', key='reasoning_controls_reset',
					width='stretch', on_click=request_text_generation_reset,
					args=('reasoning_controls_reset',) )
			
			# ------------------------------------------------------------------
			# Coding Controls
			# ------------------------------------------------------------------
			with st.expander( label='Coding Controls', icon='🧾', expanded=False ):
				code_c1, code_c2, code_c3, code_c4, code_c5 = st.columns( [ 0.2, 0.2, 0.2, 0.2, 0.2 ],
					border=True, gap='medium' )
				
				with code_c1:
					st.selectbox( label='Code Language', options=[ 'Python', 'C#', 'SQL', 'VBA',
							'JavaScript', 'Markdown' ], key='coding_language' )
				
				with code_c2:
					st.selectbox( label='Coding Task',
						options=[ 'Generate', 'Refactor', 'Explain', 'Debug', 'Review' ],
						key='coding_task' )
				
				with code_c3:
					st.toggle( label='Include Comments',
						value=bool( st.session_state.get( 'coding_include_comments', True ) ),
						key='coding_include_comments' )
				
				with code_c4:
					st.toggle( label='Use Editor Format',
						value=bool( st.session_state.get( 'coding_editor_format', True ) ),
						key='coding_editor_format' )
				
				with code_c5:
					st.toggle( label='Emit Fenced Code',
						value=bool( st.session_state.get( 'coding_fenced_output', True ) ),
						key='coding_fenced_output' )
				
				translation_col_left, translation_col_right = st.columns( [ 0.5, 0.5 ] )
				
				with translation_col_left:
					st.text_input( label='Translation Target Language',
						key='translation_target_language' )
				
				with translation_col_right:
					st.markdown( '<br>', unsafe_allow_html=True )
					st.button( label='Reset', key='coding_controls_reset',
						width='stretch', on_click=request_text_generation_reset,
						args=('coding_controls_reset',) )
			
			# ------------------------------------------------------------------
			# Advanced Model Capabilities
			# ------------------------------------------------------------------
			with st.expander( label='Advanced Capabilities', icon='🧬', expanded=False ):
				capabilities = get_active_model_capabilities( )
				advanced_c1, advanced_c2, advanced_c3, advanced_c4 = st.columns(
					[ 0.25, 0.25, 0.25, 0.25 ], border=True, gap='medium' )
				
				with advanced_c1:
					st.toggle( label='Enable Thinking',
						value=bool( st.session_state.get( 'thinking_mode_enabled', False ) ),
						key='thinking_mode_enabled',
						disabled=not model_supports_capability( 'thinking' ) )
				
				with advanced_c2:
					st.selectbox( label='Thinking Effort', options=[ 'Low', 'Medium', 'High' ],
						key='thinking_effort', disabled=not model_supports_capability( 'thinking' ) )
				
				with advanced_c3:
					st.toggle( label='Reasoning Summary',
						value=bool( st.session_state.get( 'thinking_summary_enabled', True ) ),
						key='thinking_summary_enabled',
						disabled=not model_supports_capability( 'thinking' ) )
				
				with advanced_c4:
					st.toggle( label='Enable Advanced Coding',
						value=bool( st.session_state.get( 'coding_mode_enabled', False ) ),
						key='coding_mode_enabled',
						disabled=not model_supports_capability( 'coding' ) )
				
				coding_ext_c1, coding_ext_c2 = st.columns( [ 0.5, 0.5 ], border=True, gap='medium' )
				
				with coding_ext_c1:
					st.toggle( label='Include Test Strategy',
						value=bool( st.session_state.get( 'coding_test_request', False ) ),
						key='coding_test_request',
						disabled=not model_supports_capability( 'coding' ) )
				
				with coding_ext_c2:
					st.toggle( label='Explain Implementation',
						value=bool( st.session_state.get( 'coding_explain_request', False ) ),
						key='coding_explain_request',
						disabled=not model_supports_capability( 'coding' ) )
				
				st.divider( )
				
				function_c1, function_c2 = st.columns( [ 0.35, 0.65 ], border=True, gap='medium' )
				
				with function_c1:
					st.toggle( label='Enable Function Calling',
						value=bool( st.session_state.get( 'function_call_enabled', False ) ),
						key='function_call_enabled',
						disabled=not model_supports_capability( 'function_calling' ) )
					
					st.text_area( label='Function Call Guidance', height=120,
						key='function_call_guidance',
						placeholder='Optional guidance for when the model should emit a tool call.',
						disabled=not model_supports_capability( 'function_calling' ) )
				
				with function_c2:
					st.text_area( label='Function Schema JSON', height=220,
						key='function_schema_text',
						disabled=not model_supports_capability( 'function_calling' ) )
				
				status_rows: List[ Dict[ str, Any ] ] = [ { 'Capability': 'Thinking',
						'Enabled': bool( st.session_state.get( 'thinking_mode_enabled', False ) ),
						'Supported': bool( capabilities.get( 'thinking', False ) ) },
						{ 'Capability': 'Advanced Coding',
								'Enabled': bool( st.session_state.get( 'coding_mode_enabled',
									False ) ),
								'Supported': bool( capabilities.get( 'coding', False ) ) },
						{ 'Capability': 'Function Calling',
								'Enabled': bool( st.session_state.get( 'function_call_enabled',
									False ) ),
								'Supported': bool( capabilities.get( 'function_calling', False )
								) },
						{ 'Capability': 'Web Browsing', 'Enabled': False,
								'Supported': bool( capabilities.get( 'web_browsing', False ) ) } ]
				
				df_capabilities = pd.DataFrame( status_rows )
				st.dataframe( df_capabilities, use_container_width=True, hide_index=True )
				
				if not model_supports_capability( 'thinking' ):
					st.caption( get_capability_status_message( 'thinking' ) )
				
				if not model_supports_capability( 'coding' ):
					st.caption( get_capability_status_message( 'coding' ) )
				
				if not model_supports_capability( 'function_calling' ):
					st.caption( get_capability_status_message( 'function_calling' ) )
				
				if st.button( label='Reset Advanced Capabilities',
						key='advanced_capabilities_reset', width='stretch' ):
					st.session_state[ 'thinking_mode_enabled' ] = False
					st.session_state[ 'thinking_effort' ] = 'Medium'
					st.session_state[ 'thinking_summary_enabled' ] = True
					st.session_state[ 'coding_mode_enabled' ] = False
					st.session_state[ 'coding_test_request' ] = False
					st.session_state[ 'coding_explain_request' ] = False
					st.session_state[ 'function_call_enabled' ] = False
					st.session_state[ 'function_call_prompt' ] = ''
					st.session_state[ 'function_call_guidance' ] = ''
					st.session_state[ 'function_schema_text' ] = get_default_function_schema_text( )
					st.session_state[ 'function_call_response' ] = ''
					st.session_state[ 'function_call_result' ] = ''
					st.session_state[ 'function_call_status' ] = ''
					st.rerun( )
			
			# ------------------------------------------------------------------
			# Gipity Tools and Web Browsing
			# ------------------------------------------------------------------
			with st.expander( label='Tools & Web', icon='🧰', expanded=False ):
				tools_supported = bool( model_supports_capability( 'function_calling' ) )
				web_supported = bool( model_supports_capability( 'web_browsing' ) )
				
				tool_status_c1, tool_status_c2, tool_status_c3 = st.columns(
					[ 0.34, 0.33, 0.33 ], border=True, gap='medium' )
				
				with tool_status_c1:
					st.metric( 'Selected Model', get_selected_model_name( ) )
				
				with tool_status_c2:
					st.metric( 'Function Calling',
						'Enabled' if tools_supported else 'Unavailable' )
				
				with tool_status_c3:
					st.metric( 'Web Browsing',
						'Enabled' if web_supported else 'Unavailable' )
				
				if not tools_supported:
					st.info( get_capability_status_message( 'function_calling' ) )
				
				if not web_supported:
					st.info( get_capability_status_message( 'web_browsing' ) )
				
				tool_c1, tool_c2 = st.columns( [ 0.45, 0.55 ], border=True, gap='medium' )
				
				with tool_c1:
					st.text_area(
						label='Tool Task',
						height=150,
						key='function_call_prompt',
						placeholder=(
								'Example: Fetch https://example.com and summarize the page '
								'in five bullets.'),
						disabled=not tools_supported
					)
					
					generate_tool = st.button(
						label='Generate Function Call',
						key='generate_function_call_button',
						width='stretch',
						disabled=not tools_supported
					)
					
					execute_tool = st.button(
						label='Execute Function Call',
						key='execute_function_call_button',
						width='stretch',
						disabled=not tools_supported
					)
					
					final_answer_tool = st.button(
						label='Generate Final Answer from Tool Result',
						key='tool_final_answer_button',
						width='stretch',
						disabled=not tools_supported
					)
				
				with tool_c2:
					st.text_area(
						label='Generated Function Call JSON',
						height=180,
						key='function_call_model_json',
						disabled=not tools_supported
					)
					
					st.text_area(
						label='Tool Execution Result',
						height=180,
						key='function_call_result',
						disabled=True
					)
				
				if generate_tool:
					try:
						task_value = str(
							st.session_state.get( 'function_call_prompt', '' ) or '' ).strip( )
						
						response = generate_function_call_json( task_value )
						st.session_state[ 'function_call_model_json' ] = response
						st.session_state[ 'function_call_response' ] = response
						st.session_state[ 'function_call_status' ] = (
								'Function call generated.')
						st.success( st.session_state[ 'function_call_status' ] )
					except Exception as e:
						st.session_state[ 'function_call_status' ] = (
								f'Function-call generation failed: {e}')
						st.error( st.session_state[ 'function_call_status' ] )
				
				if execute_tool:
					try:
						tool_text = str(
							st.session_state.get( 'function_call_model_json', '' ) or '' ).strip( )
						result = execute_tool_call_text( tool_text )
						
						st.session_state[ 'function_call_result' ] = str(
							result.get( 'result', '' ) or '' )
						st.session_state[ 'function_call_status' ] = (
								f'Executed allowlisted function: {result.get( "name", "" )}')
						
						if str( result.get( 'name', '' ) ) == 'web_browse_url':
							st.session_state[ 'web_browse_context_buffer' ] = str(
								result.get( 'result', '' ) or '' )
						
						st.success( st.session_state[ 'function_call_status' ] )
					except Exception as e:
						st.session_state[ 'function_call_status' ] = (
								f'Function-call execution failed: {e}')
						st.error( st.session_state[ 'function_call_status' ] )
				
				if final_answer_tool:
					try:
						task_value = str(
							st.session_state.get( 'function_call_prompt', '' ) or '' ).strip( )
						result_text = str(
							st.session_state.get( 'function_call_result', '' ) or '' ).strip( )
						
						if not result_text:
							st.info( 'Execute a function call before generating the final answer.' )
						else:
							tool_result = {
									'name': 'app_tool',
									'arguments': { },
									'result': result_text
							}
							final_answer = generate_tool_grounded_final_answer(
								user_task=task_value,
								tool_result=tool_result
							)
							
							st.session_state[ 'function_call_response' ] = final_answer
							st.session_state[ 'function_call_status' ] = (
									'Final answer generated from tool result.')
							st.markdown( '### Tool-Grounded Final Answer' )
							st.markdown( final_answer )
					except Exception as e:
						st.session_state[ 'function_call_status' ] = (
								f'Final answer generation failed: {e}')
						st.error( st.session_state[ 'function_call_status' ] )
				
				status_value = str(
					st.session_state.get( 'function_call_status', '' ) or '' ).strip( )
				if status_value:
					st.caption( status_value )
				
				st.divider( )
				
				web_c1, web_c2 = st.columns( [ 0.45, 0.55 ], border=True, gap='medium' )
				
				with web_c1:
					st.text_input(
						label='Web URL',
						key='web_browse_url',
						placeholder='https://example.com/article',
						disabled=not web_supported
					)
					
					st.text_input(
						label='Allowed Domain Optional',
						key='web_browse_allow_domain',
						placeholder='example.com',
						disabled=not web_supported
					)
					
					st.text_area(
						label='Web Browse Prompt',
						height=120,
						key='web_browse_prompt',
						placeholder='Summarize the page and identify the most important facts.',
						disabled=not web_supported
					)
					
					browse_web = st.button(
						label='Fetch Web Context',
						key='web_browse_fetch_button',
						width='stretch',
						disabled=not web_supported
					)
					
					send_web = st.button(
						label='Send Web Context to Text Generation',
						key='web_browse_send_context_button',
						width='stretch',
						disabled=not web_supported
					)
				
				with web_c2:
					st.text_area(
						label='Fetched Web Context',
						height=340,
						key='web_browse_result',
						disabled=True
					)
				
				if browse_web:
					try:
						context_text = web_browse_url_tool(
							url=str( st.session_state.get( 'web_browse_url', '' ) or '' ),
							prompt=str( st.session_state.get( 'web_browse_prompt', '' ) or '' ),
							allowed_domain=str(
								st.session_state.get( 'web_browse_allow_domain', '' ) or '' ),
							max_chars=12000
						)
						
						st.session_state[ 'web_browse_result' ] = context_text
						st.session_state[ 'web_browse_context_buffer' ] = context_text
						st.session_state[ 'web_browse_status' ] = 'Web context fetched.'
						st.success( st.session_state[ 'web_browse_status' ] )
					except Exception as e:
						st.session_state[ 'web_browse_status' ] = f'Web fetch failed: {e}'
						st.error( st.session_state[ 'web_browse_status' ] )
				
				if send_web:
					context_text = str(
						st.session_state.get( 'web_browse_context_buffer', '' ) or '' ).strip( )
					
					if context_text:
						send_web_context_to_text_generation( context_text )
						st.success( 'Web context added to Text Generation context.' )
					else:
						st.info( 'No web context is available to send.' )
				
				web_status = str(
					st.session_state.get( 'web_browse_status', '' ) or '' ).strip( )
				if web_status:
					st.caption( web_status )
				
				if st.button( label='Clear Tool and Web State',
						key='clear_tool_web_state_button', width='stretch' ):
					st.session_state[ 'function_call_prompt' ] = ''
					st.session_state[ 'function_call_response' ] = ''
					st.session_state[ 'function_call_result' ] = ''
					st.session_state[ 'function_call_status' ] = ''
					st.session_state[ 'function_call_model_json' ] = ''
					st.session_state[ 'web_browse_url' ] = ''
					st.session_state[ 'web_browse_allow_domain' ] = ''
					st.session_state[ 'web_browse_prompt' ] = ''
					st.session_state[ 'web_browse_result' ] = ''
					st.session_state[ 'web_browse_status' ] = ''
					st.session_state[ 'web_browse_context_buffer' ] = ''
					st.rerun( )
					
			# ------------------------------------------------------------------
			# Response Controls
			# ------------------------------------------------------------------
			with st.expander( label='Response Controls', icon='↔️', expanded=False ):
				mind_c1, mind_c2, mind_c3, mind_c4 = st.columns( [ 0.25, 0.25, 0.25, 0.25 ],
					border=True, gap='medium' )
				
				with mind_c1:
					st.slider( label='Temperature', min_value=0.0,
						max_value=1.0, help=cfg.TEMPERATURE,
						key='temperature' )
					temperature = st.session_state[ 'temperature' ]
				
				with mind_c2:
					st.slider( label='Top-P', min_value=0.0, max_value=1.0,
						step=0.01, key='top_percent', help=cfg.TOP_P )
					top_percent = st.session_state[ 'top_percent' ]
				
				with mind_c3:
					st.slider( label='Top-K', min_value=0, max_value=50,
						step=1, key='top_k', help=cfg.TOP_K )
					top_k = st.session_state[ 'top_k' ]
				
				with mind_c4:
					st.toggle( label='Use Grounding',
						value=bool( st.session_state.get( 'is_grounded', False ) ),
						key='is_grounded' )
					is_grounded = st.session_state[ 'is_grounded' ]
				
				st.button( label='Reset', key='response_controls_reset', width='stretch',
					on_click=request_text_generation_reset, args=('response_controls_reset',) )
			
			# ------------------------------------------------------------------
			# Inference Settings
			# ------------------------------------------------------------------
			with st.expander( label='Inference Settings', icon='🎚️', expanded=False ):
				prob_c1, prob_c2, prob_c3, prob_c4 = st.columns( [ 0.25, 0.25, 0.25, 0.25 ],
					border=True, gap='medium' )
				
				with prob_c1:
					st.slider( label='Repeat Window', min_value=0, max_value=1024,
						step=16, key='repeat_window',
						help=cfg.REPEAT_WINDOW )
					repeat_window = st.session_state[ 'repeat_window' ]
				
				with prob_c2:
					st.slider( label='Repeat Penalty', min_value=0.0, max_value=2.0,
						key='repeat_penalty', step=0.05,
						help=cfg.REPEAT_PENALTY )
					repeat_penalty = st.session_state[ 'repeat_penalty' ]
				
				with prob_c3:
					st.slider( label='Presence Penalty', min_value=0.0,
						max_value=2.0, key='presense_penalty', step=0.05,
						help=cfg.PRESENCE_PENALTY )
					presense_penalty = st.session_state[ 'presense_penalty' ]
				
				with prob_c4:
					st.slider( label='Frequency Penalty', min_value=0.0, max_value=2.0,
						key='frequency_penalty', step=0.05, help=cfg.FREQUENCY_PENALTY )
					frequency_penalty = st.session_state[ 'frequency_penalty' ]
				
				st.button( label='Reset', key='probability_controls_reset', width='stretch',
					on_click=request_text_generation_reset, args=('probability_controls_reset',) )
			
			# ------------------------------------------------------------------
			# Context Controls
			# ------------------------------------------------------------------
			with st.expander( label='Context Controls', icon='🎛️', expanded=False ):
				ctx_c1, ctx_c2, ctx_c3, ctx_c4 = st.columns( [ 0.25, 0.25, 0.25, 0.25 ], border=True,
					gap='medium' )
				
				with ctx_c1:
					st.slider( label='Context Window', min_value=0, max_value=8192,
						key='context_window', step=512,
						help=cfg.CONTEXT_WINDOW )
					context_window = st.session_state[ 'context_window' ]
				
				with ctx_c2:
					st.slider( label='CPU Threads', min_value=0, max_value=cfg.CORES,
						key='cpu_threads', step=1, help=cfg.CPU_CORES )
					cpu_threads = st.session_state[ 'cpu_threads' ]
				
				with ctx_c3:
					st.slider( label='Max Tokens', min_value=0, max_value=4096,
						step=128, key='max_tokens', help=cfg.MAX_TOKENS )
					max_tokens = st.session_state[ 'max_tokens' ]
				
				with ctx_c4:
					st.slider( label='Random Seed', min_value=0, max_value=4096,
						step=1, key='random_seed', help=cfg.SEED )
				
				st.button( label='Reset', key='context_controls_reset', width='stretch',
					on_click=request_text_generation_reset, args=('context_controls_reset',) )
				
		# ------------------------------------------------------------------
		# Expander — System Instructions
		# ------------------------------------------------------------------
		with st.expander( label='System Instructions', icon='🖥️', expanded=False, width='stretch' ):
			render_system_instructions( prefix='text', include_apply_preset=True,
				include_preview=True )
			
		st.markdown( cfg.BLUE_DIVIDER, unsafe_allow_html=True )
	
		# ------------------------------------------------------------------
		# Chat History Render
		# ------------------------------------------------------------------
		if 'messages' not in st.session_state or not isinstance( st.session_state.messages, list ):
			st.session_state.messages = [ ]
		
		for msg in st.session_state.messages:
			role = ''
			content = ''
			
			if isinstance( msg, dict ):
				role = str( msg.get( 'role', '' ) or '' ).strip( )
				content = msg.get( 'content', '' )
			
			elif isinstance( msg, tuple ) or isinstance( msg, list ):
				if len( msg ) == 2:
					role = str( msg[ 0 ] or '' ).strip( )
					content = msg[ 1 ]
			
			if role not in ('user', 'assistant', 'system'):
				continue
			
			if content is None:
				content = ''
			elif not isinstance( content, str ):
				content = str( content )
			
			with st.chat_message( role ):
				st.markdown( content )
		
		# ------------------------------------------------------------------
		# Chat Input
		# ------------------------------------------------------------------
		user_input = st.chat_input( 'Ask Loca…' )
		if user_input and isinstance( user_input, str ) and user_input.strip( ):
			user_input = user_input.strip( )
			st.session_state[ 'last_preview_input' ] = str( user_input )
			
			if 'messages' not in st.session_state or not isinstance( st.session_state.messages, list ):
				st.session_state.messages = [ ]
			
			save_message( 'user', user_input )
			st.session_state.messages.append( ('user', user_input) )
			
			with st.chat_message( 'user' ):
				st.markdown( user_input )
			
			with st.chat_message( 'assistant' ):
				out = st.empty( )
				buf = run_llm_turn( user_input=user_input,
					temperature=float( st.session_state.get( 'temperature', 0.0 ) ),
					top_p=float( st.session_state.get( 'top_percent', 0.95 ) ),
					repeat_penalty=float( st.session_state.get( 'repeat_penalty', 1.1 ) ),
					max_tokens=int( st.session_state.get( 'max_tokens', 1024 ) ) or 1024,
					stream=True, output=out )
			
			if buf is None:
				buf = ''
			elif not isinstance( buf, str ):
				buf = str( buf )
			
			buf = buf.strip( )
			save_message( 'assistant', buf )
			st.session_state.messages.append( ('assistant', buf) )
		
		if st.button( '🧹 Clear Chat' ):
			clear_history( )
			st.session_state.messages = [ ]
			st.rerun( )

# ==============================================================================
# DOCNQ&A
# ==============================================================================
elif mode == 'Document Q&A':
	messages = st.session_state.get( 'messages', [ ] )
	uploaded = st.session_state.get( 'uploaded', [ ] )
	active_docs = st.session_state.get( 'active_docs', [ ] )
	doc_bytes = st.session_state.get( 'doc_bytes', { } )
	max_tokens = st.session_state.get( 'max_tokens', 0 )
	top_percent = st.session_state.get( 'top_percent', 0.0 )
	top_k = st.session_state.get( 'top_k', 0 )
	temperature = st.session_state.get( 'temperature', 0.0 )
	frequency_penalty = st.session_state.get( 'frequency_penalty', 0.0 )
	presense_penalty = st.session_state.get( 'presense_penalty', 0.0 )
	repeat_penalty = st.session_state.get( 'repeat_penalty', 0.0 )
	repeat_window = st.session_state.get( 'repeat_window', 0.0 )
	cpu_threads = st.session_state.get( 'cpu_threads', cfg.CORES )
	context_window = st.session_state.get( 'context_window', cfg.DEFAULT_CTX )
	
	left, center, right = st.columns( [ 0.05, 0.9, 0.05 ] )
	with center:
		st.subheader( '📚 Retrieval Augementation', help=cfg.RETRIEVAL_AUGMENTATION )
		st.divider( )
		
		# ------------------------------------------------------------------
		# Expander — Mind Controls
		# ------------------------------------------------------------------
		with st.expander( label='Mind Controls', icon='🧠', expanded=False ):
			# ------------------------------------------------------------------
			# Document Q&A reset request processing.
			#
			# Important:
			# ---------
			# Streamlit widget-owned keys must be reset before their widgets are
			# instantiated in this script pass. Buttons below only set the pending
			# request key. The actual reset is processed here at the top of this
			# expander before any controls are created.
			# ------------------------------------------------------------------
			docqna_reset_defaults: Dict[ str, Dict[ str, Any ] ] = {
					'doc_retrieval_controls_reset':
						{
								'retrieval_k': 6,
								'retrieval_chunk_size': 1200,
								'retrieval_chunk_overlap': 200,
								'show_retrieved_chunks': True,
								'require_grounding': True,
								'answer_from_excerpts_only': True,
								'prefer_sqlite_vec': True,
								'allow_similarity_fallback': True
						},
					'doc_parsing_controls_reset':
						{
								'ocr_enabled': False,
								'prefer_native_pdf_text': True,
								'include_page_markers': False,
								'show_docqna_diagnostics': False
						},
					'doc_response_controls_reset':
						{
								'top_k': 0,
								'top_percent': 0.95,
								'temperature': 0.0
						},
					'doc_probability_controls_reset':
						{
								'frequency_penalty': 0.0,
								'presense_penalty': 0.0,
								'repeat_penalty': 1.1,
								'repeat_window': 0
						},
					'doc_context_controls_reset':
						{
								'random_seed': 0,
								'max_tokens': 1024,
								'cpu_threads': int( cfg.CORES ),
								'context_window': int( cfg.DEFAULT_CTX )
						}
			}
			
			def request_docqna_reset( reset_name: str ) -> None:
				"""
				Purpose:
				--------
				Request a Document Q&A control reset without directly modifying any
				widget-owned keys after their widgets have been instantiated.
		
				Parameters:
				-----------
				reset_name : str
					Name of the reset group to process on the next safe script pass.
		
				Returns:
				--------
				None
				"""
				st.session_state[ 'pending_docqna_reset' ] = str( reset_name or '' )
			
			pending_docqna_reset = st.session_state.pop( 'pending_docqna_reset', None )
			if pending_docqna_reset:
				reset_values = docqna_reset_defaults.get( str( pending_docqna_reset ), { } )
				
				if isinstance( reset_values, dict ):
					for reset_key, reset_value in reset_values.items( ):
						st.session_state[ reset_key ] = reset_value
			
			# ------------------------------------------------------------------
			# Retrieval Controls
			# ------------------------------------------------------------------
			with st.expander( label='Retrieval Controls', icon='🧲', expanded=False ):
				ret_c1, ret_c2, ret_c3, ret_c4 = st.columns(
					[ 0.25, 0.25, 0.25, 0.25 ],
					border=True,
					gap='medium'
				)
				
				with ret_c1:
					st.slider(
						label='Chunks to Retrieve',
						min_value=1,
						max_value=20,
						step=1,
						key='retrieval_k'
					)
				
				with ret_c2:
					st.slider(
						label='Chunk Size',
						min_value=256,
						max_value=4000,
						step=64,
						key='retrieval_chunk_size'
					)
				
				with ret_c3:
					st.slider(
						label='Chunk Overlap',
						min_value=0,
						max_value=1000,
						step=25,
						key='retrieval_chunk_overlap'
					)
				
				with ret_c4:
					st.toggle(
						label='Show Retrieved Chunks',
						value=bool( st.session_state.get( 'show_retrieved_chunks', True ) ),
						key='show_retrieved_chunks'
					)
				
				ret_c5, ret_c6, ret_c7, ret_c8 = st.columns(
					[ 0.25, 0.25, 0.25, 0.25 ],
					border=True,
					gap='medium'
				)
				
				with ret_c5:
					st.toggle(
						label='Require Grounding',
						value=bool( st.session_state.get( 'require_grounding', True ) ),
						key='require_grounding'
					)
				
				with ret_c6:
					st.toggle(
						label='Answer From Excerpts Only',
						value=bool( st.session_state.get( 'answer_from_excerpts_only', True ) ),
						key='answer_from_excerpts_only'
					)
				
				with ret_c7:
					st.toggle(
						label='Use sqlite-vec',
						value=bool( st.session_state.get( 'prefer_sqlite_vec', True ) ),
						key='prefer_sqlite_vec'
					)
				
				with ret_c8:
					st.toggle(
						label='Fallback Cosine Search',
						value=bool( st.session_state.get( 'allow_similarity_fallback', True ) ),
						key='allow_similarity_fallback'
					)
				
				st.button(
					label='Reset',
					key='doc_retrieval_controls_reset',
					width='stretch',
					on_click=request_docqna_reset,
					args=('doc_retrieval_controls_reset',)
				)
			
			# ------------------------------------------------------------------
			# Document Actions
			# ------------------------------------------------------------------
			with st.expander( label='Document Actions', icon='🗂️', expanded=False ):
				action_c1, action_c2 = st.columns( [ 0.6, 0.4 ], border=True )
				
				with action_c1:
					st.selectbox(
						label='Action',
						options=[
								'Answer Question',
								'Summarize Active Document',
								'Extract Key Points',
								'Generate Outline',
								'Extract Entities',
								'Extract Tables',
								'Compare Active Documents'
						],
						key='docqna_action'
					)
				
				with action_c2:
					st.markdown( '<br>', unsafe_allow_html=True )
					
					if st.button( 'Run Action', key='doc_run_action', width='stretch' ):
						action_name = str(
							st.session_state.get( 'docqna_action', 'Answer Question' ) or
							'Answer Question'
						).strip( )
						
						action_prompts = {
								'Summarize Active Document':
									'Summarize the active document set clearly and faithfully.',
								'Extract Key Points':
									'Extract the key points from the active document set.',
								'Generate Outline':
									'Generate an outline of the active document set.',
								'Extract Entities':
									'Extract named entities, dates, organizations, and references from the active document set.',
								'Extract Tables':
									'Describe the tabular or structured information visible in the active document set.',
								'Compare Active Documents':
									'Compare the active documents and explain major agreements, differences, and gaps.'
						}
						
						if action_name != 'Answer Question':
							action_prompt = action_prompts.get(
								action_name,
								'Summarize the active document set.'
							)
							
							with st.chat_message( 'assistant' ):
								out = st.empty( )
								response = run_llm_turn(
									user_input=build_docqna_input(
										user_query=action_prompt,
										k=int( st.session_state.get( 'retrieval_k', 6 ) )
									),
									temperature=float( st.session_state.get( 'temperature', 0.0 ) ),
									top_p=float( st.session_state.get( 'top_percent', 0.95 ) ),
									repeat_penalty=float(
										st.session_state.get( 'repeat_penalty', 1.1 )
									),
									max_tokens=int(
										st.session_state.get( 'max_tokens', 1024 )
									) or 1024,
									stream=True,
									output=out
								)
							
							save_message( 'assistant', response )
							st.session_state.messages.append( ('assistant', response) )
			
			# ------------------------------------------------------------------
			# Document Parsing
			# ------------------------------------------------------------------
			with st.expander( label='Document Parsing', icon='📄', expanded=False ):
				parse_c1, parse_c2, parse_c3, parse_c4 = st.columns(
					[ 0.25, 0.25, 0.25, 0.25 ],
					border=True,
					gap='medium'
				)
				
				with parse_c1:
					st.toggle(
						label='Enable OCR',
						value=bool( st.session_state.get( 'ocr_enabled', False ) ),
						key='ocr_enabled'
					)
				
				with parse_c2:
					st.toggle(
						label='Prefer Native PDF Text',
						value=bool( st.session_state.get( 'prefer_native_pdf_text', True ) ),
						key='prefer_native_pdf_text'
					)
				
				with parse_c3:
					st.toggle(
						label='Include Page Markers',
						value=bool( st.session_state.get( 'include_page_markers', False ) ),
						key='include_page_markers'
					)
				
				with parse_c4:
					st.toggle(
						label='Show Diagnostics',
						value=bool( st.session_state.get( 'show_docqna_diagnostics', False ) ),
						key='show_docqna_diagnostics'
					)
				
				st.button(
					label='Reset',
					key='doc_parsing_controls_reset',
					width='stretch',
					on_click=request_docqna_reset,
					args=('doc_parsing_controls_reset',)
				)
			
			# ------------------------------------------------------------------
			# Response Settings
			# ------------------------------------------------------------------
			with st.expander( label='Response Settings', icon='↔️', expanded=False ):
				mind_c1, mind_c2, mind_c3 = st.columns(
					[ 0.33, 0.33, 0.33 ],
					border=True,
					gap='medium'
				)
				
				with mind_c1:
					st.slider(
						label='Temperature',
						min_value=0.0,
						max_value=1.0,
						value=float( st.session_state.get( 'temperature', 0.0 ) ),
						help=cfg.TEMPERATURE,
						key='temperature'
					)
					temperature = st.session_state[ 'temperature' ]
				
				with mind_c2:
					st.slider(
						label='Top-P',
						min_value=0.0,
						max_value=1.0,
						step=0.01,
						key='top_percent',
						help=cfg.TOP_P
					)
					top_percent = st.session_state[ 'top_percent' ]
				
				with mind_c3:
					st.slider(
						label='Top-K',
						min_value=0,
						max_value=50,
						step=1,
						key='top_k',
						help=cfg.TOP_K
					)
					top_k = st.session_state[ 'top_k' ]
				
				st.button(
					label='Reset',
					key='doc_response_controls_reset',
					width='stretch',
					on_click=request_docqna_reset,
					args=('doc_response_controls_reset',)
				)
			
			# ------------------------------------------------------------------
			# Inference Settings
			# ------------------------------------------------------------------
			with st.expander( label='Inference Settings', icon='🎚️', expanded=False ):
				prob_c1, prob_c2, prob_c3, prob_c4 = st.columns(
					[ 0.25, 0.25, 0.25, 0.25 ],
					border=True,
					gap='medium'
				)
				
				with prob_c1:
					st.slider(
						label='Repeat Window',
						min_value=0,
						max_value=1024,
						step=16,
						key='repeat_window',
						help=cfg.REPEAT_WINDOW
					)
					repeat_window = st.session_state[ 'repeat_window' ]
				
				with prob_c2:
					st.slider(
						label='Repeat Penalty',
						min_value=0.0,
						max_value=2.0,
						key='repeat_penalty',
						step=0.05,
						help=cfg.REPEAT_PENALTY
					)
					repeat_penalty = st.session_state[ 'repeat_penalty' ]
				
				with prob_c3:
					st.slider(
						label='Presence Penalty',
						min_value=0.0,
						max_value=2.0,
						key='presense_penalty',
						step=0.05,
						help=cfg.PRESENCE_PENALTY
					)
					presense_penalty = st.session_state[ 'presense_penalty' ]
				
				with prob_c4:
					st.slider(
						label='Frequency Penalty',
						min_value=0.0,
						max_value=2.0,
						key='frequency_penalty',
						step=0.05,
						help=cfg.FREQUENCY_PENALTY
					)
					frequency_penalty = st.session_state[ 'frequency_penalty' ]
				
				st.button(
					label='Reset',
					key='doc_probability_controls_reset',
					width='stretch',
					on_click=request_docqna_reset,
					args=('doc_probability_controls_reset',)
				)
			
			# ------------------------------------------------------------------
			# Context Controls
			# ------------------------------------------------------------------
			with st.expander( label='Context Controls', icon='🎛️', expanded=False ):
				ctx_c1, ctx_c2, ctx_c3, ctx_c4 = st.columns(
					[ 0.25, 0.25, 0.25, 0.25 ],
					border=True,
					gap='medium'
				)
				
				with ctx_c1:
					st.slider(
						label='Context Window',
						min_value=0,
						max_value=8192,
						key='context_window',
						step=512,
						help=cfg.CONTEXT_WINDOW
					)
					context_window = st.session_state[ 'context_window' ]
				
				with ctx_c2:
					st.slider(
						label='CPU Threads',
						min_value=0,
						max_value=cfg.CORES,
						key='cpu_threads',
						step=1,
						help=cfg.CPU_CORES
					)
					cpu_threads = st.session_state[ 'cpu_threads' ]
				
				with ctx_c3:
					st.slider(
						label='Max Tokens',
						min_value=0,
						max_value=4096,
						step=128,
						key='max_tokens',
						help=cfg.MAX_TOKENS
					)
					max_tokens = st.session_state[ 'max_tokens' ]
				
				with ctx_c4:
					st.slider(
						label='Random Seed',
						min_value=0,
						max_value=4096,
						step=1,
						key='random_seed',
						help=cfg.SEED
					)
				
				st.button(
					label='Reset',
					key='doc_context_controls_reset',
					width='stretch',
					on_click=request_docqna_reset,
					args=('doc_context_controls_reset',)
				)
	
		# ------------------------------------------------------------------
		# Expander — System Instructions
		# ------------------------------------------------------------------
		with st.expander( label='System Instructions', icon='🖥️', expanded=False, width='stretch' ):
			render_system_instructions(
				prefix='docqna',
				include_apply_preset=False,
				include_preview=False
			)
		
		# ------------------------------------------------------------------
		# Document Selection UI
		# ------------------------------------------------------------------
		with st.expander( label='Document Loader', icon='📥', expanded=False, width='stretch' ):
			# ------------------------------------------------------------------
			# Document loader pending-action processing.
			#
			# Important:
			# ---------
			# active_docs is owned by the multiselect widget below. Any reset/unload
			# of active_docs must happen before that widget is instantiated.
			# ------------------------------------------------------------------
			pending_doc_loader_action = st.session_state.pop( 'pending_doc_loader_action', None )
			
			if pending_doc_loader_action == 'unload':
				st.session_state[ 'uploaded' ] = [ ]
				st.session_state[ 'active_docs' ] = [ ]
				st.session_state[ 'doc_bytes' ] = { }
				st.session_state[ 'docqna_inventory_rows' ] = [ ]
				st.session_state[ 'docqna_fingerprint' ] = ''
				st.session_state[ 'docqna_chunk_count' ] = 0
				st.session_state[ 'docqna_fallback_rows' ] = [ ]
				st.session_state[ 'docqna_last_retrieval' ] = [ ]
				st.session_state[ 'docqna_vec_ready' ] = False
			
			def request_document_unload( ) -> None:
				"""
				Purpose:
				--------
				Request that active uploaded documents be unloaded on the next safe script pass.
		
				Parameters:
				-----------
				None
		
				Returns:
				--------
				None
				"""
				st.session_state[ 'pending_doc_loader_action' ] = 'unload'
			
			def render_pdf_preview( file_bytes: bytes, preview_name: str ) -> None:
				"""
				Purpose:
				--------
				Render a PDF preview using st.pdf when available, otherwise fall back to a
				base64 iframe and, if needed, extracted text.
		
				Parameters:
				-----------
				file_bytes : bytes
					PDF file bytes.
		
				preview_name : str
					Display name for the active PDF.
		
				Returns:
				--------
				None
				"""
				if not file_bytes:
					st.info( 'PDF preview unavailable.' )
					return
				
				try:
					if hasattr( st, 'pdf' ):
						st.pdf( file_bytes, height=420 )
						return
				except Exception:
					pass
				
				try:
					encoded_pdf = base64.b64encode( file_bytes ).decode( 'utf-8' )
					pdf_html = (
							f'<iframe src="data:application/pdf;base64,{encoded_pdf}" '
							f'width="100%" height="420" type="application/pdf"></iframe>'
					)
					st.markdown( pdf_html, unsafe_allow_html=True )
					return
				except Exception:
					pass
				
				preview_text = extract_text( file_bytes, preview_name )
				if preview_text:
					st.text_area(
						label=f'Preview: {preview_name}',
						value=preview_text[ :4000 ],
						height=420,
						disabled=True,
						key='doc_loader_pdf_text_fallback'
					)
				else:
					st.info( 'PDF preview unavailable.' )
			
			doc_left, doc_right = st.columns( [ 0.5, 0.5 ], gap='medium', border=True )
			
			with doc_left:
				st.radio( label='Document Source', options=[ 'uploadlocal' ], index=0,
					horizontal=True, key='doc_source' )
				
				uploaded = st.file_uploader( label='Upload document(s) (PDF, TXT, DOCX)',
					type=[ 'pdf', 'txt', 'docx' ], accept_multiple_files=True,
					label_visibility='visible', key='doc_loader_uploader' )
				
				if uploaded is not None and isinstance( uploaded, list ) and len( uploaded ) > 0:
					st.session_state[ 'uploaded' ] = uploaded
					
					names: List[ str ] = [ str( f.name ) for f in uploaded if getattr( f, 'name', None ) ]
					
					if 'doc_bytes' not in st.session_state or not isinstance(
							st.session_state.get( 'doc_bytes' ), dict ):
						st.session_state[ 'doc_bytes' ] = { }
					
					for f in uploaded:
						try:
							file_name = str( getattr( f, 'name', '' ) or '' ).strip( )
							if file_name:
								st.session_state[ 'doc_bytes' ][ file_name ] = f.getvalue( )
						except Exception:
							continue
					
					if 'active_docs' not in st.session_state:
						st.session_state[ 'active_docs' ] = names
					else:
						current_active_docs = st.session_state.get( 'active_docs', [ ] )
						if not isinstance( current_active_docs, list ) or len( current_active_docs ) == 0:
							st.session_state[ 'active_docs' ] = names
						else:
							st.session_state[ 'active_docs' ] = [ name for name in current_active_docs
							                                      if name in names ] or names
					
					st.session_state[ 'docqna_inventory_rows' ] = build_docqna_inventory( )
				
				else:
					st.info( 'Load a document.' )
				
				uploaded_names = [ str( f.name ) for f in st.session_state.get( 'uploaded', [ ] )
						if getattr( f, 'name', None ) ]
				
				if uploaded_names:
					active_default = st.session_state.get( 'active_docs', [ ] )
					
					if not isinstance( active_default, list ):
						active_default = [ ]
					
					active_default = [ name for name in active_default if name in uploaded_names ]
					
					st.multiselect( label='Active Documents', options=uploaded_names,
						default=active_default, key='active_docs' )
				
				st.button( label='Unload Document(s)', width='stretch',
					key='doc_loader_unload_documents', on_click=request_document_unload )
				
				if bool( st.session_state.get( 'show_docqna_diagnostics', False ) ):
					retrieval_chunk_size = int( st.session_state.get( 'retrieval_chunk_size', 1200 ) )
					retrieval_chunk_overlap = int( st.session_state.get( 'retrieval_chunk_overlap', 200 ) )
					docqna_vec_ready = bool( st.session_state.get( 'docqna_vec_ready', False ) )
					docqna_chunk_count = int( st.session_state.get( 'docqna_chunk_count', 0 ) )
					
					st.caption( f'Chunk Size: {retrieval_chunk_size} '
						f'| Chunk Overlap: {retrieval_chunk_overlap} '
						f'| Index Ready: {docqna_vec_ready} '
						f'| Chunk Count: {docqna_chunk_count}' )
			
			with doc_right:
				active_docs = st.session_state.get( 'active_docs', [ ] )
				doc_bytes = st.session_state.get( 'doc_bytes', { } )
				
				if isinstance( active_docs, list ) and len( active_docs ) > 0:
					preview_name = str( active_docs[ 0 ] )
					file_bytes = doc_bytes.get( preview_name, b'' )
					
					if file_bytes and preview_name.lower( ).endswith( '.pdf' ):
						render_pdf_preview( file_bytes, preview_name )
					
					elif file_bytes:
						preview_text = extract_text( file_bytes, preview_name )
						st.text_area( label=f'Preview: {preview_name}', value=preview_text[ :4000 ],
							height=420, disabled=True, key='doc_loader_text_preview' )
					
					else:
						st.info( 'Document loaded but preview unavailable.' )
				
				else:
					st.info( 'No document loaded.' )
			
			if st.session_state.get( 'docqna_inventory_rows' ):
				st.markdown( '### Active Document Inventory' )
				st.dataframe( pd.DataFrame( st.session_state.get( 'docqna_inventory_rows', [ ] ) ),
					use_container_width=True )
		
		# ------------------------------------------------------------------
		# Chat History Render
		# ------------------------------------------------------------------
		if 'messages' not in st.session_state or not isinstance( st.session_state.messages, list ):
			st.session_state.messages = [ ]
		
		for msg in st.session_state.messages:
			role = ''
			content = ''
			
			if isinstance( msg, dict ):
				role = str( msg.get( 'role', '' ) or '' ).strip( )
				content = msg.get( 'content', '' )
			else:
				if isinstance( msg, tuple ) or isinstance( msg, list ):
					if len( msg ) == 2:
						role = str( msg[ 0 ] or '' ).strip( )
						content = msg[ 1 ]
					else:
						role = ''
						content = ''
				else:
					role = ''
					content = ''
			
			if role not in ('user', 'assistant', 'system'):
				continue
			
			if content is None:
				content = ''
			elif not isinstance( content, str ):
				content = str( content )
			
			with st.chat_message( role ):
				st.markdown( content )
		
		st.markdown( cfg.BLUE_DIVIDER, unsafe_allow_html=True )
		
		# ------------------------------------------------------------------
		# Chat Input
		# ------------------------------------------------------------------
		user_input = st.chat_input( 'Ask a question about the document' )
		if user_input and isinstance( user_input, str ) and user_input.strip( ):
			user_input = user_input.strip( )
			
			if 'messages' not in st.session_state or not isinstance(
					st.session_state.messages, list ):
				st.session_state.messages = [ ]
			
			save_message( 'user', user_input )
			st.session_state.messages.append( ('user', user_input) )
			
			with st.chat_message( 'user' ):
				st.markdown( user_input )
			
			doc_user_input = build_docqna_input( user_query=user_input,
				k=int( st.session_state.get( 'retrieval_k', 6 ) ) )
			
			if not doc_user_input or not isinstance( doc_user_input,
					str ) or not doc_user_input.strip( ):
				doc_user_input = user_input
			
			with st.chat_message( 'assistant' ):
				out = st.empty( )
				response = run_llm_turn( user_input=doc_user_input,
					temperature=float( st.session_state.get( 'temperature', 0.0 ) ),
					top_p=float( st.session_state.get( 'top_percent', 0.95 ) ),
					repeat_penalty=float( st.session_state.get( 'repeat_penalty', 1.1 ) ),
					max_tokens=int( st.session_state.get( 'max_tokens', 1024 ) ) or 1024,
					stream=True, output=out )
			
			if response is None:
				response = ''
			elif not isinstance( response, str ):
				response = str( response )
			
			response = response.strip( )
			save_message( 'assistant', response )
			st.session_state.messages.append( ('assistant', response) )
			if bool( st.session_state.get( 'show_retrieved_chunks', True ) ):
				hits = st.session_state.get( 'docqna_last_retrieval', [ ] )
				if hits:
					with st.expander( 'Retrieved Chunks', expanded=False ):
						for idx, hit in enumerate( hits, start=1 ):
							doc_name = str( hit[ 0 ] )
							chunk_text_value = str( hit[ 1 ] )
							score_value = hit[ 2 ]
							
							st.markdown( f'**{idx}. {doc_name}**' )
							st.caption( f'Score / Distance: {score_value}' )
							st.text_area( label=f'Chunk {idx}', value=chunk_text_value,
								height=140, disabled=True, key=f'doc_hit_{idx}' )
		
		if st.button( '🧹 Clear Chat', key='doc_clear_chat' ):
			clear_history( )
			st.session_state.messages = [ ]
			st.rerun( )

# ==============================================================================
# IMAGES API
# ==============================================================================
elif mode == get_mode_constant( 'IMAGE_MODE', 'Images API' ):
	
	def run_image_mode_adapter( image_bytes: bytes, image_name: str, prompt: str ) -> str:
		"""
			Purpose:
			--------
			Run an optional image-analysis adapter when one has been wired into app.py.
			The function fails closed when the selected model or runtime does not support
			image analysis.

			Parameters:
			-----------
			image_bytes : bytes
				Uploaded image bytes.

			image_name : str
				Uploaded image filename.

			prompt : str
				User prompt for image analysis.

			Returns:
			--------
			str
				Image analysis response text.
		"""
		try:
			if not model_supports_capability( 'image_mode' ):
				return get_capability_status_message( 'image_mode' )
			
			runtime_status = get_runtime_multimodal_status( )
			if not bool( runtime_status.get( 'image_runtime_available', False ) ):
				return get_capability_status_message( 'image_mode' )
			
			adapter = globals( ).get( 'analyze_image_with_model', None )
			if not callable( adapter ):
				return (
						'Image Mode is configured for this model, but no image adapter named '
						'analyze_image_with_model is wired into app.py yet.')
			
			try:
				result = adapter(
					model_path=get_selected_model_path( ),
					model_name=get_selected_model_name( ),
					image_bytes=image_bytes,
					image_name=image_name,
					prompt=prompt
				)
			except TypeError:
				result = adapter( image_bytes, prompt )
			
			if result is None:
				return ''
			
			return str( result )
		except Exception as e:
			return f'Image analysis failed: {e}'
	
	def build_image_context_text( image_name: str, prompt: str, response: str ) -> str:
		"""
			Purpose:
			--------
			Build reusable image context text for Text Generation, Document Q&A, or Prompt
			Engineering workflows.

			Parameters:
			-----------
			image_name : str
				Uploaded image filename.

			prompt : str
				User image-analysis prompt.

			response : str
				Image-analysis response or runtime status.

			Returns:
			--------
			str
				Reusable image context text.
		"""
		name_value = str( image_name or '' ).strip( )
		prompt_value = str( prompt or '' ).strip( )
		response_value = str( response or '' ).strip( )
		
		parts: List[ str ] = [ ]
		if name_value:
			parts.append( f'Image Source: {name_value}' )
		if prompt_value:
			parts.append( f'Image Prompt: {prompt_value}' )
		if response_value:
			parts.append( f'Image Analysis:\n{response_value}' )
		
		return '\n\n'.join( parts ).strip( )
	
	left, center, right = st.columns( [ 0.05, 0.90, 0.05 ] )
	with center:
		image_help = get_mode_definition_text(
			get_mode_constant( 'IMAGE_MODE', 'Images API' ) )
		
		st.subheader( '🖼️ Images API', help=image_help )
		st.divider( )
		
		capabilities = get_active_model_capabilities( )
		runtime_status = get_runtime_multimodal_status( )
		
		status_c1, status_c2, status_c3 = st.columns( [ 0.34, 0.33, 0.33 ], border=True )
		with status_c1:
			st.metric( 'Selected Model', get_selected_model_name( ) )
		with status_c2:
			st.metric( 'Base Model', str( capabilities.get( 'base_model', '' ) or '' ) )
		with status_c3:
			runtime_label = 'Available' if bool(
				runtime_status.get( 'image_runtime_available', False ) ) else 'Text Only'
			st.metric( 'Image Runtime', runtime_label )
		
		if not model_supports_capability( 'image_mode' ):
			st.warning( get_capability_status_message( 'image_mode' ) )
		
		elif not bool( runtime_status.get( 'image_runtime_available', False ) ):
			st.info( get_capability_status_message( 'image_mode' ) )
		
		with st.expander( label='Image Input', icon='📥', expanded=True ):
			img_c1, img_c2 = st.columns( [ 0.45, 0.55 ], gap='medium' )
			
			with img_c1:
				image_file = st.file_uploader(
					label='Upload Image',
					type=[ 'png', 'jpg', 'jpeg', 'webp', 'bmp' ],
					accept_multiple_files=False,
					key='image_file_uploader'
				)
				
				if image_file is not None:
					try:
						image_bytes = image_file.getvalue( )
						image_name = str( getattr( image_file, 'name', '' ) or '' )
						st.session_state[ 'image_uploaded_name' ] = image_name
						st.image( image_bytes, caption=image_name, use_container_width=True )
						st.caption( f'Image Size: {len( image_bytes ):,} bytes' )
					except Exception as e:
						image_bytes = b''
						image_name = ''
						st.session_state[ 'image_status' ] = f'Image preview failed: {e}'
						st.error( st.session_state[ 'image_status' ] )
				else:
					image_bytes = b''
					image_name = ''
			
			with img_c2:
				st.text_area(
					label='Image Prompt',
					height=180,
					key='image_prompt',
					placeholder='Describe what you want the model to analyze in the image.'
				)
				
				st.toggle(
					label='Send Image Context to Text Generation',
					key='image_send_to_text',
					value=bool( st.session_state.get( 'image_send_to_text', False ) )
				)
				
				st.caption(
					'Image analysis requires both a model that advertises image capability '
					'and a runtime adapter that can pass image bytes to the local model.' )
		
		with st.expander( label='Image Analysis', icon='🔎', expanded=True ):
			run_disabled = image_file is None
			if st.button( 'Analyze Image', key='image_analyze_button',
					width='stretch', disabled=run_disabled ):
				prompt_value = str( st.session_state.get( 'image_prompt', '' ) or '' ).strip( )
				if not prompt_value:
					prompt_value = (
							'Analyze the uploaded image and provide a concise, structured '
							'description of the visible content.')
				
				response = run_image_mode_adapter(
					image_bytes=image_bytes,
					image_name=image_name,
					prompt=prompt_value
				)
				
				st.session_state[ 'image_response' ] = response
				st.session_state[ 'image_status' ] = 'Image analysis request completed.'
				st.session_state[ 'image_context_buffer' ] = build_image_context_text(
					image_name=image_name,
					prompt=prompt_value,
					response=response
				)
				
				if bool( st.session_state.get( 'image_send_to_text', False ) ):
					existing_docs = st.session_state.get( 'basic_docs', [ ] )
					if not isinstance( existing_docs, list ):
						existing_docs = [ ]
					
					context_text = st.session_state.get( 'image_context_buffer', '' )
					if context_text:
						existing_docs.append( context_text )
						st.session_state[ 'basic_docs' ] = existing_docs
						st.session_state[ 'use_document_context' ] = True
			
			response_value = str( st.session_state.get( 'image_response', '' ) or '' )
			if response_value:
				st.markdown( '### Image Response' )
				st.markdown( response_value )
			
			status_value = str( st.session_state.get( 'image_status', '' ) or '' )
			if status_value:
				st.caption( status_value )
		
		with st.expander( label='Actions', icon='🔀', expanded=False ):
			action_c1, action_c2, action_c3 = st.columns( [ 0.34, 0.33, 0.33 ] )
			
			with action_c1:
				if st.button( 'Send Context to Text Generation',
						key='image_send_context_button', width='stretch' ):
					context_text = str(
						st.session_state.get( 'image_context_buffer', '' ) or '' ).strip( )
					
					if context_text:
						existing_docs = st.session_state.get( 'basic_docs', [ ] )
						if not isinstance( existing_docs, list ):
							existing_docs = [ ]
						
						existing_docs.append( context_text )
						st.session_state[ 'basic_docs' ] = existing_docs
						st.session_state[ 'use_document_context' ] = True
						st.success( 'Image context added to Text Generation context.' )
					else:
						st.info( 'No image context is available to send.' )
			
			with action_c2:
				if st.button( 'Save as Prompt Context',
						key='image_save_prompt_context_button', width='stretch' ):
					context_text = str(
						st.session_state.get( 'image_context_buffer', '' ) or '' ).strip( )
					
					if context_text:
						existing_docs = st.session_state.get( 'basic_docs', [ ] )
						if not isinstance( existing_docs, list ):
							existing_docs = [ ]
						
						existing_docs.append( context_text )
						st.session_state[ 'basic_docs' ] = existing_docs
						st.success( 'Image context saved as shared prompt context.' )
					else:
						st.info( 'No image context is available to save.' )
			
			with action_c3:
				if st.button( 'Clear Image State',
						key='image_clear_state_button', width='stretch' ):
					st.session_state[ 'image_prompt' ] = ''
					st.session_state[ 'image_uploaded_name' ] = ''
					st.session_state[ 'image_response' ] = ''
					st.session_state[ 'image_status' ] = ''
					st.session_state[ 'image_context_buffer' ] = ''
					st.session_state[ 'image_send_to_text' ] = False
					st.rerun( )

# ==============================================================================
# AUDIO API
# ==============================================================================
elif mode == get_mode_constant( 'AUDIO_MODE', 'Audio API' ):
	
	def run_audio_mode_adapter( audio_bytes: bytes, audio_name: str, prompt: str ) -> str:
		"""
			Purpose:
			--------
			Run an optional audio-analysis adapter when one has been wired into app.py.
			The function fails closed when the selected model or runtime does not support
			audio transcription, translation, or audio analysis.

			Parameters:
			-----------
			audio_bytes : bytes
				Uploaded audio bytes.

			audio_name : str
				Uploaded audio filename.

			prompt : str
				User prompt for audio transcription or analysis.

			Returns:
			--------
			str
				Audio transcription, translation, analysis, or runtime status text.
		"""
		try:
			if not model_supports_capability( 'audio_mode' ):
				return get_capability_status_message( 'audio_mode' )
			
			runtime_status = get_runtime_multimodal_status( )
			if not bool( runtime_status.get( 'audio_runtime_available', False ) ):
				return get_capability_status_message( 'audio_mode' )
			
			adapter = globals( ).get( 'analyze_audio_with_model', None )
			if not callable( adapter ):
				return (
						'Audio Mode is configured for this model, but no audio adapter named '
						'analyze_audio_with_model is wired into app.py yet.')
			
			try:
				result = adapter(
					model_path=get_selected_model_path( ),
					model_name=get_selected_model_name( ),
					audio_bytes=audio_bytes,
					audio_name=audio_name,
					prompt=prompt
				)
			except TypeError:
				result = adapter( audio_bytes, prompt )
			
			if result is None:
				return ''
			
			return str( result )
		except Exception as e:
			return f'Audio analysis failed: {e}'
	
	def build_audio_context_text( audio_name: str, prompt: str, response: str ) -> str:
		"""
			Purpose:
			--------
			Build reusable audio context text for Text Generation, Document Q&A, Semantic
			Search, or Prompt Engineering workflows.

			Parameters:
			-----------
			audio_name : str
				Uploaded audio filename.

			prompt : str
				User audio-analysis prompt.

			response : str
				Audio transcription, translation, analysis, or runtime status.

			Returns:
			--------
			str
				Reusable audio context text.
		"""
		name_value = str( audio_name or '' ).strip( )
		prompt_value = str( prompt or '' ).strip( )
		response_value = str( response or '' ).strip( )
		
		parts: List[ str ] = [ ]
		if name_value:
			parts.append( f'Audio Source: {name_value}' )
		if prompt_value:
			parts.append( f'Audio Prompt: {prompt_value}' )
		if response_value:
			parts.append( f'Audio Transcript / Analysis:\n{response_value}' )
		
		return '\n\n'.join( parts ).strip( )
	
	def get_audio_mime_type( audio_name: str ) -> str:
		"""
			Purpose:
			--------
			Return a browser-friendly MIME type for Streamlit audio preview based on the
			uploaded audio filename.

			Parameters:
			-----------
			audio_name : str
				Uploaded audio filename.

			Returns:
			--------
			str
				Audio MIME type.
		"""
		name_value = str( audio_name or '' ).strip( ).lower( )
		
		if name_value.endswith( '.mp3' ):
			return 'audio/mpeg'
		if name_value.endswith( '.m4a' ):
			return 'audio/mp4'
		if name_value.endswith( '.flac' ):
			return 'audio/flac'
		if name_value.endswith( '.ogg' ):
			return 'audio/ogg'
		if name_value.endswith( '.wav' ):
			return 'audio/wav'
		
		return 'audio/wav'
	
	left, center, right = st.columns( [ 0.05, 0.90, 0.05 ] )
	with center:
		audio_help = get_mode_definition_text(
			get_mode_constant( 'AUDIO_MODE', 'Audio API' ) )
		
		st.subheader( '🎧 Audio API', help=audio_help )
		st.divider( )
		
		capabilities = get_active_model_capabilities( )
		runtime_status = get_runtime_multimodal_status( )
		
		status_c1, status_c2, status_c3 = st.columns( [ 0.34, 0.33, 0.33 ], border=True )
		with status_c1:
			st.metric( 'Selected Model', get_selected_model_name( ) )
		with status_c2:
			st.metric( 'Base Model', str( capabilities.get( 'base_model', '' ) or '' ) )
		with status_c3:
			runtime_label = 'Available' if bool(
				runtime_status.get( 'audio_runtime_available', False ) ) else 'Text Only'
			st.metric( 'Audio Runtime', runtime_label )
		
		if not model_supports_capability( 'audio_mode' ):
			st.warning( get_capability_status_message( 'audio_mode' ) )
		
		elif not bool( runtime_status.get( 'audio_runtime_available', False ) ):
			st.info( get_capability_status_message( 'audio_mode' ) )
		
		with st.expander( label='Audio Input', icon='📥', expanded=True ):
			audio_c1, audio_c2 = st.columns( [ 0.45, 0.55 ], gap='medium' )
			
			with audio_c1:
				audio_file = st.file_uploader(
					label='Upload Audio',
					type=[ 'wav', 'mp3', 'm4a', 'flac', 'ogg' ],
					accept_multiple_files=False,
					key='audio_file_uploader'
				)
				
				if audio_file is not None:
					try:
						audio_bytes = audio_file.getvalue( )
						audio_name = str( getattr( audio_file, 'name', '' ) or '' )
						audio_mime = get_audio_mime_type( audio_name )
						st.session_state[ 'audio_uploaded_name' ] = audio_name
						st.audio( audio_bytes, format=audio_mime )
						st.caption( f'Audio File: {audio_name}' )
						st.caption( f'Audio Size: {len( audio_bytes ):,} bytes' )
					except Exception as e:
						audio_bytes = b''
						audio_name = ''
						st.session_state[ 'audio_status' ] = f'Audio preview failed: {e}'
						st.error( st.session_state[ 'audio_status' ] )
				else:
					audio_bytes = b''
					audio_name = ''
			
			with audio_c2:
				st.text_area(
					label='Audio Prompt',
					height=180,
					key='audio_prompt',
					placeholder=(
							'Transcribe this audio, summarize the speaker’s key points, '
							'or translate the spoken content into English.')
				)
				
				st.toggle(
					label='Send Audio Context to Text Generation',
					key='audio_send_to_text',
					value=bool( st.session_state.get( 'audio_send_to_text', False ) )
				)
				
				st.caption(
					'Audio analysis requires both a model that advertises audio capability '
					'and a runtime adapter that can pass audio bytes to the local model.' )
		
		with st.expander( label='Audio Analysis', icon='🔎', expanded=True ):
			run_disabled = audio_file is None
			if st.button( 'Analyze Audio', key='audio_analyze_button',
					width='stretch', disabled=run_disabled ):
				prompt_value = str( st.session_state.get( 'audio_prompt', '' ) or '' ).strip( )
				if not prompt_value:
					prompt_value = (
							'Transcribe the uploaded audio and provide a concise summary '
							'of the spoken content.')
				
				response = run_audio_mode_adapter(
					audio_bytes=audio_bytes,
					audio_name=audio_name,
					prompt=prompt_value
				)
				
				st.session_state[ 'audio_response' ] = response
				st.session_state[ 'audio_transcript' ] = response
				st.session_state[ 'audio_status' ] = 'Audio analysis request completed.'
				st.session_state[ 'audio_context_buffer' ] = build_audio_context_text(
					audio_name=audio_name,
					prompt=prompt_value,
					response=response
				)
				
				if bool( st.session_state.get( 'audio_send_to_text', False ) ):
					existing_docs = st.session_state.get( 'basic_docs', [ ] )
					if not isinstance( existing_docs, list ):
						existing_docs = [ ]
					
					context_text = st.session_state.get( 'audio_context_buffer', '' )
					if context_text:
						existing_docs.append( context_text )
						st.session_state[ 'basic_docs' ] = existing_docs
						st.session_state[ 'use_document_context' ] = True
			
			response_value = str( st.session_state.get( 'audio_response', '' ) or '' )
			if response_value:
				st.markdown( '### Audio Response' )
				st.markdown( response_value )
			
			status_value = str( st.session_state.get( 'audio_status', '' ) or '' )
			if status_value:
				st.caption( status_value )
		
		with st.expander( label='Actions', icon='🔀', expanded=False ):
			action_c1, action_c2, action_c3 = st.columns( [ 0.34, 0.33, 0.33 ] )
			
			with action_c1:
				if st.button( 'Send Context to Text Generation',
						key='audio_send_context_button', width='stretch' ):
					context_text = str(
						st.session_state.get( 'audio_context_buffer', '' ) or '' ).strip( )
					
					if context_text:
						existing_docs = st.session_state.get( 'basic_docs', [ ] )
						if not isinstance( existing_docs, list ):
							existing_docs = [ ]
						
						existing_docs.append( context_text )
						st.session_state[ 'basic_docs' ] = existing_docs
						st.session_state[ 'use_document_context' ] = True
						st.success( 'Audio context added to Text Generation context.' )
					else:
						st.info( 'No audio context is available to send.' )
			
			with action_c2:
				if st.button( 'Save as Prompt Context',
						key='audio_save_prompt_context_button', width='stretch' ):
					context_text = str(
						st.session_state.get( 'audio_context_buffer', '' ) or '' ).strip( )
					
					if context_text:
						existing_docs = st.session_state.get( 'basic_docs', [ ] )
						if not isinstance( existing_docs, list ):
							existing_docs = [ ]
						
						existing_docs.append( context_text )
						st.session_state[ 'basic_docs' ] = existing_docs
						st.success( 'Audio context saved as shared prompt context.' )
					else:
						st.info( 'No audio context is available to save.' )
			
			with action_c3:
				if st.button( 'Clear Audio State',
						key='audio_clear_state_button', width='stretch' ):
					st.session_state[ 'audio_prompt' ] = ''
					st.session_state[ 'audio_uploaded_name' ] = ''
					st.session_state[ 'audio_response' ] = ''
					st.session_state[ 'audio_status' ] = ''
					st.session_state[ 'audio_transcript' ] = ''
					st.session_state[ 'audio_context_buffer' ] = ''
					st.session_state[ 'audio_send_to_text' ] = False
					st.rerun( )
					
# ==============================================================================
# SEMANTIC SEARCH
# ==============================================================================
elif mode == 'Semantic Search':
	left, center, right = st.columns( [ 0.05, 0.9, 0.05 ] )
	with center:
		st.subheader( '🔍 Semantic Search', help=cfg.SEMANTIC_SEARCH )
		st.divider( )
		
		with st.expander( label='Index Builder', icon='🧱', expanded=False ):
			idx_c1, idx_c2, idx_c3, idx_c4 = st.columns( [ 0.25, 0.25, 0.25, 0.25 ],
				border=True, gap='medium' )
			
			with idx_c1:
				st.slider( label='Chunk Size', min_value=256, max_value=4000, step=64,
					key='semantic_chunk_size' )
			
			with idx_c2:
				st.slider( label='Chunk Overlap', min_value=0, max_value=1000, step=25,
					key='semantic_chunk_overlap' )
			
			with idx_c3:
				st.toggle( label='Clear Existing Index',
					value=bool( st.session_state.get( 'semantic_clear_existing', True ) ),
					key='semantic_clear_existing' )
			
			with idx_c4:
				st.toggle( label='Append to Existing Index',
					value=bool( st.session_state.get( 'semantic_append_existing', False ) ),
					key='semantic_append_existing' )
			
			st.toggle( label='Show Embedding Diagnostics',
				value=bool( st.session_state.get( 'semantic_show_diagnostics', True ) ),
				key='semantic_show_diagnostics' )
			
			semantic_files = st.file_uploader( label='Upload for embedding',
				accept_multiple_files=True, type=[ 'pdf', 'txt', 'docx' ],
				key='semantic_file_uploader' )
			
			if st.button( 'Build Index', key='semantic_build_index', width='stretch' ):
				if semantic_files:
					result = build_semantic_index( semantic_files )
					if bool( result.get( 'success', False ) ):
						st.success( str( result.get( 'message', '' ) ) )
					else:
						st.error( str( result.get( 'message', 'Index build failed.' ) ) )
				else:
					st.info( 'Upload one or more files before building the index.' )
			
			if bool( st.session_state.get( 'semantic_show_diagnostics', True ) ):
				diag_c1, diag_c2, diag_c3 = st.columns( [ 0.33, 0.33, 0.34 ], border=True )
				with diag_c1:
					st.metric( 'Indexed Documents',
						int( st.session_state.get( 'semantic_index_doc_count', 0 ) ) )
					
				with diag_c2:
					st.metric( 'Indexed Chunks',
						int( st.session_state.get( 'semantic_index_chunk_count', 0 ) ) )
				with diag_c3:
					st.metric( 'Vector Dimension',
						int( st.session_state.get( 'semantic_index_dim', 0 ) ) )
		
		with st.expander( label='Semantic Query', icon='🧠', expanded=False ):
			query_c1, query_c2, query_c3 = st.columns( [ 0.34, 0.33, 0.33 ], border=True,
				gap='medium' )
			
			with query_c1:
				st.slider( label='Top K', min_value=1, max_value=25, step=1,
					key='semantic_top_k' )
			
			with query_c2:
				st.slider( label='Minimum Similarity', min_value=0.0, max_value=1.0, step=0.01,
					key='semantic_min_similarity' )
			
			with query_c3:
				st.toggle( label='Group by Document',
					value=bool( st.session_state.get( 'semantic_group_by_document', False ) ),
					key='semantic_group_by_document' )
			
			semantic_query = st.text_area( label='Semantic Query', height=120,
				key='semantic_query_text' )
			
			if st.button( 'Run Semantic Search', key='semantic_run_query', width='stretch' ):
				rows = query_semantic_index( semantic_query )
				if len( rows ) == 0:
					st.info( 'No semantic matches found.' )
			
			result_rows = st.session_state.get( 'semantic_result_rows', [ ] )
			if isinstance( result_rows, list ) and len( result_rows ) > 0:
				edited_rows = st.data_editor( result_rows, hide_index=True, use_container_width=True,
					key='semantic_results_editor' )
				
				selected_rows = extract_selected_rows( edited_rows )
				st.session_state[ 'semantic_selected_rows' ] = selected_rows
				if len( selected_rows ) > 0:
					st.caption( f'Selected Chunks: {len( selected_rows )}' )
		
		with st.expander( label='Actions', icon='🔀', expanded=False ):
			act_c1, act_c2, act_c3 = st.columns( [ 0.34, 0.33, 0.33 ] )
			
			with act_c1:
				if st.button( 'Send Selected Chunks to Text Generation', width='stretch' ):
					send_text_chunks( )
					st.success( 'Selected chunks added to shared Text Generation context.' )
			
			with act_c2:
				if st.button( 'Send Selected Chunks to Document Q&A', width='stretch' ):
					send_docqna_chunks( )
					st.success( 'Selected chunks added to the shared Document Q&A context buffer.' )
			
			with act_c3:
				if st.button( 'Save Selected Chunks as Prompt Context', width='stretch' ):
					context_text = create_semantic_context( )
					if context_text:
						existing_docs = st.session_state.get( 'basic_docs', [ ] )
						if not isinstance( existing_docs, list ):
							existing_docs = [ ]
						existing_docs.append( context_text )
						st.session_state[ 'basic_docs' ] = existing_docs
						st.success( 'Selected chunks saved to shared prompt context.' )
					else:
						st.info( 'Select one or more chunks first.' )
			
			selected_rows = st.session_state.get( 'semantic_selected_rows', [ ] )
			if isinstance( selected_rows, list ) and len( selected_rows ) > 0:
				st.markdown( '### Selected Semantic Context Preview' )
				st.text_area( label='Selected Context',
					value=create_semantic_context( ),
					height=220, disabled=True )
		
		with st.expander( label='Index Maintenance', icon='🛠️', expanded=False ):
			maint_c1, maint_c2, maint_c3 = st.columns( [ 0.34, 0.33, 0.33 ] )
			
			with maint_c1:
				if st.button( 'Delete Index', width='stretch' ):
					clear_semantic_index( )
					st.success( 'Semantic index deleted.' )
			
			with maint_c2:
				if st.button( 'Recompute Diagnostics', width='stretch' ):
					rows = decode_embedding_rows( )
					st.session_state[ 'semantic_index_chunk_count' ] = len( rows )
					if len( rows ) > 0:
						st.session_state[ 'semantic_index_dim' ] = int( rows[ 0 ][ 1 ].shape[ 0 ] )
					else:
						st.session_state[ 'semantic_index_dim' ] = 0
					st.success( 'Diagnostics refreshed.' )
			
			with maint_c3:
				if st.button( 'Clear Query Results', width='stretch' ):
					st.session_state[ 'semantic_result_rows' ] = [ ]
					st.session_state[ 'semantic_selected_rows' ] = [ ]
					st.session_state[ 'semantic_last_query' ] = ''
					st.success( 'Query results cleared.' )
			
			if bool( st.session_state.get( 'semantic_show_diagnostics', True ) ):
				st.caption( f'Last Query: {str( st.session_state.get( "semantic_last_query", "" ) )} '
					f'| Uploaded Sources: {len( st.session_state.get( "semantic_uploaded_names", [ ] ) )}' )

# ==============================================================================
# PROMPT ENGINEERING MODE
# ==============================================================================
elif mode == 'Prompt Engineering':
	import sqlite3
	import math
	
	TABLE = 'Prompts'
	PAGE_SIZE = 10
	st.session_state.setdefault( 'pe_cascade_enabled', False )
	left, center, right = st.columns( [ 0.05, 0.90, 0.05 ] )
	with center:
		st.subheader( '📝 Prompt Engineering', help=cfg.PROMPT_ENGINEERING )
		st.divider( )
		
		st.checkbox( 'Cascade selection into shared System Instructions and task settings',
			key='pe_cascade_enabled' )
		
		# ------------------------------------------------------------------
		# Session state
		# ------------------------------------------------------------------
		st.session_state.setdefault( 'pe_page', 1 )
		st.session_state.setdefault( 'pe_search', '' )
		st.session_state.setdefault( 'pe_sort_col', 'PromptsId' )
		st.session_state.setdefault( 'pe_sort_dir', 'ASC' )
		st.session_state.setdefault( 'pe_selected_id', None )
		st.session_state.setdefault( 'pe_caption', '' )
		st.session_state.setdefault( 'pe_name', '' )
		st.session_state.setdefault( 'pe_text', '' )
		st.session_state.setdefault( 'pe_version', '' )
		st.session_state.setdefault( 'pe_id', 0 )
		
		# ------------------------------------------------------------------
		# DB helpers
		# ------------------------------------------------------------------
		def get_conn( ):
			return sqlite3.connect( cfg.DB_PATH )
		
		def reset_selection( ):
			st.session_state.pe_selected_id = None
			st.session_state.pe_caption = ''
			st.session_state.pe_name = ''
			st.session_state.pe_text = ''
			st.session_state.pe_version = ''
			st.session_state.pe_id = 0
		
		def load_prompt( pid: int ) -> None:
			with get_conn( ) as conn:
				_select = f'''
					SELECT PromptsId, Caption, Name, Text, Version, ID
					FROM {TABLE}
					WHERE PromptsId=?
				'''
				cur = conn.execute( _select, (pid,) )
				row = cur.fetchone( )
				if not row:
					return
				
				st.session_state.pe_selected_id = row[ 0 ]
				st.session_state.pe_caption = row[ 1 ]
				st.session_state.pe_name = row[ 2 ]
				st.session_state.pe_text = row[ 3 ]
				st.session_state.pe_version = row[ 4 ]
				st.session_state.pe_id = row[ 5 ]
				
				prompt_row = {
						'PromptsId': row[ 0 ],
						'Caption': row[ 1 ],
						'Name': row[ 2 ],
						'Text': row[ 3 ],
						'Version': row[ 4 ],
						'ID': row[ 5 ]
				}
				
				st.session_state[ 'prompt_category' ] = infer_prompt_category( prompt_row )
		
		# ------------------------------------------------------------------
		# Filters
		# ------------------------------------------------------------------
		c1, c2, c3, c4, c5 = st.columns( [ 3, 2, 2, 2, 3 ], border=True )
		
		with c1:
			st.text_input( 'Search (Caption / Name / Text)', key='pe_search' )
		
		with c2:
			st.selectbox( 'Category', get_prompt_categories( ), key='prompt_category_selection' )
		
		with c3:
			st.selectbox( 'Sort by',
				[ 'PromptsId', 'Caption', 'Name', 'Text', 'Version', 'ID' ], key='pe_sort_col' )
		
		with c4:
			st.selectbox( 'Direction', [ 'ASC', 'DESC' ], key='pe_sort_dir' )
		
		with c5:
			st.markdown(
				"<div style='font-size:0.95rem;font-weight:600;margin-bottom:0.25rem;'>Go to ID</div>",
				unsafe_allow_html=True )
			
			a1, a2, a3 = st.columns( [ 2, 1, 1 ] )
			with a1:
				jump_id = st.number_input( 'Go to ID', min_value=1, step=1,
					label_visibility='collapsed' )
			
			with a2:
				if st.button( 'Go' ):
					st.session_state.pe_selected_id = int( jump_id )
					load_prompt( int( jump_id ) )
			
			with a3:
				st.button( 'Clear', on_click=reset_selection )
		
		# ------------------------------------------------------------------
		# Load prompt table
		# ------------------------------------------------------------------
		where_clauses: List[ str ] = [ ]
		params: List[ Any ] = [ ]
		
		if st.session_state.pe_search:
			where_clauses.append( '(Caption LIKE ? OR Name LIKE ? OR Text LIKE ?)' )
			s = f"%{st.session_state.pe_search}%"
			params.extend( [ s, s, s ] )
		
		where = ''
		if len( where_clauses ) > 0:
			where = 'WHERE ' + ' AND '.join( where_clauses )
		
		offset = (st.session_state.pe_page - 1) * PAGE_SIZE
		
		query = f"""
	        SELECT PromptsId, Caption, Name, Text, Version, ID
	        FROM {TABLE}
	        {where}
	        ORDER BY {st.session_state.pe_sort_col} {st.session_state.pe_sort_dir}
	        LIMIT {PAGE_SIZE} OFFSET {offset}
	    """
		
		count_query = f"SELECT COUNT(*) FROM {TABLE} {where}"
		
		with get_conn( ) as conn:
			rows = conn.execute( query, params ).fetchall( )
			total_rows = conn.execute( count_query, params ).fetchone( )[ 0 ]
		
		total_pages = max( 1, math.ceil( total_rows / PAGE_SIZE ) )
		
		# ------------------------------------------------------------------
		# Prompt table
		# ------------------------------------------------------------------
		table_rows: List[ Dict[ str, Any ] ] = [ ]
		selected_category = str( st.session_state.get( 'prompt_category_table', 'General Chat' ) or 'General Chat' )
		
		for r in rows:
			prompt_row = {
					'PromptsId': r[ 0 ],
					'Caption': r[ 1 ],
					'Name': r[ 2 ],
					'Text': r[ 3 ],
					'Version': r[ 4 ],
					'ID': r[ 5 ]
			}
			
			inferred_category = infer_prompt_category( prompt_row )
			if selected_category and inferred_category != selected_category:
				continue
			
			table_rows.append( {
						'Selected': r[ 0 ] == st.session_state.pe_selected_id,
						'PromptsId': r[ 0 ],
						'Category': inferred_category,
						'Caption': r[ 1 ],
						'Name': r[ 2 ],
						'Text': r[ 3 ],
						'Version': r[ 4 ],
						'ID': r[ 5 ]
				} )
		
		edited = st.data_editor( table_rows, hide_index=True, use_container_width=True,
			key='prompt_table' )
		
		# ------------------------------------------------------------------
		# Selection processing
		# ------------------------------------------------------------------
		selected = [ r for r in edited if isinstance( r, dict ) and r.get( 'Selected' ) ]
		if len( selected ) == 1:
			pid = int( selected[ 0 ][ 'PromptsId' ] )
			if pid != st.session_state.pe_selected_id:
				load_prompt( pid )
				if bool( st.session_state.get( 'pe_cascade_enabled', False ) ):
					apply_prompt_to_text_generation( st.session_state.pe_text )
					apply_prompt_metadata_to_shared_state(
						category=selected[ 0 ].get( 'Category', 'General Chat' ),
						task_type=st.session_state.get( 'prompt_task', 'Chat' ),
						response_format=st.session_state.get( 'prompt_response_format', 'Markdown' ),
						language=st.session_state.get( 'pe_language', 'English' ) )
		
		elif len( selected ) == 0:
			pass
		
		elif len( selected ) > 1:
			st.warning( 'Select exactly one prompt row.' )
		
		# ------------------------------------------------------------------
		# Paging
		# ------------------------------------------------------------------
		p1, p2, p3 = st.columns( [ 0.25, 3.5, 0.25 ] )
		with p1:
			if st.button( '◀ Prev' ) and st.session_state.pe_page > 1:
				st.session_state.pe_page -= 1
		
		with p2:
			st.markdown( f'Page **{st.session_state.pe_page}** of **{total_pages}**' )
		
		with p3:
			if st.button( 'Next ▶' ) and st.session_state.pe_page < total_pages:
				st.session_state.pe_page += 1
		
		st.markdown( cfg.BLUE_DIVIDER, unsafe_allow_html=True )
		
		# ------------------------------------------------------------------
		# Prompt actions
		# ------------------------------------------------------------------
		with st.expander( '⚙️ Prompt Actions', expanded=False ):
			act_c1, act_c2, act_c3, act_c4 = st.columns( [ 0.25, 0.25, 0.25, 0.25 ] )
			
			with act_c1:
				if st.button( 'Apply to Text Generation', width='stretch' ):
					apply_prompt_to_text_generation( st.session_state.get( 'pe_text', '' ) )
					apply_prompt_metadata_to_shared_state(
						category=st.session_state.get( 'prompt_category_apply', 'General Chat' ),
						task_type=st.session_state.get( 'prompt_task', 'Chat' ),
						response_format=st.session_state.get( 'prompt_response_format', 'Markdown' ),
						language=st.session_state.get( 'pe_language', 'English' ) )
					st.success( 'Applied to shared Text Generation settings.' )
			
			with act_c2:
				if st.button( 'Apply to Document Q&A', width='stretch' ):
					apply_prompt_to_document_qna( st.session_state.get( 'pe_text', '' ) )
					apply_prompt_metadata_to_shared_state(
						category=st.session_state.get( 'prompt_category_meta', 'General Chat' ),
						task_type=st.session_state.get( 'prompt_task', 'Chat' ),
						response_format=st.session_state.get( 'prompt_response_format', 'Markdown' ),
						language=st.session_state.get( 'pe_language', 'English' ) )
					st.success( 'Applied to shared Document Q&A settings.' )
			
			with act_c3:
				if st.button( 'Clone as New Template', width='stretch' ):
					source_prompt = {
							'PromptsId': st.session_state.get( 'pe_selected_id' ),
							'Caption': st.session_state.get( 'pe_caption', '' ),
							'Name': st.session_state.get( 'pe_name', '' ),
							'Text': st.session_state.get( 'pe_text', '' ),
							'Version': st.session_state.get( 'pe_version', '' ),
							'ID': st.session_state.get( 'pe_id', 0 )
					}
					clone_prompt_record( source_prompt )
					st.success( 'Prompt cloned into a new editable draft.' )
			
			with act_c4:
				if st.button( 'Generate Starter Prompt', width='stretch' ):
					st.session_state.pe_text = build_starter_prompt_template(
						category=st.session_state.get( 'prompt_category', 'General Chat' ),
						task_type=st.session_state.get( 'prompt_task', 'Chat' ),
						response_format=st.session_state.get( 'prompt_response_format', 'Markdown' ),
						language=st.session_state.get( 'pe_language', 'English' ) )
					st.success( 'Starter prompt generated into the edit surface.' )
		
		# ------------------------------------------------------------------
		# Prompt generator
		# ------------------------------------------------------------------
		with st.expander( '🧪 Prompt Generator', expanded=False ):
			gen_c1, gen_c2, gen_c3, gen_c4 = st.columns( [ 0.25, 0.25, 0.25, 0.25 ], border=True )
			
			with gen_c1:
				st.selectbox( 'Task Type', get_prompt_task_types( ), key='prompt_task_generator' )
			
			with gen_c2:
				st.selectbox( 'Response Format',
					[ 'Plain Text', 'Markdown', 'Bullet Summary', 'JSON' ],
					key='prompt_format' )
			
			with gen_c3:
				st.text_input( 'Language', key='pe_language' )
			
			with gen_c4:
				st.selectbox( 'Generator Style', [ 'Practical', 'Formal', 'Analytical', 'Concise' ],
					key='pe_generator_style' )
			
			st.text_input( 'Goal', key='pe_generator_goal' )
			
			st.text_area( 'Constraints', height=120, key='pe_generator_constraints' )
			
			if st.button( 'Generate Template Draft', width='stretch' ):
				draft = generate_prompt_template_draft(
					goal=st.session_state.get( 'pe_generator_goal', '' ),
					constraints=st.session_state.get( 'pe_generator_constraints', '' ),
					style=st.session_state.get( 'pe_generator_style', 'Practical' ),
					category=st.session_state.get( 'prompt_category_draft', 'General Chat' ),
					task_type=st.session_state.get( 'prompt_task', 'Chat' ),
					response_format=st.session_state.get( 'prompt_response_format', 'Markdown' ),
					language=st.session_state.get( 'pe_language', 'English' ) )
				st.session_state[ 'pe_generated_template' ] = draft
				st.session_state.pe_text = draft
			
			if st.session_state.get( 'pe_generated_template', '' ):
				st.text_area( 'Generated Draft',
					value=st.session_state.get( 'pe_generated_template', '' ),
					height=180, disabled=True )
		
		# ------------------------------------------------------------------
		# Edit Prompt
		# ------------------------------------------------------------------
		with st.expander( '🖊️ Edit Prompt', expanded=False ):
			meta_c1, meta_c2, meta_c3, meta_c4 = st.columns( [ 0.25, 0.25, 0.25, 0.25 ] )
			
			with meta_c1:
				st.text_input( 'PromptsId', value=st.session_state.pe_selected_id or '',
					disabled=True )
			
			with meta_c2:
				st.selectbox( 'Category', get_prompt_categories( ),
					key='prompt_category_edit' )
			
			with meta_c3:
				st.selectbox( 'Task Type', get_prompt_task_types( ),
					key='pe_task_type_edit' )
			
			with meta_c4:
				st.selectbox( 'Response Format', [ 'Plain Text', 'Markdown', 'Bullet Summary', 'JSON' ],
					key='prompt_response_format' )
			
			st.text_input( 'Caption', key='pe_caption' )
			st.text_input( 'Name', key='pe_name' )
			st.text_input( 'Language', key='pe_language_edit' )
			st.text_area( 'Text', key='pe_text', height=260 )
			st.text_input( 'Version', key='pe_version' )
			
			c1, c2, c3 = st.columns( 3 )
			with c1:
				save_label = '💾 Save Changes' if st.session_state.pe_selected_id else '➕ Create Prompt'
				if st.button( save_label ):
					with get_conn( ) as conn:
						if st.session_state.pe_selected_id:
							conn.execute( f"""
	                            UPDATE {TABLE}
	                            SET Caption=?, Name=?, Text=?, Version=?, ID=?
	                            WHERE PromptsId=?
	                            """,
								( st.session_state.pe_caption, st.session_state.pe_name,
								  st.session_state.pe_text, st.session_state.pe_version,
								  st.session_state.pe_id, st.session_state.pe_selected_id ) )
						else:
							conn.execute( f"""
	                            INSERT INTO {TABLE} (Caption, Name, Text, Version, ID)
	                            VALUES (?, ?, ?, ?, ?)
	                            """, ( st.session_state.pe_caption,
										st.session_state.pe_name,
										st.session_state.pe_text,
										st.session_state.pe_version,
										st.session_state.pe_id ) )
						conn.commit( )
					
					st.success( 'Saved.' )
			
			with c2:
				if st.session_state.pe_selected_id and st.button( 'Delete' ):
					with get_conn( ) as conn:
						conn.execute(
							f'DELETE FROM {TABLE} WHERE PromptsId=?',
							(st.session_state.pe_selected_id,)
						)
						conn.commit( )
					
					reset_selection( )
					st.success( 'Deleted.' )
			
			with c3:
				st.button( '🧹 Clear Selection', on_click=reset_selection )

# ==============================================================================
# DATA MANAGEMENT MODE
# ==============================================================================
elif mode == 'Data Management':
	left, center, right = st.columns( [ 0.05, 0.90, 0.05 ] )
	with center:
		st.subheader( '🏛️ Data Management', help=cfg.DATA_MANAGEMENT )
		st.divider( )
		
		tabs = st.tabs( [ '📥 Import', '🗂 Browse', '💉 CRUD', '📊 Explore', '🔎 Filter',
			                  '🧮 Aggregate', '📈 Visualize', '⚙ Admin', '🧠 SQL' ] )
		
		tables = list_tables( )
		if not tables:
			st.info( 'No tables available.' )
		else:
			table = st.selectbox( 'Table', tables )
			df_full = read_table( table )
		
		# ----------------------------------------------------------------------
		# IMPORT TAB
		# ----------------------------------------------------------------------
		with tabs[ 0 ]:
			st.subheader( 'Structured Data Import' )
			uploaded_file = st.file_uploader( 'Upload Excel File', type=[ 'xlsx' ] )
			overwrite = st.checkbox( 'Overwrite existing tables', value=True )
			
			if uploaded_file:
				try:
					sheets = pd.read_excel( uploaded_file, sheet_name=None )
					with create_connection( ) as conn:
						conn.execute( 'BEGIN' )
						for sheet_name, df in sheets.items( ):
							table_name = create_identifier( sheet_name )
							if overwrite:
								conn.execute( f'DROP TABLE IF EXISTS "{table_name}"' )
							
							columns = [ ]
							df.columns = [ create_identifier( c ) for c in df.columns ]
							for col in df.columns:
								sql_type = get_sqlite_type( df[ col ].dtype )
								columns.append( f'"{col}" {sql_type}' )
							
							create_stmt = ( f'CREATE TABLE "{table_name}" '
									f'({", ".join( columns )});' )
							
							conn.execute( create_stmt )
							
							placeholders = ", ".join( [ "?" ] * len( df.columns ) )
							insert_stmt = ( f'INSERT INTO "{table_name}" '
									f'VALUES ({placeholders});' )
							
							conn.executemany( insert_stmt,
								df.where( pd.notnull( df ), None ).values.tolist( ) )
						
						conn.commit( )
					
					st.success( 'Import completed successfully (transaction committed).' )
					st.rerun( )
				
				except Exception as e:
					try:
						conn.rollback( )
					except Exception:
						pass
					st.error( f'Import failed — transaction rolled back.\n\n{e}' )
			
			st.markdown( cfg.BLUE_DIVIDER, unsafe_allow_html=True )
			st.subheader( 'AI Asset Registration' )
			asset_c1, asset_c2 = st.columns( [ 0.5, 0.5 ], border=True )
			with asset_c1:
				if st.button( 'Register Active Documents', width='stretch' ):
					doc_result = register_session_documents( )
					chunk_result = register_session_chunks( )
					embed_result = register_session_embeddings( )
					
					st.session_state[ 'dm_asset_sync_status' ] = (
							f'Documents inserted: {doc_result[ "inserted" ]}, '
							f'updated: {doc_result[ "updated" ]}, '
							f'chunks inserted: {chunk_result[ "inserted" ]}, '
							f'embeddings inserted: {embed_result[ "inserted" ]}'
					)
					st.success( st.session_state[ 'dm_asset_sync_status' ] )
			
			with asset_c2:
				image_uploads = st.file_uploader( 'Upload images for metadata registration',
					type=[ 'png', 'jpg', 'jpeg', 'webp' ], accept_multiple_files=True,
					key='dm_image_uploads' )
				
				if st.button( 'Register Uploaded Images', width='stretch' ):
					if image_uploads:
						image_result = register_upload_images( image_uploads )
						st.session_state[ 'dm_asset_sync_status' ] = (
								f'Images inserted: {image_result[ "inserted" ]}, '
								f'updated: {image_result[ "updated" ]}' )
						st.success( st.session_state[ 'dm_asset_sync_status' ] )
					else:
						st.info( 'Upload one or more images first.' )
			
			if st.session_state.get( 'dm_asset_sync_status', '' ):
				st.caption( st.session_state.get( 'dm_asset_sync_status', '' ) )
		
		# ----------------------------------------------------------------------
		# BROWSE TAB
		# ----------------------------------------------------------------------
		with tabs[ 1 ]:
			tables = list_tables( )
			if tables:
				table = st.selectbox( 'Table', tables, key='table_name' )
				df = read_table( table )
				st.dataframe( df, use_container_width=True )
			else:
				st.info( 'No tables available.' )
		
		# ----------------------------------------------------------------------
		# CRUD TAB
		# ----------------------------------------------------------------------
		with tabs[ 2 ]:
			tables = list_tables( )
			if not tables:
				st.info( 'No tables available.' )
			else:
				table = st.selectbox( 'Select Table', tables, key='crud_table' )
				df = read_table( table )
				schema = create_schema( table )
				
				type_map = { col[ 1 ]: col[ 2 ].upper( ) for col in schema if col[ 1 ] != 'rowid' }
				
				st.subheader( 'Insert Row' )
				insert_data = { }
				for column, col_type in type_map.items( ):
					if 'INT' in col_type:
						insert_data[ column ] = st.number_input( column, step=1, key=f'ins_{column}' )
					elif 'REAL' in col_type:
						insert_data[ column ] = st.number_input( column, format='%.6f', key=f'ins_{column}' )
					elif 'BOOL' in col_type:
						insert_data[ column ] = 1 if st.checkbox( column, key=f'ins_{column}' ) else 0
					else:
						insert_data[ column ] = st.text_input( column, key=f'ins_{column}' )
				
				if st.button( 'Insert Row' ):
					cols = list( insert_data.keys( ) )
					placeholders = ', '.join( [ '?' ] * len( cols ) )
					stmt = f'INSERT INTO "{table}" ({", ".join( cols )}) VALUES ({placeholders});'
					
					with create_connection( ) as conn:
						conn.execute( stmt, list( insert_data.values( ) ) )
						conn.commit( )
					
					st.success( 'Row inserted.' )
					st.rerun( )
				
				st.subheader( 'Update Row' )
				rowid = st.number_input( 'Row ID', min_value=1, step=1 )
				update_data = { }
				for column, col_type in type_map.items( ):
					if 'INT' in col_type:
						val = st.number_input( column, step=1, key=f'upd_{column}' )
						update_data[ column ] = val
					elif 'REAL' in col_type:
						val = st.number_input( column, format='%.6f', key=f'upd_{column}' )
						update_data[ column ] = val
					elif 'BOOL' in col_type:
						val = 1 if st.checkbox( column, key=f'upd_{column}' ) else 0
						update_data[ column ] = val
					else:
						val = st.text_input( column, key=f'upd_{column}' )
						update_data[ column ] = val
				
				if st.button( 'Update Row' ):
					set_clause = ', '.join( [ f'{c}=?' for c in update_data ] )
					stmt = f'UPDATE {table} SET {set_clause} WHERE rowid=?;'
					
					with create_connection( ) as conn:
						conn.execute( stmt, list( update_data.values( ) ) + [ rowid ] )
						conn.commit( )
					
					st.success( 'Row updated.' )
					st.rerun( )
				
				st.subheader( 'Delete Row' )
				delete_id = st.number_input( 'Row ID to Delete', min_value=1, step=1 )
				if st.button( 'Delete Row' ):
					with create_connection( ) as conn:
						conn.execute( f'DELETE FROM {table} WHERE rowid=?;', (delete_id,) )
						conn.commit( )
					
					st.success( 'Row deleted.' )
					st.rerun( )
		
		# ----------------------------------------------------------------------
		# EXPLORE TAB
		# ----------------------------------------------------------------------
		with tabs[ 3 ]:
			tables = list_tables( )
			if tables:
				table = st.selectbox( 'Table', tables, key='explore_table' )
				page_size = st.slider( 'Rows per page', 10, 500, 50 )
				page = st.number_input( 'Page', min_value=1, step=1 )
				offset = (page - 1) * page_size
				df_page = read_table( table, page_size, offset )
				st.dataframe( df_page, use_container_width=True )
		
		# ----------------------------------------------------------------------
		# FILTER TAB
		# ----------------------------------------------------------------------
		with tabs[ 4 ]:
			tables = list_tables( )
			if tables:
				table = st.selectbox( 'Table', tables, key='filter_table' )
				df = read_table( table )
				column = st.selectbox( 'Column', df.columns )
				value = st.text_input( 'Contains' )
				if value:
					df = df[ df[ column ].astype( str ).str.contains( value ) ]
				st.dataframe( df, use_container_width=True )
		
		# ----------------------------------------------------------------------
		# AGGREGATE TAB
		# ----------------------------------------------------------------------
		with tabs[ 5 ]:
			tables = list_tables( )
			if tables:
				table = st.selectbox( 'Table', tables, key='agg_table' )
				df = read_table( table )
				numeric_cols = df.select_dtypes( include=[ 'number' ] ).columns.tolist( )
				if numeric_cols:
					col = st.selectbox( 'Column', numeric_cols )
					agg = st.selectbox( 'Function', [ 'SUM', 'AVG', 'COUNT' ] )
					if agg == 'SUM':
						st.metric( 'Result', df[ col ].sum( ) )
					elif agg == 'AVG':
						st.metric( 'Result', df[ col ].mean( ) )
					elif agg == 'COUNT':
						st.metric( 'Result', df[ col ].count( ) )
		
		# ----------------------------------------------------------------------
		# VISUALIZE TAB
		# ----------------------------------------------------------------------
		with tabs[ 6 ]:
			tables = list_tables( )
			if tables:
				table = st.selectbox( 'Table', tables, key='viz_table' )
				df = read_table( table )
				numeric_cols = df.select_dtypes( include=[ 'number' ] ).columns.tolist( )
				if numeric_cols:
					col = st.selectbox( 'Column', numeric_cols, key='viz_column' )
					fig = px.histogram( df, x=col )
					st.plotly_chart( fig, use_container_width=True )
		
		# ----------------------------------------------------------------------
		# ADMIN TAB
		# ----------------------------------------------------------------------
		with tabs[ 7 ]:
			tables = list_tables( )
			if tables:
				table = st.selectbox( 'Table', tables, key='admin_table' )
			
			st.divider( )
			
			st.subheader( 'AI Asset Governance' )
			
			if st.button( 'Refresh AI Asset Counts', width='stretch' ):
				st.session_state[ 'dm_asset_counts' ] = get_ai_asset_counts( )
			
			asset_counts = st.session_state.get( 'dm_asset_counts', { } )
			if asset_counts:
				ac1, ac2, ac3, ac4 = st.columns( [ 0.25, 0.25, 0.25, 0.25 ] )
				with ac1:
					st.metric( 'Documents', int( asset_counts.get( 'documents', 0 ) ) )
				with ac2:
					st.metric( 'Document Chunks',
						int( asset_counts.get( 'document_chunks', 0 ) ) )
				with ac3:
					st.metric( 'Document Embeddings',
						int( asset_counts.get( 'document_embeddings', 0 ) ) )
				with ac4:
					st.metric( 'Images', int( asset_counts.get( 'images', 0 ) ) )
			
			asset_admin_c1, asset_admin_c2 = st.columns( [ 0.5, 0.5 ], border=True )
			
			with asset_admin_c1:
				if st.button( 'Rebuild Active Document Asset Rows', width='stretch' ):
					doc_result = register_session_documents( )
					chunk_result = register_session_chunks( )
					embed_result = register_session_embeddings( )
					
					st.success(
						f'Documents inserted: {doc_result[ "inserted" ]}, '
						f'updated: {doc_result[ "updated" ]}, '
						f'chunks inserted: {chunk_result[ "inserted" ]}, '
						f'embeddings inserted: {embed_result[ "inserted" ]}' )
			
			with asset_admin_c2:
				if st.button( 'Purge Orphaned AI Assets', width='stretch' ):
					purge_result = purge_orphaned_ai_assets( )
					st.success( f'Deleted chunks: {purge_result[ "deleted_chunks" ]}, '
						f'deleted embeddings: {purge_result[ "deleted_embeddings" ]}' )
			
			st.markdown( cfg.BLUE_DIVIDER, unsafe_allow_html=True )
			
			st.subheader( 'Data Profiling' )
			tables = list_tables( )
			if tables:
				table = st.selectbox( 'Select Table', tables, key='profile_table' )
				if st.button( 'Generate Profile' ):
					profile_df = create_profile_table( table )
					st.dataframe( profile_df, use_container_width=True )
			
			st.subheader( 'Drop Table' )
			
			tables = list_tables( )
			if tables:
				table = st.selectbox( 'Select Table to Drop', tables, key='admin_drop_table' )
				
				if 'dm_confirm_drop' not in st.session_state:
					st.session_state.dm_confirm_drop = False
				
				if st.button( 'Drop Table', key='admin_drop_button' ):
					st.session_state.dm_confirm_drop = True
				
				if st.session_state.dm_confirm_drop:
					st.warning( f'You are about to permanently delete table {table}. '
						'This action cannot be undone.' )
					
					col1, col2 = st.columns( 2 )
					
					if col1.button( 'Confirm Drop', key='admin_confirm_drop' ):
						try:
							drop_table( table )
							st.success( f'Table {table} dropped successfully.' )
						except Exception as e:
							st.error( f'Drop failed: {e}' )
						
						st.session_state.dm_confirm_drop = False
						st.rerun( )
					
					if col2.button( 'Cancel', key='admin_cancel_drop' ):
						st.session_state.dm_confirm_drop = False
						st.rerun( )
				
				df = read_table( table )
				col = st.selectbox( 'Create Index On', df.columns )
				
				if st.button( 'Create Index' ):
					create_index( table, col )
					st.success( 'Index created.' )
			
			st.divider( )
			
			st.subheader( 'Create Custom Table' )
			new_table_name = st.text_input( 'Table Name' )
			column_count = st.number_input( 'Number of Columns', min_value=1, max_value=20,
				value=1 )
			columns = [ ]
			for i in range( column_count ):
				st.markdown( f'### Column {i + 1}' )
				col_name = st.text_input( 'Column Name', key=f'col_name_{i}' )
				col_type = st.selectbox( 'Column Type', [ 'INTEGER', 'REAL', 'TEXT' ],
					key=f'col_type_{i}' )
				
				not_null = st.checkbox( 'NOT NULL', key=f'not_null_{i}' )
				primary_key = st.checkbox( 'PRIMARY KEY', key=f'pk_{i}' )
				auto_inc = st.checkbox( 'AUTOINCREMENT (INTEGER only)', key=f'ai_{i}' )
				
				columns.append( {
							'name': col_name,
							'type': col_type,
							'not_null': not_null,
							'primary_key': primary_key,
							'auto_increment': auto_inc
					} )
			
			if st.button( 'Create Table' ):
				try:
					create_custom_table( new_table_name, columns )
					st.success( 'Table created successfully.' )
					st.rerun( )
				except Exception as e:
					st.error( f'Error: {e}' )
			
			st.divider( )
			st.subheader( 'Schema Viewer' )
			
			tables = list_tables( )
			if tables:
				table = st.selectbox( 'Select Table', tables, key='schema_view_table' )
				schema = create_schema( table )
				schema_df = pd.DataFrame( schema,
					columns=[ 'cid', 'name', 'type', 'notnull', 'default', 'pk' ] )
				
				st.markdown( '### Columns' )
				st.dataframe( schema_df, use_container_width=True )
				with create_connection( ) as conn:
					count = conn.execute( f'SELECT COUNT(*) FROM "{table}"' ).fetchone( )[ 0 ]
				
				st.metric( 'Row Count', f'{count:,}' )
				indexes = get_indexes( table )
				if indexes:
					idx_df = pd.DataFrame( indexes,
						columns=[ 'seq', 'name', 'unique', 'origin', 'partial' ] )
					st.markdown( '### Indexes' )
					st.dataframe( idx_df, use_container_width=True )
				else:
					st.info( 'No indexes defined.' )
			
			st.divider( )
			st.subheader( 'ALTER TABLE Operations' )
			tables = list_tables( )
			if tables:
				table = st.selectbox( 'Select Table', tables, key='alter_table_select' )
				operation = st.selectbox( 'Operation',
					[ 'Add Column', 'Rename Column', 'Rename Table', 'Drop Column' ] )
				
				if operation == 'Add Column':
					new_col = st.text_input( 'Column Name' )
					col_type = st.selectbox( 'Column Type', [ 'INTEGER', 'REAL', 'TEXT' ] )
					
					if st.button( 'Add Column' ):
						add_column( table, new_col, col_type )
						st.success( 'Column added.' )
						st.rerun( )
				
				elif operation == 'Rename Column':
					schema = create_schema( table )
					col_names = [ col[ 1 ] for col in schema ]
					old_col = st.selectbox( 'Column to Rename', col_names )
					new_col = st.text_input( 'New Column Name' )
					
					if st.button( 'Rename Column' ):
						rename_column( table, old_col, new_col )
						st.success( 'Column renamed.' )
						st.rerun( )
				
				elif operation == 'Rename Table':
					new_name = st.text_input( 'New Table Name' )
					
					if st.button( 'Rename Table' ):
						rename_table( table, new_name )
						st.success( 'Table renamed.' )
						st.rerun( )
				
				elif operation == 'Drop Column':
					schema = create_schema( table )
					col_names = [ col[ 1 ] for col in schema ]
					drop_col = st.selectbox( 'Column to Drop', col_names )
					
					if st.button( 'Drop Column' ):
						drop_column( table, drop_col )
						st.success( 'Column dropped.' )
						st.rerun( )
		
		# ----------------------------------------------------------------------
		# SQL TAB
		# ----------------------------------------------------------------------
		with tabs[ 8 ]:
			st.subheader( 'SQL Console' )
			query = st.text_area( 'Enter SQL Query' )
			if st.button( 'Run Query' ):
				if not is_safe_query( query ):
					st.error( 'Query blocked: Only read-only SELECT statements are allowed.' )
				else:
					try:
						start_time = time.perf_counter( )
						with create_connection( ) as conn:
							result = pd.read_sql_query( query, conn )
						
						end_time = time.perf_counter( )
						elapsed = end_time - start_time
						st.dataframe( result, use_container_width=True )
						row_count = len( result )
						col1, col2 = st.columns( 2 )
						col1.metric( 'Rows Returned', f'{row_count:,}' )
						col2.metric( 'Execution Time (seconds)', f'{elapsed:.6f}' )
						
						if elapsed > 2.0:
							st.warning( 'Slow query detected (> 2 seconds). Consider indexing.' )
						
						if not result.empty:
							csv = result.to_csv( index=False ).encode( 'utf-8' )
							st.download_button( 'Download CSV', csv, 'query_results.csv',
								'text/csv' )
					
					except Exception as e:
						st.error( f'Execution failed: {e}' )

# ==============================================================================
# FOOTER — SECTION
# ==============================================================================
st.markdown(
	"""
	<style>
	.block-container {
		padding-bottom: 3rem;
	}
	</style>
	""",
	unsafe_allow_html=True,
)

# ---- Fixed Container
st.markdown(
	"""
	<style>
	.boo-status-bar {
		position: fixed;
		bottom: 0;
		left: 0;
		width: 100%;
		background-color: rgba(17, 17, 17, 0.95);
		border-top: 1px solid #2a2a2a;
		padding: 10px 16px;
		font-size: 0.80rem;
		color: #5292f7;
		z-index: 1000;
	}
	.boo-status-inner {
		display: flex;
		justify-content: space-between;
		align-items: center;
		max-width: 100%;
	}
	</style>
	""", unsafe_allow_html=True,)

# ======================================================================================
# FOOTER RENDERING
# ======================================================================================

right_parts: List[ str ] = [ ]
model = str( st.session_state.get( 'selected_model_name',
		get_default_model_name( ) ) or get_default_model_name( ) )

mode_value = mode if mode is not None else st.session_state.get( 'mode' )
if mode_value:
	right_parts.append( f'Mode: {mode_value}' )
	
selected_model_path = str( st.session_state.get( 'selected_model_path', '' ) or '' )
if selected_model_path and Path( selected_model_path ).exists( ):
	right_parts.append( 'Model File: Available' )
elif selected_model_path:
	right_parts.append( 'Model File: Missing' )
else:
	right_parts.append( 'Model File: Not Configured' )
	
temperature = st.session_state.get( 'temperature' )
top_p = st.session_state.get( 'top_percent' )
top_k = st.session_state.get( 'top_k' )
frequency = st.session_state.get( 'frequency_penalty' )
presense = st.session_state.get( 'presense_penalty' )
repeat_penalty = st.session_state.get( 'repeat_penalty' )
max_tokens = st.session_state.get( 'max_tokens' )
context_window = st.session_state.get( 'context_window' )
cpu_threads = st.session_state.get( 'cpu_threads' )
repeat_window = st.session_state.get( 'repeat_window' )
use_semantic = st.session_state.get( 'use_semantic' )
basic_docs = st.session_state.get( 'basic_docs' )

# ------------------------------------------------------------------
# Parameter summary (show 0 values; suppress only when None)
# ------------------------------------------------------------------
if temperature is not None:
	right_parts.append( f'Temp: {float( temperature ):0.2f}' )

if top_p is not None:
	right_parts.append( f'Top-P: {float( top_p ):0.2f}' )

if top_k is not None:
	right_parts.append( f'Top-K: {int( top_k )}' )

if frequency is not None:
	right_parts.append( f'Freq: {float( frequency ):0.2f}' )

if presense is not None:
	right_parts.append( f'Presence: {float( presense ):0.2f}' )

if repeat_penalty is not None:
	right_parts.append( f'Repeat: {float( repeat_penalty ):0.2f}' )

if repeat_window is not None:
	right_parts.append( f'Repeat Window: {int( repeat_window )}' )

if max_tokens is not None:
	right_parts.append( f'Max Tokens: {int( max_tokens )}' )

if context_window is not None:
	right_parts.append( f'Context: {int( context_window )}' )

if cpu_threads is not None:
	right_parts.append( f'Threads: {int( cpu_threads )}' )

# ------------------------------------------------------------------
# Context flags (optional but useful)
# ------------------------------------------------------------------
if use_semantic is not None:
	right_parts.append( f'Semantic: {"On" if use_semantic else "Off"}' )

if isinstance( basic_docs, list ):
	right_parts.append( f'Docs: {len( basic_docs )}' )

right_text = ' ◽ '.join( right_parts ) if right_parts else '—'

# ---- Rendering Method
st.markdown(
	f"""
    <div class="boo-status-bar">
        <div class="boo-status-inner">
            <span>{model}</span>
            <span>{right_text}</span>
        </div>
    </div>
    """,
	unsafe_allow_html=True, )